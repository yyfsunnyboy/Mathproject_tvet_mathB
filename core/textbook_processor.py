# -*- coding: utf-8 -*-
"""
=============================================================================
模組名稱 (Module Name): core/textbook_processor.py
功能說明 (Description): 教科書內容分析與 AI 結構化匯入模組。將 PDF 與 Word 檔案轉為
                      標準多項式與 LaTeX 題庫結構，並利用 Gemini LLM 進行技能單元
                      切分與資料庫入庫作業。
使用說明 (Usage): 由後台管理端路由或批次腳本調用。
版本資訊 (Version): V2.0
更新日期 (Date): 2026-01-13
維護團隊 (Maintainer): Math AI Project Team
=============================================================================
"""
"""
教科書內容分析與 AI 結構化匯入工具 (Textbook Processor & AI Analyzer) - Final Complete Version

本模組的核心任務為讀取各種格式之教材（PDF 與 Word 檔案），精確切分出內文中的觀念
引導、例題、隨堂練習與章節習題，並透過 Google Gemini LLM 進行自動化語意分析，
將切分出的數學題目與對應的技能單元進行關聯，最終結構化寫入系統資料庫。

版本修訂履歷：
1. 完整修復並維護原有系統之例外處理、公式 assets 自動對齊與防禦型回填機制 (Restore full logic)。
2. 整合針對 Word/Pandoc 導出文本之特定字元清洗器 (clean_pandoc_output)。
3. 更新針對技高數學、高職數B 特化之 Prompt 引導樣板與驗證流程。
"""

import json
import re
import os
import hashlib
import zipfile
import xml.etree.ElementTree as ET
import uuid
import platform
# import fitz  # PyMuPDF -> Moved to inside function
import time
import io
from typing import Any
# import pypandoc -> Moved to inside function
# from pypandoc.pandoc_download import download_pandoc
from google.api_core.exceptions import ResourceExhausted
from models import db, SkillInfo, SkillCurriculum, TextbookExample
from core.ai_analyzer import get_model
from flask import current_app, has_app_context
import traceback
from core.code_generator import auto_generate_skill_code
from core.math_formula_normalizer import (
    detect_suspicious_formula,
    normalize_converted_docx_latex_text,
    normalize_math_text,
)
from core.math_expression_formatter import standardize_problem_latex
from core.question_image_assets import (
    attach_image_metadata,
    build_question_asset_filename,
    build_question_asset_dir,
    build_question_assets_dir,
    build_question_code,
    convert_vector_image_to_png,
    detect_image_reason,
    find_best_page_index,
    infer_source_page_for_question,
    make_page_image_asset,
    question_needs_image,
    render_pdf_page_to_image,
)
from core.textbook_filename_parser import (
    parse_textbook_filename_metadata,
    resolve_parse_filename_for_import,
)
from core.textbook_structure_parser import get_structure_map
from core.utils import normalize_vocational_math_skill_id

_DOCX_IMPORT_CONTEXT: dict[str, Any] = {}

FORMULA_PLACEHOLDER_RE = re.compile(r"\[FORMULA_IMAGE_\d+\]|\[FORMULA_MISSING\]|\[WORD_EQUATION_UNPARSED\]")
TEXT_MOJIBAKE_CHARS = "嚗嚙踐□■◆＊"
TEXT_MOJIBAKE_RE = re.compile("[" + re.escape(TEXT_MOJIBAKE_CHARS) + r"]")
LATEX_SIGNAL_GUARD_RE = re.compile(r"\\\(|\\\)|\\\[|\\\]|\\(?:frac|sqrt|left|right)\b|[\^_]")
GARBLE_MARKERS = ("\ufffd", "銝", "憒", "瘜", "\uf010", "", "嚗", "踐", "□", "■", "◆", "＊")
PERM_COMB_NOTATION_RE = re.compile(
    r"(?:\bP\s*\(|\bC\s*\(|\bP\s*\^|\bC\s*\^|\bP_|\bC_)",
    re.IGNORECASE,
)
LATEX_DEBUG_SIGNAL_RE = re.compile(
    r"\\(?:frac|sqrt|left|right|times|leq|geq|binom)\b|\$[^$]+\$"
)


def _safe_regex_search(pattern: str, text: str, flags: int = 0) -> bool:
    try:
        return re.search(pattern, str(text or ""), flags) is not None
    except re.error:
        try:
            if has_app_context():
                current_app.logger.warning(f"[REGEX GUARD] invalid regex skipped: {pattern!r}")
        except Exception:
            pass
        return False


def _norm_title_spaces(title: str) -> str:
    return re.sub(r"\s+", " ", str(title or "").strip())


def _strip_title_brackets(s: str) -> str:
    return s.strip("〔〕[]()（）")


_CH_SA_CH_MARKER_RE = re.compile(r"CH\s*(\d+)\s*自我評量", re.IGNORECASE)
_CH_SA_ZH_CHAPTER_RE = re.compile(r"第\s*(\d+)\s*章")
_CH_SA_SECTION_HEADING_RE = re.compile(r"^\s*(\d+-\d+)\s+(.+)$")
_CH_SA_QUESTION_LINE_RE = re.compile(r"^\s*(\d{1,2})(?:[\.、\)\t]|\s+)(.+)")
_CH_SA_PAGE_ONLY_RE = re.compile(r"^\s*\d{2,3}\s*$")


def detect_chapter_self_assessment_context(extracted_text: str) -> dict[str, Any] | None:
    """Detect CH{n}自我評量 / 章末題庫 layout with multiple section headings (1-1, 1-2, …)."""
    text = str(extracted_text or "")
    if not text.strip():
        return None
    has_ch_marker = bool(_CH_SA_CH_MARKER_RE.search(text))
    has_sa_label = bool(re.search(r"(?m)^\s*自我評量\s*$", text)) or bool(
        re.search(r"自我評量", text)
    )
    if not (has_ch_marker or has_sa_label):
        return None

    section_codes: list[str] = []
    for ln in text.splitlines():
        line = str(ln or "").strip()
        if not line or "習題" in line:
            continue
        sec_m = _CH_SA_SECTION_HEADING_RE.match(line)
        if sec_m:
            section_codes.append(sec_m.group(1))

    unique_sections = sorted(
        set(section_codes),
        key=lambda x: tuple(int(p) for p in str(x).split("-")),
    )
    if len(unique_sections) < 2:
        return None

    ch_num: int | None = None
    m_ch = _CH_SA_CH_MARKER_RE.search(text)
    if m_ch:
        ch_num = int(m_ch.group(1))
    else:
        m_zh = _CH_SA_ZH_CHAPTER_RE.search(text)
        if m_zh:
            ch_num = int(m_zh.group(1))
    if ch_num is None:
        ch_num = 1

    title_prefix = (
        f"CH{ch_num}自我評量" if has_ch_marker else f"第{ch_num}章自我評量"
    )
    return {
        "mode": "chapter_self_assessment",
        "chapter_num": ch_num,
        "title_prefix": title_prefix,
        "section_codes": unique_sections,
    }


def _scan_chapter_self_assessment_inventory(
    lines: list[str],
    ctx: dict[str, Any],
) -> list[dict[str, Any]]:
    prefix = str(ctx.get("title_prefix") or "CH1自我評量")
    items: list[dict[str, Any]] = []
    current_section_code = ""
    current_section_title = ""
    started = False

    for ln in lines:
        line = str(ln or "").strip()
        if not line or _CH_SA_PAGE_ONLY_RE.match(line):
            continue
        if _CH_SA_CH_MARKER_RE.search(line) or line == "自我評量":
            started = True
            continue
        if _CH_SA_ZH_CHAPTER_RE.match(line) and "習題" not in line:
            continue
        sec_m = _CH_SA_SECTION_HEADING_RE.match(line)
        if sec_m and "習題" not in line and not _CH_SA_QUESTION_LINE_RE.match(line):
            current_section_code = sec_m.group(1)
            current_section_title = line
            started = True
            continue
        if not started:
            continue
        qm = _CH_SA_QUESTION_LINE_RE.match(line)
        if not qm:
            continue
        n = int(qm.group(1))
        if n < 1 or n > 99:
            continue
        canon = f"{prefix} 題{n}"
        items.append(
            {
                "raw_title": qm.group(0).strip()[:48],
                "canonical_title": canon,
                "kind": "self_assessment",
                "section_code": current_section_code,
                "section_title": current_section_title,
                "exercise_block": prefix,
                "zone": "",
                "number": str(n),
                "source_span_preview": str(qm.group(2) or "").strip()[:120],
            }
        )
    return items


def _scan_chapter_self_assessment_blocks(
    lines: list[str],
    ctx: dict[str, Any],
) -> dict[str, str]:
    prefix = str(ctx.get("title_prefix") or "CH1自我評量")
    blocks: dict[str, str] = {}
    current_key: str | None = None
    buffer: list[str] = []
    started = False

    def flush() -> None:
        nonlocal current_key, buffer
        _scan_flush_question_block(blocks, current_key, buffer)
        current_key = None
        buffer = []

    def start_question(num: int, first_line: str) -> None:
        nonlocal current_key, buffer
        flush()
        current_key = f"{prefix} 題{num}"
        buffer = [first_line]

    for raw in lines:
        line = str(raw or "").strip()
        if not line:
            if current_key:
                buffer.append("")
            continue
        if _CH_SA_PAGE_ONLY_RE.match(line):
            continue
        if _SCAN_KEY_LINE_RE.match(line):
            flush()
            continue
        if _CH_SA_CH_MARKER_RE.search(line) or line == "自我評量":
            started = True
            continue
        if _CH_SA_ZH_CHAPTER_RE.match(line) and "習題" not in line:
            continue
        sec_m = _CH_SA_SECTION_HEADING_RE.match(line)
        if sec_m and "習題" not in line and not _CH_SA_QUESTION_LINE_RE.match(line):
            flush()
            started = True
            continue
        if not started:
            continue
        qm = _CH_SA_QUESTION_LINE_RE.match(line)
        if qm:
            start_question(int(qm.group(1)), line)
            continue
        if current_key:
            if _CH_SA_QUESTION_LINE_RE.match(line) or (
                _CH_SA_SECTION_HEADING_RE.match(line)
                and "習題" not in line
                and not _CH_SA_QUESTION_LINE_RE.match(line)
            ):
                flush()
                qm2 = _CH_SA_QUESTION_LINE_RE.match(line)
                if qm2:
                    start_question(int(qm2.group(1)), line)
                continue
            buffer.append(line)
            continue

    flush()
    return blocks


def scan_docx_title_inventory(extracted_text: str, section_code: str | None = None) -> list[dict[str, Any]]:
    """Deterministic scan of example / practice / chapter exercise / exam titles from DOCX text."""
    text = str(extracted_text or "")
    lines = text.splitlines()
    sa_ctx = detect_chapter_self_assessment_context(text)
    if sa_ctx:
        return _scan_chapter_self_assessment_inventory(lines, sa_ctx)

    items: list[dict[str, Any]] = []

    def _infer_section_from_line(line: str) -> str | None:
        m = re.search(r"(\d+-\d+)", line)
        return m.group(1) if m else None

    inferred_global = None
    for ln in lines:
        g = _infer_section_from_line(ln)
        if g:
            inferred_global = g
            break

    effective_section = (str(section_code).strip() if section_code else "") or (inferred_global or "")

    # --- 例題 (document-wide) ---
    for m in re.finditer(r"例(?:題)?\s*(\d{1,2})\b", text):
        n = int(m.group(1))
        raw = m.group(0).strip()
        preview = text[max(0, m.start() - 20) : m.end() + 40].replace("\n", " ")
        items.append(
            {
                "raw_title": raw,
                "canonical_title": f"例題{n}",
                "kind": "example",
                "section_code": effective_section or "",
                "exercise_block": "",
                "zone": "",
                "number": str(n),
                "source_span_preview": preview[:120],
            }
        )

    # ==============================================================================
    # 【Regex 補強】相容隨堂練習後方帶有連續點點點（Dot leaders）或空格之排版
    # ==============================================================================
    for m in re.finditer(r"隨堂練習[\s\.…·]*(\d{1,2})\b", text):
        n = int(m.group(1))
        raw = m.group(0).strip()
        preview = text[max(0, m.start() - 12) : m.end() + 30].replace("\n", " ")
        items.append(
            {
                "raw_title": raw,
                "canonical_title": f"隨堂練習{n}",
                "kind": "in_class_practice",
                "section_code": effective_section or "",
                "exercise_block": "",
                "zone": "",
                "number": str(n),
                "source_span_preview": preview[:120],
            }
        )
    # ==============================================================================
    suitang_m = re.search(r"隨堂練習(?=\s*$)", text, flags=re.MULTILINE)
    if not suitang_m:
        suitang_m = re.search(r"(?m)^\s*隨堂練習\s*$", text)
    if suitang_m:
        tail = text[suitang_m.end() :]
        end_m = re.search(r"\d+-\d+習題|[〔\[]?\s*\d{2,3}\s*統測", tail)
        region = tail[: end_m.start()] if end_m else tail
        for lm in re.finditer(r"(?m)^\s*(\d{1,2})(?:[\.、\)\t]|\s+)", region):
            n = int(lm.group(1))
            raw = lm.group(0).strip()
            preview = region[max(0, lm.start() - 10) : lm.end() + 40].replace("\n", " ")
            items.append(
                {
                    "raw_title": raw,
                    "canonical_title": f"隨堂練習{n}",
                    "kind": "in_class_practice",
                    "section_code": effective_section or "",
                    "exercise_block": "",
                    "zone": "",
                    "number": str(n),
                    "source_span_preview": preview[:120],
                }
            )

    # --- 章節習題區（zone 完全依原文掃描，不依題號推測）---
    ZONE_HEADERS = ("基礎題", "進階題", "自我評量")

    i = 0
    while i < len(lines):
        line = lines[i]
        blk = re.search(r"(\d+-\d+)習題", line)
        if not blk:
            i += 1
            continue
        sec = blk.group(1)
        exercise_block = f"{sec}習題"
        current_zone = "其他"
        i += 1
        while i < len(lines):
            ln = lines[i]
            hdr = re.match(r"^\s*(\d+-\d+)習題", ln)
            if hdr and hdr.group(1) != sec:
                break
            if re.search(r"[〔\[]?\s*\d{2,3}\s*統測", ln, flags=re.IGNORECASE):
                break
            stripped = ln.strip()
            zone_hit = None
            for z in ZONE_HEADERS:
                if stripped == z or stripped.startswith(z + " ") or stripped.startswith(z + "　"):
                    zone_hit = z
                    break
            if zone_hit:
                current_zone = zone_hit
                i += 1
                continue
            mnum = re.match(r"^\s*(\d{1,2})(?:[\.、\)\t]|\s+)", ln)
            if mnum:
                n = int(mnum.group(1))
                raw = mnum.group(0).strip()
                preview = ln[:100]
                canon = f"{sec}習題 {current_zone}{n}"
                items.append(
                    {
                        "raw_title": raw,
                        "canonical_title": canon,
                        "kind": "chapter_exercise",
                        "section_code": sec,
                        "exercise_block": exercise_block,
                        "zone": current_zone,
                        "number": str(n),
                        "source_span_preview": preview,
                    }
                )
            i += 1
        continue

    # --- 統測 ---
    for m in re.finditer(r"[〔\[]?\s*(\d{2,3})\s*統測\s*([AB])\s*[〕\]]?", text, flags=re.IGNORECASE):
        y = int(m.group(1))
        suf = m.group(2).upper()
        raw = m.group(0).strip()
        preview = text[max(0, m.start() - 10) : m.end() + 30].replace("\n", " ")
        items.append(
            {
                "raw_title": raw,
                "canonical_title": f"{y}統測{suf}",
                "kind": "exam_practice",
                "section_code": effective_section or "",
                "exercise_block": "",
                "zone": "",
                "number": "",
                "source_span_preview": preview[:120],
            }
        )

    return items


CONVERTED_DOCX_LATEX_JSON_RULES = (
    "【converted_docx_latex metadata-only 模式 — JSON 硬性規則】\n"
    "1. 本模式題幹由 DOCX 決定性掃描補回；Gemini 只輸出章節結構與題目 metadata。\n"
    "2. 所有 JSON 字串禁止實際換行；需換行時僅能寫 \\\\n。\n"
    "3. 反斜線必須 JSON 雙重跳脫（例如 LaTeX 的 \\\\frac 在 JSON 內寫成 \\\\\\\\frac）。\n"
    "4. problem_text 僅能填與 title 相同的短字串或極短摘要（單行、≤40 字），禁止多行長題幹。\n"
    "5. correct_answer 與 detailed_solution 請填 \"\"；禁止輸出 null。\n"
    "6. 每一題仍須包含：title、source_description、source_type、concept_name、concept_en_id、"
    "problem_text、correct_answer、detailed_solution。\n"
    "7. title 與 source_description 必須相同，且與題目標題清單一致。\n"
)


def scan_converted_docx_question_blocks(
    extracted_text: str,
    *,
    source_scope: str = "",
) -> dict[str, str]:
    """Deterministic per-title question blocks from converted LaTeX DOCX plain text."""
    paragraph_blocks = [
        {"type": "paragraph", "text": str(ln or "").strip()}
        for ln in str(extracted_text or "").splitlines()
        if str(ln or "").strip()
    ]
    return build_docx_question_formula_context(paragraph_blocks, source_scope=source_scope)


def build_converted_docx_latex_gemini_outline_payload(
    extracted_text: str,
    section_code: str | None = None,
) -> str:
    """Compact Gemini input: title inventory + section hints only (no full LaTeX stems)."""
    items = scan_docx_title_inventory(extracted_text, section_code=section_code)
    sa_ctx = detect_chapter_self_assessment_context(extracted_text)
    parts = [
        "【converted_docx_latex metadata-only — 題目標題清單（請逐題輸出對應 metadata，勿重寫題幹）】",
    ]
    if sa_ctx:
        parts.append(
            f"[CHAPTER_SELF_ASSESSMENT] prefix={sa_ctx.get('title_prefix')} "
            f"sections={','.join(sa_ctx.get('section_codes') or [])}"
        )
        parts.append(
            "每題放入 examples 陣列、source_type=self_assessment；"
            "title 與 source_description 必須等於 canonical_title（如 CH1自我評量 題1）；"
            "practice_questions 留空；problem_text 僅填 title。"
        )
    for it in items:
        parts.append(
            f"- canonical_title={it.get('canonical_title')} kind={it.get('kind')} "
            f"section={it.get('section_code')} zone={it.get('zone')}"
        )
    seen_sections: set[str] = set()
    for m in re.finditer(r"(\d+-\d+)\s+([^\n]{2,48})", str(extracted_text or "")):
        hint = f"{m.group(1)} {m.group(2).strip()}"
        if hint not in seen_sections:
            seen_sections.add(hint)
            parts.append(f"[SECTION_HINT] {hint[:72]}")
    if not items:
        parts.append("(no titles detected — still output JSON skeleton with empty arrays)")
    return "\n".join(parts)


_HYDRATE_QUESTION_CUE_RE = re.compile(
    r"求|解|試求|計算|判斷|選|何者|下列|若|已知|設|問|算|證明|化簡|作圖|求出|比較"
)
_HYDRATE_LATEX_SIGNAL_RE = re.compile(
    r"\$|\\\(|\\\[|\\frac|\\sqrt|\^|_|\\left|\\right",
    re.IGNORECASE,
)
_INVALID_HYDRATED_BLOCK_PATTERNS = (
    re.compile(r"^第\s*\d+\s*章.*$"),
    re.compile(r"^\d+\s*-\s*\d+(?:\s+\S{1,30})?$"),
    re.compile(r"^\d+\s*-\s*\d+\s*習題$"),
    re.compile(r"^(基礎題|進階題|自我評量|綜合題|歷屆試題|統測歷屆試題)$"),
)
_METADATA_ONLY_HYDRATED_BLOCK_RE = re.compile(
    r"^(?:例(?:題)?|隨堂練習|習題)\d+$|^\d{2,3}統測[AB]$",
    re.IGNORECASE,
)


def _append_item_parse_warning(item: dict, tag: str) -> None:
    tag = str(tag or "").strip()
    if not tag:
        return
    existing = str(item.get("parse_warning") or "").strip()
    parts = [p.strip() for p in existing.split(";") if p.strip()] if existing else []
    if tag not in parts:
        parts.append(tag)
    item["parse_warning"] = ";".join(parts)


def _is_invalid_hydrated_block(block: str, title: str = "", section_code: str = "") -> bool:
    """Return True when a DOCX block is a heading/metadata shell, not a question stem."""
    b = str(block or "").strip()
    if not b:
        return True

    b_norm = b.replace("　", " ").strip()
    b_compact = re.sub(r"\s+", "", b_norm)
    t_compact = re.sub(r"\s+", "", str(title or "").strip())
    if t_compact and b_compact == t_compact:
        return True

    sc_compact = re.sub(r"\s+", "", str(section_code or "").strip())
    if sc_compact and b_compact in {sc_compact, f"{sc_compact}習題"}:
        return True

    for pat in _INVALID_HYDRATED_BLOCK_PATTERNS:
        if pat.match(b_norm):
            return True

    if _METADATA_ONLY_HYDRATED_BLOCK_RE.match(b_compact):
        return True

    has_cue = bool(_HYDRATE_QUESTION_CUE_RE.search(b_norm))
    has_latex = bool(_HYDRATE_LATEX_SIGNAL_RE.search(b_norm))
    if len(b_norm) < 12 and not has_cue and not has_latex:
        return True

    if not has_cue and not has_latex and len(b_norm) < 24:
        if re.match(r"^[\d\-\s習題基礎進階自我評量綜合歷屆統測學測例隨堂]+$", b_compact):
            return True

    return False


_DIAGRAM_REQUIRED_RE = re.compile(
    r"如圖|下圖|上圖|右圖|左圖|依圖|根據圖|由圖可知|圖中|圖形中|如圖所示|"
    r"座標圖|函數圖形|數線上|陰影部分|幾何圖形|統計圖|長條圖|折線圖|圓形圖|"
    r"填入圖中|判斷圖形|觀察圖|如右圖|如左圖|如下圖|如上圖|附圖|"
    r"線段\s*AB\s*如圖|著色|路線圖|棋盤式街道圖|樹狀圖|圓環|示意圖|題圖"
)
_DIAGRAM_DEPENDENCY_PHRASE_RE = re.compile(
    r"如圖|下圖|上圖|右圖|左圖|依圖|根據圖|由圖可知|圖中|圖形中|如圖所示|"
    r"如右圖|如左圖|如下圖|如上圖|附圖"
)
_LIFE_CONTEXT_RE = re.compile(
    r"咖啡車|咖啡|商店|商品|成本|收入|利潤|票價|租金|車資|購買|販賣|製作|販售|售價|進貨"
)
_MATH_SELF_SUFFICIENT_RE = re.compile(
    r"成本|收入|利潤|方程式|不等式|試問|已知|"
    r"\\frac|\\sqrt|\$[^$\n]+\$|[><≥≤＝=]|\d+\s*元|"
    r"設\s*[^，。]{0,30}(?:為|是)|最少|至少|最多"
)


def classify_image_dependency(
    question_text: str,
    nearby_caption: str = "",
    image_hint: str = "",
) -> dict[str, Any]:
    """Classify whether a question stem requires an image asset vs decorative nearby art."""
    q = str(question_text or "").strip()
    cap = str(nearby_caption or "").strip()
    hint = str(image_hint or "").strip()
    combined = "\n".join(p for p in (q, cap, hint) if p)

    if _DIAGRAM_REQUIRED_RE.search(combined):
        return {
            "needs_image": True,
            "image_role": "required_diagram",
            "reason": "diagram_reference_or_chart_type_in_stem",
        }

    if _LIFE_CONTEXT_RE.search(q) and _MATH_SELF_SUFFICIENT_RE.search(q):
        return {
            "needs_image": False,
            "image_role": "decorative_context",
            "reason": "word_problem_with_complete_math_conditions",
        }

    if re.search(r"圖", combined) and not _DIAGRAM_DEPENDENCY_PHRASE_RE.search(combined):
        if _MATH_SELF_SUFFICIENT_RE.search(q):
            return {
                "needs_image": False,
                "image_role": "decorative_context",
                "reason": "figure_word_without_diagram_dependency_and_self_contained_stem",
            }
        return {
            "needs_image": False,
            "image_role": "unknown",
            "reason": "figure_mention_without_clear_diagram_dependency",
        }

    if question_needs_image(q, ai_has_image=False):
        return {
            "needs_image": True,
            "image_role": "required_diagram",
            "reason": "legacy_image_keyword_heuristic",
        }

    return {
        "needs_image": False,
        "image_role": "unknown",
        "reason": "no_diagram_dependency_detected",
    }


def _image_dependency_notes_metadata(dep: dict[str, Any]) -> dict[str, Any] | None:
    role = str(dep.get("image_role") or "")
    if role in ("decorative_context", "unknown") and not dep.get("needs_image"):
        return {
            "image_role": role,
            "image_dependency_reason": str(dep.get("reason") or ""),
            "needs_image": False,
        }
    return None


def hydrate_converted_docx_latex_parsed_data(
    parsed_data: dict,
    *,
    extracted_text: str,
    section_code: str | None = None,
    inventory_items: list[dict[str, Any]] | None = None,
    question_blocks: dict[str, str] | None = None,
) -> tuple[dict, int, int]:
    """Fill problem_text from DOCX blocks; normalize empty answer/solution fields."""
    blocks = dict(question_blocks or scan_converted_docx_question_blocks(extracted_text))
    filled = 0
    if not isinstance(parsed_data, dict):
        return parsed_data, 0, len(blocks)

    for chapter in parsed_data.get("chapters", []) or []:
        if not isinstance(chapter, dict):
            continue
        for section in chapter.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            sec_title = str(section.get("section_title", "") or "")
            sec_code = section_code
            if not sec_code:
                m = re.search(r"(\d+-\d+)", sec_title)
                sec_code = m.group(1) if m else None
            for concept in section.get("concepts", []) or []:
                if not isinstance(concept, dict):
                    continue
                for bucket in ("examples", "practice_questions"):
                    for item in concept.get(bucket, []) or []:
                        if not isinstance(item, dict):
                            continue
                        title = get_question_title(item) or ""
                        canon = canonicalize_import_title(
                            title,
                            section_code=sec_code,
                            inventory_items=inventory_items,
                        )
                        canon_title = str(canon or title or "").strip()
                        block = _lookup_docx_formula_block(
                            canon_title,
                            blocks,
                            allow_section_fallback=False,
                            allow_bare_number=False,
                        ) or _lookup_docx_formula_block(
                            title,
                            blocks,
                            allow_section_fallback=False,
                            allow_bare_number=False,
                        )
                        if title:
                            item["title"] = title
                            item["source_description"] = str(item.get("source_description") or title).strip()
                        stub_text = str(title or item.get("problem_text") or "").strip()
                        label = canon_title or title
                        if block:
                            if _is_invalid_hydrated_block(
                                block,
                                title=label,
                                section_code=str(sec_code or ""),
                            ):
                                item["problem_text"] = stub_text
                                item["needs_review"] = True
                                item["hydrate_missing_block"] = True
                                item["hydrate_invalid_block"] = True
                                _append_item_parse_warning(item, "invalid_hydrated_block")
                                current_app.logger.warning(
                                    "[DOCX HYDRATE WARNING] title=%s reason=invalid_hydrated_block block=%s",
                                    label,
                                    str(block)[:80],
                                )
                            else:
                                item["problem_text"] = block
                                filled += 1
                        else:
                            item["problem_text"] = stub_text
                            item["needs_review"] = True
                            item["hydrate_missing_block"] = True
                            _append_item_parse_warning(item, "missing_hydrated_block")
                        if item.get("correct_answer") is None:
                            item["correct_answer"] = ""
                        else:
                            item["correct_answer"] = str(item.get("correct_answer") or "")
                        if item.get("detailed_solution") is None:
                            item["detailed_solution"] = ""
                        else:
                            item["detailed_solution"] = str(item.get("detailed_solution") or "")
    return parsed_data, filled, len(blocks)


def map_returned_import_title(
    title: str,
    *,
    section_code: str | None = None,
    inventory_items: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Map AI-returned title string to canonical form; never guess exercise zone from number alone."""
    raw = _norm_title_spaces(str(title or ""))
    inv = list(inventory_items or [])
    s_compact = re.sub(r"\s+", "", raw)
    s_compact = _strip_title_brackets(s_compact)

    m = re.match(r"^例(?:題)?(\d+)$", s_compact)
    if m:
        return {
            "returned_raw": raw,
            "returned_canonical": f"例題{int(m.group(1))}",
            "mapping_method": "direct_example",
            "needs_review": False,
        }

    m = re.match(r"^隨堂練習(\d+)$", s_compact)
    if m:
        return {
            "returned_raw": raw,
            "returned_canonical": f"隨堂練習{int(m.group(1))}",
            "mapping_method": "direct_practice",
            "needs_review": False,
        }

    m = re.match(r"^(\d{2,3})統測([AB])$", s_compact, flags=re.IGNORECASE)
    if m:
        return {
            "returned_raw": raw,
            "returned_canonical": f"{int(m.group(1))}統測{m.group(2).upper()}",
            "mapping_method": "direct_exam",
            "needs_review": False,
        }

    merged_exam_bucket_names = frozenset({
        "統測歷屆試題",
        "歷屆試題",
        "統測題",
        "統測試題",
        "學測歷屆試題",
    })
    if s_compact in merged_exam_bucket_names:
        exam_in_inv = [
            it for it in inv
            if str(it.get("kind", "")) == "exam_practice"
            or "統測" in str(it.get("canonical_title", ""))
        ]
        if exam_in_inv:
            return {
                "returned_raw": raw,
                "returned_canonical": raw,
                "mapping_method": "merged_exam_bucket_rejected",
                "needs_review": True,
            }

    sc = str(section_code or "").strip()
    if sc and sc != "unknown":
        zone_label_map = (
            ("基礎題", "基礎題"),
            ("進階題", "進階題"),
            ("自我評量", "自我評量"),
        )
        for zone_key, zone_label in zone_label_map:
            m_zone = re.match(rf"^{re.escape(zone_key)}(\d+)$", s_compact)
            if m_zone:
                num = int(m_zone.group(1))
                return {
                    "returned_raw": raw,
                    "returned_canonical": f"{sc}習題 {zone_label}{num}",
                    "mapping_method": "section_zone_labeled_exercise",
                    "needs_review": False,
                }

    m = re.match(r"^習題(\d+)$", s_compact)
    if m:
        num = m.group(1)
        cands = [it for it in inv if it.get("kind") == "chapter_exercise" and str(it.get("number", "")) == num]
        if len(cands) == 1:
            return {
                "returned_raw": raw,
                "returned_canonical": str(cands[0].get("canonical_title", "")),
                "mapping_method": "exercise_context_map",
                "needs_review": False,
            }
        sc = str(section_code or "").strip() or ""
        fallback = f"{sc}習題 題{num}" if sc else f"習題 題{num}"
        return {
            "returned_raw": raw,
            "returned_canonical": fallback,
            "mapping_method": "fallback_unresolved",
            "needs_review": True,
        }

    m = re.match(r"^(\d+-\d+)習題(基礎題|進階題|自我評量|其他)(\d+)$", s_compact)
    if m:
        return {
            "returned_raw": raw,
            "returned_canonical": f"{m.group(1)}習題 {m.group(2)}{int(m.group(3))}",
            "mapping_method": "exercise_context_map",
            "needs_review": False,
        }

    m_sa_prefix = re.match(r"^(CH\d+自我評量|第\d+章自我評量)題(\d+)$", s_compact, flags=re.IGNORECASE)
    if m_sa_prefix:
        num = m_sa_prefix.group(2)
        cands = [
            it
            for it in inv
            if str(it.get("kind", "")) == "self_assessment" and str(it.get("number", "")) == num
        ]
        if len(cands) == 1:
            return {
                "returned_raw": raw,
                "returned_canonical": str(cands[0].get("canonical_title", "")),
                "mapping_method": "chapter_self_assessment_map",
                "needs_review": False,
            }
        return {
            "returned_raw": raw,
            "returned_canonical": f"{m_sa_prefix.group(1)} 題{int(num)}",
            "mapping_method": "chapter_self_assessment_direct",
            "needs_review": False,
        }

    m_sa_bare = re.match(r"^題(\d+)$", s_compact)
    if m_sa_bare:
        num = m_sa_bare.group(1)
        cands = [
            it
            for it in inv
            if str(it.get("kind", "")) == "self_assessment" and str(it.get("number", "")) == num
        ]
        if len(cands) == 1:
            return {
                "returned_raw": raw,
                "returned_canonical": str(cands[0].get("canonical_title", "")),
                "mapping_method": "chapter_self_assessment_map",
                "needs_review": False,
            }

    return {
        "returned_raw": raw,
        "returned_canonical": raw,
        "mapping_method": "fallback_unresolved",
        "needs_review": True,
    }


def canonicalize_import_title(
    title: str,
    section_code: str | None = None,
    inventory_items: list[dict[str, Any]] | None = None,
    *,
    exercise_zone_map: dict[int, str] | None = None,
) -> str:
    """Normalize title for inventory; optional inventory_items disambiguates 「習題 n」."""
    meta = map_returned_import_title(
        title,
        section_code=section_code,
        inventory_items=inventory_items,
    )
    if exercise_zone_map and meta.get("mapping_method") == "fallback_unresolved":
        s_compact = re.sub(r"\s+", "", _norm_title_spaces(title))
        s_compact = _strip_title_brackets(s_compact)
        m = re.match(r"^習題(\d+)$", s_compact)
        if m and section_code:
            num = int(m.group(1))
            zone = exercise_zone_map.get(num)
            if zone:
                return f"{section_code}習題 {zone}{num}"
    return str(meta.get("returned_canonical") or "").strip() or str(title or "").strip()


def is_official_vh_skill_id(skill_id: str) -> bool:
    """True when skill_id is a formal student-facing vocational skill (vh_, not outline)."""
    sid = str(skill_id or "").strip()
    return bool(sid.startswith("vh_") and not sid.startswith("outline_"))


def is_outline_skill_id(skill_id: str) -> bool:
    return str(skill_id or "").strip().startswith("outline_")


def extract_section_code_from_title(section_title: str) -> str:
    m = re.search(r"(\d+-\d+)", str(section_title or ""))
    return m.group(1) if m else ""


def _norm_section_label(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip())


def _fetch_skill_curriculum_rows_for_section(
    curriculum: str,
    volume: str,
    section_code: str,
) -> list[Any]:
    """Load SkillCurriculum rows for same 課綱/冊 + section_code prefix."""
    curr = str(curriculum or "").strip()
    vol = str(volume or "").strip()
    code = str(section_code or "").strip()
    if not curr or not vol or not code:
        return []
    try:
        q = SkillCurriculum.query.filter(
            SkillCurriculum.curriculum == curr,
            SkillCurriculum.volume == vol,
        )
        rows = list(q.filter(SkillCurriculum.section.like(f"{code} %")).all())
        for row in q.filter(SkillCurriculum.section == code).all():
            rid = getattr(row, "id", None)
            if rid is None or rid not in {getattr(r, "id", None) for r in rows}:
                rows.append(row)
    except Exception:
        return []
    seen: set[int] = set()
    out: list[Any] = []
    for row in rows or []:
        rid = getattr(row, "id", None)
        if rid is not None and rid in seen:
            continue
        if rid is not None:
            seen.add(rid)
        out.append(row)
    return out


def _pick_anchor_curriculum_row(
    rows: list[Any],
    *,
    preferred_section_title: str,
    section_code: str,
) -> tuple[Any | None, str]:
    """Pick section anchor (main) skill, not a fine-grained sub-skill row."""
    pref = _norm_section_label(preferred_section_title)
    code = str(section_code or "").strip()
    official = [
        r
        for r in (rows or [])
        if is_official_vh_skill_id(str(getattr(r, "skill_id", "") or ""))
    ]
    if not official:
        return None, ""

    if pref:
        exact = [r for r in official if _norm_section_label(getattr(r, "section", "")) == pref]
        if exact:
            exact.sort(
                key=lambda r: (
                    int(getattr(r, "display_order", 0) or 0),
                    int(getattr(r, "id", 0) or 0),
                )
            )
            return exact[0], "exact_section_title"

    scoped: list[Any] = []
    for row in official:
        sec = _norm_section_label(getattr(row, "section", ""))
        if code and (sec == code or sec.startswith(f"{code} ")):
            scoped.append(row)
    if not scoped:
        return None, ""

    def _sort_key(row: Any) -> tuple:
        sec = _norm_section_label(getattr(row, "section", ""))
        exact_rank = 0 if pref and sec == pref else 1
        return (
            exact_rank,
            len(sec),
            int(getattr(row, "display_order", 0) or 0),
            int(getattr(row, "id", 0) or 0),
        )

    scoped.sort(key=_sort_key)
    reason = "exact_section_title" if pref and _norm_section_label(getattr(scoped[0], "section", "")) == pref else "section_code_anchor"
    return scoped[0], reason


def _pick_outline_curriculum_row(rows: list[Any], section_code: str) -> Any | None:
    code = str(section_code or "").strip()
    outline = [
        r
        for r in (rows or [])
        if is_outline_skill_id(str(getattr(r, "skill_id", "") or ""))
    ]
    if not outline:
        return None
    outline.sort(
        key=lambda r: (
            len(_norm_section_label(getattr(r, "section", ""))),
            int(getattr(r, "display_order", 0) or 0),
            int(getattr(r, "id", 0) or 0),
        )
    )
    return outline[0]


def resolve_question_bank_section_skill_id(
    curriculum: str,
    volume: str,
    section_code: str,
    section_title: str = "",
) -> tuple[str | None, str]:
    """Resolve anchor skill_id for chapter question-bank / self-assessment by section."""
    code = str(section_code or "").strip() or extract_section_code_from_title(section_title)
    anchor_title = _norm_section_label(section_title) or code
    if not code:
        return None, "missing_section_code"

    curr = str(curriculum or "").strip()
    vol = str(volume or "").strip()
    if anchor_title:
        try:
            exact_rows = (
                SkillCurriculum.query.filter(
                    SkillCurriculum.curriculum == curr,
                    SkillCurriculum.volume == vol,
                    SkillCurriculum.section == anchor_title,
                    SkillCurriculum.skill_id.like("vh_%"),
                    ~SkillCurriculum.skill_id.like("outline_%"),
                )
                .order_by(SkillCurriculum.display_order.asc(), SkillCurriculum.id.asc())
                .all()
            )
        except Exception:
            exact_rows = []
        if exact_rows:
            sid = str(getattr(exact_rows[0], "skill_id", "") or "").strip()
            if has_app_context():
                current_app.logger.info(
                    "[QUESTION BANK SKILL ALIGN] section_code=%s section_title=%s "
                    "selected_skill_id=%s source=exact_section_title",
                    code,
                    anchor_title,
                    sid,
                )
            return sid, "exact_section_title"

    rows = _fetch_skill_curriculum_rows_for_section(curriculum, volume, code)
    row, reason = _pick_anchor_curriculum_row(
        rows,
        preferred_section_title=anchor_title,
        section_code=code,
    )
    if row is not None:
        sid = str(getattr(row, "skill_id", "") or "").strip()
        if has_app_context():
            current_app.logger.info(
                "[QUESTION BANK SKILL ALIGN] section_code=%s section_title=%s "
                "selected_skill_id=%s source=%s",
                code,
                anchor_title,
                sid,
                reason,
            )
        return sid, reason

    outline_row = _pick_outline_curriculum_row(rows, code)
    if outline_row is not None:
        sid = str(getattr(outline_row, "skill_id", "") or "").strip()
        if has_app_context():
            current_app.logger.warning(
                "[QUESTION BANK SKILL ALIGN WARNING] section_code=%s section_title=%s "
                "reason=fallback_outline fallback_outline_skill_id=%s",
                code,
                anchor_title,
                sid,
            )
        return sid, "fallback_outline"

    if has_app_context():
        current_app.logger.warning(
            "[QUESTION BANK SKILL ALIGN WARNING] section_code=%s section_title=%s "
            "reason=fallback_outline fallback_outline_skill_id=",
            code,
            anchor_title,
        )
    return None, "not_found"


def resolve_question_bank_skill_for_section(
    curriculum: str,
    volume: str,
    section_title: str,
    *,
    section_code: str | None = None,
) -> dict[str, Any]:
    """Align chapter question-bank / self-assessment items to official SkillCurriculum skill_id."""
    code = str(section_code or "").strip() or extract_section_code_from_title(section_title)
    sid, source = resolve_question_bank_section_skill_id(
        curriculum,
        volume,
        code,
        section_title,
    )
    return {
        "skill_id": sid or "",
        "needs_review": source in ("fallback_outline", "not_found", "missing_section_code"),
        "source": source,
        "section_code": code,
    }


def is_question_bank_import_scope(
    docx_ctx: dict[str, Any] | None,
    filename_meta: dict[str, Any] | None = None,
) -> bool:
    """True only for chapter question-bank / self-assessment imports (never section_textbook)."""
    meta = filename_meta if isinstance(filename_meta, dict) else {}
    ctx = docx_ctx if isinstance(docx_ctx, dict) else {}
    scope = str(meta.get("source_scope") or ctx.get("source_scope") or "").strip()
    if scope == "section_textbook":
        return False
    if scope in ("chapter_question_bank", "chapter_self_assessment"):
        return True
    if scope == "chapter_review":
        return True
    return bool(ctx.get("chapter_self_assessment_mode")) and scope != "section_textbook"


def is_section_textbook_import_scope(
    docx_ctx: dict[str, Any] | None,
    filename_meta: dict[str, Any] | None = None,
) -> bool:
    if is_question_bank_import_scope(docx_ctx, filename_meta):
        return False
    scope = str((filename_meta or {}).get("source_scope") or "section_textbook").strip()
    return scope == "section_textbook"


def section_textbook_allowed_needs_review(item: dict[str, Any] | None) -> bool:
    """section_textbook: only hydrate/image/parse failures may require review."""
    if not isinstance(item, dict):
        return False
    if item.get("hydrate_missing_block") or item.get("hydrate_invalid_block"):
        return True
    if str(item.get("parse_warning") or "").strip():
        return True
    if str(item.get("image_warning") or "").strip() == "missing_docx_image_asset":
        return True
    if item.get("missing_docx_image_asset"):
        return True
    repair = item.get("repair_log")
    if isinstance(repair, list):
        for entry in repair:
            if "missing_docx_image_asset" in str(entry):
                return True
    elif repair and "missing_docx_image_asset" in str(repair):
        return True
    return False


def resolve_section_textbook_skill_id(
    curriculum: str,
    volume: str,
    section_code: str,
    section_title: str = "",
    *,
    existing_anchor_section: Any = None,
    structure_meta: dict[str, Any] | None = None,
) -> tuple[str | None, str]:
    """Resolve stable section anchor skill for section_textbook imports (no outline fallback)."""
    code = str(section_code or "").strip() or extract_section_code_from_title(section_title)
    anchor_title = _norm_section_label(section_title)
    if not anchor_title and isinstance(structure_meta, dict):
        anchor_title = _norm_section_label(str(structure_meta.get("section_title") or ""))
    if not anchor_title and code:
        anchor_title = code

    curr = str(curriculum or "").strip()
    vol = str(volume or "").strip()
    if anchor_title:
        try:
            exact_rows = (
                SkillCurriculum.query.filter(
                    SkillCurriculum.curriculum == curr,
                    SkillCurriculum.volume == vol,
                    SkillCurriculum.section == anchor_title,
                    SkillCurriculum.skill_id.like("vh_%"),
                    ~SkillCurriculum.skill_id.like("outline_%"),
                )
                .order_by(SkillCurriculum.display_order.asc(), SkillCurriculum.id.asc())
                .all()
            )
        except Exception:
            exact_rows = []
        if exact_rows:
            sid = str(getattr(exact_rows[0], "skill_id", "") or "").strip()
            if has_app_context():
                current_app.logger.info(
                    "[SECTION TEXTBOOK SKILL ALIGN] section_code=%s section_title=%s "
                    "selected_skill_id=%s source=section_anchor",
                    code,
                    anchor_title,
                    sid,
                )
            return sid, "section_anchor"

    if existing_anchor_section is not None:
        sid = str(getattr(existing_anchor_section, "skill_id", "") or "").strip()
        if is_official_vh_skill_id(sid):
            sec_label = _norm_section_label(getattr(existing_anchor_section, "section", ""))
            if not anchor_title or sec_label == anchor_title:
                if has_app_context():
                    current_app.logger.info(
                        "[SECTION TEXTBOOK SKILL ALIGN] section_code=%s section_title=%s "
                        "selected_skill_id=%s source=section_anchor",
                        code,
                        anchor_title,
                        sid,
                    )
                return sid, "section_anchor"

    if code:
        rows = _fetch_skill_curriculum_rows_for_section(curriculum, volume, code)
        row, reason = _pick_anchor_curriculum_row(
            rows,
            preferred_section_title=anchor_title,
            section_code=code,
        )
        if row is not None:
            sid = str(getattr(row, "skill_id", "") or "").strip()
            if is_official_vh_skill_id(sid):
                if has_app_context():
                    current_app.logger.info(
                        "[SECTION TEXTBOOK SKILL ALIGN] section_code=%s section_title=%s "
                        "selected_skill_id=%s source=section_anchor",
                        code,
                        anchor_title,
                        sid,
                    )
                return sid, reason or "section_anchor"

    if has_app_context():
        current_app.logger.warning(
            "[SECTION TEXTBOOK SKILL ALIGN WARNING] section_code=%s section_title=%s "
            "reason=no_official_skill_found",
            code,
            anchor_title,
        )
    return None, "not_found"


def inventory_section_title_for_code(
    inventory_items: list[dict[str, Any]] | None,
    section_code: str,
) -> str:
    """Best section_title from DOCX inventory for a section_code (e.g. 1-2 平面坐標…)."""
    code = str(section_code or "").strip()
    best = ""
    for it in inventory_items or []:
        if str(it.get("section_code", "")).strip() != code:
            continue
        st = _norm_section_label(str(it.get("section_title", "") or ""))
        if len(st) > len(best):
            best = st
    return best


def scan_expected_titles_from_converted_text(extracted_text: str) -> list[str]:
    items = scan_docx_title_inventory(extracted_text)
    return sorted({str(it.get("canonical_title", "")).strip() for it in items if str(it.get("canonical_title", "")).strip()})


def collect_returned_titles_from_parsed_data(parsed_data: dict[str, Any]) -> list[str]:
    titles: list[str] = []
    for ch in (parsed_data or {}).get("chapters", []) or []:
        for sec in (ch or {}).get("sections", []) or []:
            for concept in (sec or {}).get("concepts", []) or []:
                for ex in (concept or {}).get("examples", []) or []:
                    t = ex.get("title") or ex.get("source_description") or ""
                    if str(t).strip():
                        titles.append(str(t).strip())
                for pq in (concept or {}).get("practice_questions", []) or []:
                    t = pq.get("title") or pq.get("source_description") or ""
                    if str(t).strip():
                        titles.append(str(t).strip())
    return titles


def build_title_inventory(
    expected_titles: list[str],
    returned_titles: list[str],
    section_code: str = "",
    *,
    inventory_items: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    inv_list = list(inventory_items or [])
    sc = str(section_code or "").strip()
    if sc == "unknown":
        sc = ""

    expected_raw = [str(t).strip() for t in expected_titles if str(t).strip()]
    expected_norm = sorted(
        {
            canonicalize_import_title(t, section_code=sc or None, inventory_items=inv_list)
            for t in expected_raw
        }
    )

    returned_raw = [str(t).strip() for t in returned_titles if str(t).strip()]
    returned_mappings = [
        map_returned_import_title(t, section_code=sc or None, inventory_items=inv_list) for t in returned_raw
    ]
    returned_norm = sorted(
        {str(m.get("returned_canonical") or "").strip() for m in returned_mappings if str(m.get("returned_canonical") or "").strip()}
    )

    expected_set = set(expected_norm)
    returned_set = set(returned_norm)
    missing = sorted(expected_set - returned_set)
    extra = sorted(returned_set - expected_set)
    return {
        "expected_titles_raw": expected_raw,
        "expected_titles_canonical": expected_norm,
        "returned_titles_raw": returned_raw,
        "returned_titles_canonical": returned_norm,
        "missing_titles_canonical": missing,
        "extra_titles_canonical": extra,
        "expected_titles_count": len(expected_norm),
        "returned_titles_count": len(returned_norm),
        "missing_titles_count": len(missing),
        "extra_titles_count": len(extra),
        "inventory_items": inv_list,
        "returned_title_mappings": returned_mappings,
    }


def _md_table_row(cells: list[str]) -> str:
    esc = [str(c or "").replace("|", "\\|").replace("\n", " ") for c in cells]
    return "| " + " | ".join(esc) + " |"


def write_title_inventory_report(report_path: str, *, volume: str, section: str, allow_partial_import: bool, write_aborted: bool, inv: dict[str, Any], warning: str = "") -> None:
    lines = [
        "# Title Inventory Report",
        "",
        "## Summary",
        f"- volume: `{volume}`",
        f"- section: `{section}`",
        f"- expected_titles_count: `{inv.get('expected_titles_count', 0)}`",
        f"- returned_titles_count: `{inv.get('returned_titles_count', 0)}`",
        f"- missing_titles_count: `{inv.get('missing_titles_count', 0)}`",
        f"- extra_titles_count: `{inv.get('extra_titles_count', 0)}`",
        f"- allow_partial_import: `{str(bool(allow_partial_import)).lower()}`",
        f"- write_aborted: `{str(bool(write_aborted)).lower()}`",
        f"- warning: `{warning}`",
        "",
    ]
    inv_items = inv.get("inventory_items") or []
    if inv_items:
        lines.append("## Inventory Items")
        lines.append(_md_table_row(["raw_title", "canonical_title", "kind", "section_code", "exercise_block", "zone", "number", "preview"]))
        lines.append(_md_table_row(["---"] * 8))
        for it in inv_items:
            lines.append(
                _md_table_row(
                    [
                        str(it.get("raw_title", "")),
                        str(it.get("canonical_title", "")),
                        str(it.get("kind", "")),
                        str(it.get("section_code", "")),
                        str(it.get("exercise_block", "")),
                        str(it.get("zone", "")),
                        str(it.get("number", "")),
                        str(it.get("source_span_preview", ""))[:200],
                    ]
                )
            )
        lines.append("")

    mappings = inv.get("returned_title_mappings") or []
    if mappings:
        lines.append("## Returned Title Mapping")
        lines.append(_md_table_row(["returned_raw", "returned_canonical", "mapping_method", "needs_review"]))
        lines.append(_md_table_row(["---"] * 4))
        for m in mappings:
            lines.append(
                _md_table_row(
                    [
                        str(m.get("returned_raw", "")),
                        str(m.get("returned_canonical", "")),
                        str(m.get("mapping_method", "")),
                        str(m.get("needs_review", False)).lower(),
                    ]
                )
            )
        lines.append("")

    lines.extend(
        [
            "## Expected Titles Raw",
        ]
    )
    for t in inv.get("expected_titles_raw", []):
        lines.append(f"- {t}")
    lines.append("")
    lines.append("## Expected Titles Canonical")
    for t in inv.get("expected_titles_canonical", []):
        lines.append(f"- {t}")
    lines.append("")
    lines.append("## Returned Titles Raw")
    for t in inv.get("returned_titles_raw", []):
        lines.append(f"- {t}")
    lines.append("")
    lines.append("## Returned Titles Canonical")
    for t in inv.get("returned_titles_canonical", []):
        lines.append(f"- {t}")
    lines.append("")
    lines.append("## Missing Titles Canonical")
    for t in inv.get("missing_titles_canonical", []):
        lines.append(f"- {t}")
    lines.append("")
    lines.append("## Extra Titles Canonical")
    for t in inv.get("extra_titles_canonical", []):
        lines.append(f"- {t}")
    out = os.path.abspath(report_path)
    os.makedirs(os.path.dirname(out), exist_ok=True)
    with open(out, "w", encoding="utf-8-sig") as f:
        f.write("\n".join(lines) + "\n")


def detect_curriculum_volume_warning(curriculum: str, volume: str, publisher: str = "") -> str:
    c = str(curriculum or "").strip().lower()
    v = str(volume or "").strip()
    p = str(publisher or "").strip().lower()
    if re.search(r"數學\s*B[1-4]\b", v, flags=re.IGNORECASE) and c == "general" and (not p or "longteng" in p):
        return f"[IMPORT WARNING] volume={v} but curriculum=general; vocational mathB import expected."
    return ""
def _has_text_mojibake(text: str) -> bool:
    t = str(text or "")
    if not t:
        return False
    if TEXT_MOJIBAKE_RE.search(t):
        return True
    if t.count("?") >= 3 and (t.count("?") / max(1, len(t))) > 0.04:
        return True
    return any(0xE000 <= ord(ch) <= 0xF8FF for ch in t)


def _has_garbled_text(text: str) -> bool:
    t = str(text or "")
    if not t:
        return False
    if any(marker in t for marker in GARBLE_MARKERS):
        return True
    return _has_text_mojibake(t)


def _is_low_value_import_field(value: str) -> bool:
    t = str(value or "").strip()
    if not t:
        return True
    return t in {"?", "??", "N/A", "n/a", "None", "none", "null", "-"}


def score_problem_text_quality(text) -> dict:
    """Score imported problem text so duplicate merge does not replace good LaTeX with placeholders."""
    t = str(text or "").strip()
    placeholder_count = len(FORMULA_PLACEHOLDER_RE.findall(t))
    formula_image_count = len(re.findall(r"\[FORMULA_IMAGE_\d+\]", t))
    formula_missing_count = t.count("[FORMULA_MISSING]")
    mojibake = _has_text_mojibake(t)
    latex_patterns = [
        r"\\\(.+?\\\)",
        r"\$.+?\$",
        r"\\leq?|\\geq?",
        r"\\frac",
        r"\\sqrt",
        r"\|x\|",
        r"\|[^|]+\|",
    ]
    latex_signal_count = sum(1 for pat in latex_patterns if re.search(pat, t))
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", t))
    only_placeholder = bool(t) and not re.sub(FORMULA_PLACEHOLDER_RE, "", t).strip()
    too_short = len(t) < 6

    score = 50
    score += latex_signal_count * 18
    if placeholder_count == 0:
        score += 20
    score += min(cjk_count, 20)
    if re.search(r"[？：；]", t):
        score += 4
    score -= formula_image_count * 35
    score -= formula_missing_count * 45
    if mojibake:
        score -= 60
    if too_short:
        score -= 20
    if only_placeholder:
        score -= 60

    return {
        "score": int(score),
        "placeholder_count": placeholder_count,
        "formula_image_count": formula_image_count,
        "formula_missing_count": formula_missing_count,
        "latex_signal_count": latex_signal_count,
        "mojibake_detected": mojibake,
        "too_short": too_short,
        "only_placeholder": only_placeholder,
        "length": len(t),
    }


def should_replace_problem_text(existing_text, incoming_text) -> tuple[bool, dict, dict]:
    existing_quality = score_problem_text_quality(existing_text)
    incoming_quality = score_problem_text_quality(incoming_text)
    return incoming_quality["score"] > existing_quality["score"], existing_quality, incoming_quality

# (此暫存區已刪除)

# ==============================================================================
# [修復] 處理並修正 LaTeX 的逸出字元
# ==============================================================================
def sanitize_gemini_json_text(raw: str) -> str:
    r"""
    靽桀儔 Gemini ? JSON 銝剖虜閬??澆?????

    瘜冽?嚗?
    ?撘??json.loads ????raw text??
    摰??桃??芣霈?raw text ??? JSON??
    json.loads ??敺?Python 摮葡銝剔? LaTeX ?府隞?舀迤撣詨??蝺?
    靘?鞈?摨急?敺?閰脣???\(x+1\)嚗???\\(x+1\\)??
    """
    if raw is None:
        return raw

    text = str(raw).strip()

    # 蝘駁 Markdown code fence
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)

    # ?瑕??憭惜 JSON object嚗?芋??敺?隤芣???
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]

    # 處理內部的 JSON escape
    # 這是 JSON escape 的處理：
    # \" \\ \/ \b \f \n \r \t \uXXXX
    # 避免 LaTeX 中的反斜線干擾 JSON 的解析
    # 這些反斜線在 raw JSON 中需要被雙重逸出，避免 json.loads 失敗
    # 標準的 LaTeX 寫法如果被當作 JSON 逸出，最好統一做替換
    # (靘? \binom, \frac, \times)嚗??json.loads ????嗅??憯?MathJax??
    latex_commands = (
        "binom|frac|times|cdot|sum|prod|sqrt|left|right|over|overline|underline|"
        "vec|hat|bar|lim|to|infty|sin|cos|tan|cot|sec|csc|log|ln|"
        "alpha|beta|gamma|delta|theta|lambda|mu|pi|sigma|omega|phi|rho|tau|"
        "Delta|Sigma"
    )
    text = re.sub(rf'(?<!\\)\\(?=(?:{latex_commands})\b|[()\[\]{{}}])', r'\\\\', text)

    text = re.sub(r'(?<!\\)\\(?!["\\/bfnrtu])', r'\\\\', text)

    return text


def safe_load_gemini_json(raw: str):
    r"""
    摰閫?? Gemini ? JSON??
    ???json.loads??
    ?亙仃??靽桀儔 LaTeX escape 敺? json.loads??
    """
    raw_text = str(raw).strip() if raw is not None else raw
    fixed = sanitize_gemini_json_text(raw)

    try:
        parsed = json.loads(raw)
        if fixed != raw_text:
            try:
                sanitized_parsed = json.loads(fixed)
                try:
                    current_app.logger.info("[TEXTBOOK IMPORTER] Gemini JSON parsed after LaTeX escape sanitize.")
                except RuntimeError:
                    pass
                return sanitized_parsed
            except json.JSONDecodeError:
                return parsed
        return parsed
    except json.JSONDecodeError as first_error:
        try:
            current_app.logger.debug(
                "[TEXTBOOK IMPORTER] Gemini first json.loads failed at "
                f"line {first_error.lineno}, col {first_error.colno}, pos {first_error.pos}: "
                f"{first_error.msg}"
            )
        except RuntimeError:
            pass
        try:
            parsed = json.loads(fixed)
            try:
                current_app.logger.info("[TEXTBOOK IMPORTER] Gemini JSON parsed after LaTeX escape sanitize.")
            except RuntimeError:
                pass
            return parsed
        except json.JSONDecodeError as second_error:
            try:
                _log_gemini_json_parse_failed_after_sanitize(first_error, second_error, raw)
            except RuntimeError:
                pass
            preview = str(raw)[:800] if raw is not None else ""
            raise ValueError(
                "Gemini JSON parse failed after sanitize. "
                f"First error: {first_error}. "
                f"Second error: {second_error}. "
                f"Raw preview: {preview}"
            ) from second_error


def _log_gemini_json_parse_failed_after_sanitize(first_error, second_error, raw):
    preview = str(raw)[:800] if raw is not None else ""
    current_app.logger.error(
        "[TEXTBOOK IMPORTER] Gemini JSON parse failed after sanitize. "
        f"First error: {first_error}. "
        f"Second error: {second_error}. "
        f"Raw preview: {preview}"
    )


def clean_pandoc_output(text):
    """
    清理 Word 轉出後 LaTeX 的額外修整。
    """
    if not text: return text

    # 1. 修正 Pandoc 特有的度數符號重複 (^{\^{\circ}} -> ^{\circ})
    text = text.replace(r'^{\^{\circ}}', r'^{\circ}')
    
    # 2. 將 \( ... \) 替換為 $ ... $
    text = re.sub(r"\\\((.*?)\\\)", r"$\1$", text)

    # 3. 修正 sqrt 格式
    text = re.sub(r'(?:\\)?sqrt\s+(\d+|[a-zA-Z])\b', r'\\sqrt{\1}', text)
    
    return text


def _xml_local_name(tag: str) -> str:
    if not tag:
        return ""
    return tag.split("}", 1)[-1] if "}" in tag else tag


def extract_docx_paragraph_with_equations(paragraph) -> str:
    """提取段落文字"""
    return str(paragraph.text or "").strip()


def extract_docx_table_with_equations(table) -> str:
    """提取表格文字"""
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            segs = [str(p.text or "").strip() for p in cell.paragraphs]
            cells.append(" ".join(filter(None, segs)).strip())
        lines.append(" | ".join(cells).strip())
    return "\n".join(lines).strip()


def extract_converted_latex_docx(file_path: str) -> tuple[dict[int, str], dict[str, Any]]:
    """Extract DOCX text in document order for pre-converted MathType->LaTeX DOCX."""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(file_path)
    ordered_chunks: list[str] = []
    paragraph_count = 0
    table_count = 0

    for block in doc.element.body.iterchildren():
        if block.tag.endswith("}p"):
            para = Paragraph(block, doc)
            text = str(para.text or "").strip()
            paragraph_count += 1
            if text:
                ordered_chunks.append(text)
        elif block.tag.endswith("}tbl"):
            table_count += 1
            tbl = Table(block, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = str(p.text or "").strip()
                        if text:
                            ordered_chunks.append(text)
    merged = "\n".join(ordered_chunks).strip()
    return {1: merged} if merged else {}, {
        "paragraph_count": paragraph_count,
        "table_count": table_count,
    }


def detect_converted_latex_docx(text: str) -> dict[str, Any]:
    t = str(text or "")
    latex_signal_count = 0
    latex_signal_count += len(re.findall(r"\$[^$\n]+\$", t))
    latex_signal_count += len(re.findall(r"\\\([^)\n]+\\\)", t))
    latex_signal_count += len(re.findall(r"\\\[[^\]\n]+\\\]", t))
    latex_signal_count += len(re.findall(r"\\(?:frac|sqrt|le|ge|binom|times|pm)\b", t))
    placeholder_count = len(re.findall(r"\[FORMULA_IMAGE_\d+\]|\[FORMULA_MISSING\]", t))
    is_converted = latex_signal_count >= 3 and placeholder_count <= 1
    return {
        "is_converted_latex_docx": bool(is_converted),
        "latex_signal_count": int(latex_signal_count),
        "formula_placeholder_count": int(placeholder_count),
    }


def build_docx_media_relationship_map(docx_path: str, extracted_media_dir: str) -> dict[str, dict[str, str]]:
    rel_map: dict[str, dict[str, str]] = {}
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            rel_xml = zf.read("word/_rels/document.xml.rels")
        rel_root = ET.fromstring(rel_xml)
        for rel in rel_root.findall(".//{*}Relationship"):
            rid = rel.attrib.get("Id")
            target = rel.attrib.get("Target", "")
            rtype = rel.attrib.get("Type", "")
            if not rid or "image" not in rtype.lower():
                continue
            filename = os.path.basename(target)
            extracted_path = os.path.join(extracted_media_dir, filename).replace("\\", "/")
            rel_map[rid] = {
                "target_ref": target,
                "content_type": _guess_image_content_type(filename),
                "extracted_path": extracted_path,
            }
    except Exception:
        return {}
    return rel_map


def _guess_image_content_type(filename: str) -> str:
    ext = os.path.splitext(str(filename or ""))[1].lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
        ".wmf": "image/x-wmf",
        ".emf": "image/x-emf",
    }.get(ext, "application/octet-stream")


def _docx_image_original_format(path_or_name: str, content_type: str = "") -> str:
    ext = os.path.splitext(str(path_or_name or ""))[1].lower().lstrip(".")
    if ext in ("wmf", "emf", "png", "jpg", "jpeg"):
        return "jpeg" if ext == "jpg" else ext
    ctype = str(content_type or "").lower()
    if "wmf" in ctype:
        return "wmf"
    if "emf" in ctype:
        return "emf"
    if "png" in ctype:
        return "png"
    if "jpeg" in ctype or "jpg" in ctype:
        return "jpeg"
    return "unknown"


def extract_docx_image_rids_from_paragraph(paragraph) -> list[str]:
    rids = []
    p_el = paragraph._p
    for blip in p_el.findall(".//{*}blip"):
        rid = blip.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rid:
            rids.append(rid)
    for imagedata in p_el.findall(".//{*}imagedata"):
        rid = imagedata.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        if rid:
            rids.append(rid)
    return rids


def _is_question_start_text(text: str) -> bool:
    t = str(text or "").strip()
    if not t:
        return False
    if classify_non_question_block(t) in ("concept_explanation", "figure_caption", "narration"):
        return False
    heading_patterns = [
        r"^\s*例題\s*\d+",
        r"^\s*隨堂練習\s*\d+",
        r"^\s*基礎題\s*\d+",
        r"^\s*進階題\s*\d+",
        r"^\s*(?:\d+\s*-\s*\d+\s*)?習題(?:\s*(?:基礎題|進階題))?\s*\d*",
        r"^\s*自我評量",
        r"^\s*(?:統測|學測)\s*\d*",
        r"^\s*挑戰\s*\d+",
        r"^\s*\d+[\s\.\)]",
    ]
    return any(re.search(p, t) for p in heading_patterns)


_STRUCTURAL_BOUNDARY_PATTERNS = [
    r"^\s*第\s*\d+\s*章",
    r"^\s*\d+\s*[^\d\s].*$",
    r"^\s*\d+\s*-\s*\d+\s+[^\s].*$",
    r"^\s*\d+\s*-\s*\d+\s*\.\s*\d+\s*[^\s].*$",
    r"^\s*例(?:題)?\s*\d+",
    r"^\s*隨堂練習\s*\d+",
    r"^\s*(?:\d+\s*-\s*\d+\s*)?習題",
    r"^\s*基礎題\s*\d*",
    r"^\s*進階題\s*\d*",
    r"^\s*自我評量",
    r"^\s*(?:統測|學測)",
]


def is_structural_boundary_line(line: str) -> bool:
    t = str(line or "").strip()
    if not t:
        return False
    return any(re.search(p, t) for p in _STRUCTURAL_BOUNDARY_PATTERNS)


_NON_QUESTION_EXPLANATION_CUES = (
    "說明",
    "觀念",
    "範例解析",
    "解題",
    "補充",
    "備註",
    "提示",
    "概念",
    "答案",
    "解答",
    "附註",
)

_QUESTION_VERBS = (
    "求",
    "解",
    "算",
    "證明",
    "判斷",
    "比較",
    "化簡",
    "計算",
    "作圖",
    "試求",
    "求出",
)


def _is_figure_caption_line(text: str) -> bool:
    t = str(text or "").strip()
    return bool(re.search(r"^(?:圖|表)\s*\d+", t))


def classify_non_question_block(text: str) -> str | None:
    t = str(text or "").strip()
    if not t:
        return "narration"
    if t == "[BLOCK_IMAGE]" or _is_figure_caption_line(t):
        return "figure_caption"

    has_explain = any(cue in t for cue in _NON_QUESTION_EXPLANATION_CUES) or bool(re.search(r"^(?:說明|解題|觀念)\b", t))
    has_question_verb = any(v in t for v in _QUESTION_VERBS)
    if has_explain and not has_question_verb:
        return "concept_explanation"
    if "[BLOCK_IMAGE]" in t and not has_question_verb:
        return "figure_caption"
    return None


_LEADING_PROBLEM_TITLE_RE = re.compile(
    r"^\s*(?:"
    r"例(?:題)?|隨堂練習|習題|基礎題|進階題|自我評量"
    r")\s*\d{0,3}\s*[\s\.,、，\.：:·…]+",
    flags=re.UNICODE,
)
# 章節習題排版：9\t、10.、6、 — 強分隔符（不含單一空格，避免誤傷「3 是一個奇數」）
_LEADING_CHAPTER_NUM_STRONG_RE = re.compile(
    r"^\s*\d{1,2}(?:[\.、\)\t])+\s*",
    flags=re.UNICODE,
)
# 純數字 + 單一空格：僅在後接常見題幹動詞／「已」「設」等時切除（對齊習題掃描語意）
_LEADING_CHAPTER_NUM_SPACE_RE = re.compile(
    r"^\s*\d{1,2}\s+(?=已|設|根|若|試|利|知|求|解|計|算|證|判|比|化|作|根據|利用|試問|[（(])",
    flags=re.UNICODE,
)


def clean_problem_leading_title(text: str) -> str:
    """Remove a leading question declaration (例/隨堂練習/習題 zone markers, etc.)."""
    t = str(text or "").strip()
    if not t:
        return ""
    t = re.sub(_LEADING_PROBLEM_TITLE_RE, "", t, count=1).strip()
    t = re.sub(_LEADING_CHAPTER_NUM_STRONG_RE, "", t, count=1).strip()
    t = re.sub(_LEADING_CHAPTER_NUM_SPACE_RE, "", t, count=1).strip()
    return t


def segment_question_block_text(problem_text: str, question_title: str = "") -> tuple[str, dict]:
    text = str(problem_text or "")
    lines = [ln for ln in text.splitlines()]
    kept: list[str] = []
    dropped_reason = ""
    started = False

    for idx, raw_line in enumerate(lines):
        line = str(raw_line or "").strip()
        if not line:
            if started:
                kept.append(raw_line)
            continue
        if idx == 0:
            kept.append(raw_line)
            started = True
            continue
        if is_structural_boundary_line(line):
            if question_title and _extract_question_title_from_text(line) == str(question_title).replace(" ", ""):
                kept.append(raw_line)
                started = True
                continue
            dropped_reason = f"stopped question block at structural boundary: {line}"
            break
        kind = classify_non_question_block(line)
        if kind == "figure_caption":
            dropped_reason = "detected figure caption, skipped from question text"
            break
        kept.append(raw_line)
        started = True

    result = "\n".join(kept).strip()
    meta = {"changed": result != text.strip(), "reason": dropped_reason}
    return result, meta


def _extract_question_title_from_text(text: str) -> str:
    t = str(text or "").strip()
    for pat in [
        r"([〔\[]?\s*\d{2,3}\s*統測\s*[AB]?)",
        r"(隨堂練習\s*\d+)",
        r"(例題\s*\d+)",
        r"(基礎題\s*\d+)",
        r"(進階題\s*\d+)",
        r"(自我評量\s*\d*)",
        r"((?:\d+-\d+\s*)?習題\s*\d*)",
    ]:
        m = re.search(pat, t)
        if m:
            return m.group(1).replace(" ", "")
    # Bare-numbered exercise: "5 閫??蝑?..." / "1. 閰行?..." ??return bare number string
    m = re.search(r"^(\d+)[\s\.\)]", t)
    if m:
        return m.group(1)
    return t[:20] or ""


def _is_formula_question_text(text: str) -> bool:
    t = str(text or "")
    return any(k in t for k in ("求", "解", "方程", "不等式", "函數"))


def _is_image_question_text(text: str) -> bool:
    t = str(text or "")
    return any(k in t for k in ("圖", "表", "下圖", "上圖", "右圖", "左圖", "如圖", "題圖", "示意圖"))


_QUESTION_LABEL_RE = re.compile(
    r"((?:例(?:題)?|隨堂練習|基礎題|進階題|統測|學測|"
    r"(?:\d+\s*-\s*\d+\s*)?習題)\s*\d+)"
)

# Used to validate bare-number exercise blocks vs. concept-heading blocks.
_QUESTION_VERB_IN_BLOCK_RE = re.compile(
    r"求|解|算|證明|判斷|化簡|計算|作圖|試求|求出"
)
_CONCEPT_CUES_RE = re.compile(
    r"觀念|說明|解題|補充|重點|定義|性質"
)

# Map Gemini label keywords to DOCX key label keywords (must be same type to match).
_LABEL_TYPE_MAP = {
    "例題": "例題",
    "隨堂練習": "隨堂練習",
    "基礎題": "基礎題",
    "進階題": "進階題",
    "統測": "統測",
}


def _normalize_docx_question_title_key(title: str) -> str:
    """Normalize question title key for DOCX mapping lookups."""
    t = re.sub(r"\s+", "", str(title or "").strip())
    if not t:
        return ""
    # Alias: "例1" and "例題1" are equivalent.
    t = re.sub(r"^例?!題?(\d+)$", r"例題\1", t)
    return t


def _build_docx_title_aliases(title: str, *, include_bare_number: bool = False) -> list[str]:
    """Build ordered alias candidates for title-to-block / title-to-asset matching."""
    raw = str(title or "").strip()
    ns = _normalize_docx_question_title_key(raw)
    aliases: list[str] = []

    def _add(v: str):
        vv = _normalize_docx_question_title_key(v)
        if vv and vv not in aliases:
            aliases.append(vv)

    _add(raw)
    _add(ns)

    # Prefix stripping alias: "1-1蝧? ?箇?憿?" -> "?箇?憿?"
    m = _QUESTION_LABEL_RE.search(ns)
    if m:
        _add(m.group(1))

    # 靘? / 靘?alias
    m_ex = re.search(r"^例題?(\d+)$", ns)
    if m_ex:
        _add(f"例題{m_ex.group(1)}")
    m_ex_short = re.search(r"^例(\d+)$", re.sub(r"\s+", "", raw))
    if m_ex_short:
        _add(f"例題{m_ex_short.group(1)}")

    if include_bare_number:
        # Bare number alias for exercise area fallback.
        m_num = re.search(r"(\d+)$", ns)
        if m_num:
            _add(m_num.group(1))
    return aliases


def _extract_label_and_number(title: str) -> tuple[str | None, str | None]:
    t = _normalize_docx_question_title_key(title)
    for label in ("例題", "隨堂練習", "基礎題", "進階題", "統測"):
        m = re.search(rf"^{re.escape(label)}(\d+)$", t)
        if m:
            return label, m.group(1)
    # "例" alias
    m = re.search(r"^例(\d+)$", re.sub(r"\s+", "", str(title or "").strip()))
    if m:
        return "例題", m.group(1)
    return None, None


def _normalize_docx_key_for_scan(raw_key: str) -> str:
    return _normalize_docx_question_title_key(str(raw_key or ""))


def _find_docx_prefix_match(title: str, source_map: dict):
    """Find key by normalized-prefix match, preserving label type."""
    if not isinstance(source_map, dict) or not source_map:
        return None
    label, num = _extract_label_and_number(title)
    if not label or not num:
        return None
    expected = f"{label}{num}"
    expected_short = f"例{num}" if label == "例題" else None
    for k, v in source_map.items():
        kn = _normalize_docx_key_for_scan(k)
        if not kn:
            continue
        if kn.startswith(expected):
            return v
        if expected_short and kn.startswith(expected_short):
            return v
    return None


def _find_exercise_section_key(source_map: dict) -> str:
    if not isinstance(source_map, dict):
        return ""
    for k in source_map.keys():
        kn = _normalize_docx_key_for_scan(k)
        if kn and re.search(r"(?:\d+-\d+)?習題$", kn):
            return str(k)
    return ""


def _is_safe_exercise_block(block_text: str, num: str) -> bool:
    """Return True only when *block_text* is a genuine bare-numbered exercise.

    A safe exercise block:
    - Starts with ``N<space>`` (NOT ``N.`` which signals a concept-heading).
    - Contains at least one unambiguous question verb.
    - Does NOT read as a pure concept heading (no concept cues without a verb).

    This guard prevents bare-number key "1" from being erroneously paired with
    concept paragraphs like "1.銝?撘????扯釭??.
    """
    t = str(block_text or "").strip()
    # Must begin with the exact number followed by a space (not period/comma).
    if not re.match(rf"^{re.escape(num)}\s", t):
        return False
    if not _QUESTION_VERB_IN_BLOCK_RE.search(t):
        return False
    return True


def _lookup_docx_formula_block(
    title: str,
    formula_blocks: dict,
    *,
    allow_section_fallback: bool = True,
    allow_bare_number: bool = True,
) -> str:
    """High-confidence lookup for a raw DOCX formula block by Gemini question title.

    Strategies (all high-confidence):
    1. Exact match after stripping spaces.
    2. Strip leading section prefix and match the trailing label token exactly
       (e.g. "1-1蝧? ?箇?憿?" ??look up "?箇?憿?").
    3. Same-type label number match: trailing number N and same label type
       (e.g. title "?箇?憿?" searches only for keys "?箇?憿?", never "靘?5").
    4. Safe bare-number match: key is bare number N, block content starts with
       "N " (space, NOT period) and contains a question verb.

    Returns empty string when no high-confidence match is found.
    Deliberately does NOT perform cross-type number matching or unrestricted
    bare-number lookup, which caused concept paragraphs to be misassigned.
    """
    if not title or not formula_blocks:
        return ""
    title_ns = _normalize_docx_question_title_key(title)

    for bk, bv in formula_blocks.items():
        if _normalize_docx_question_title_key(bk) == title_ns:
            return bv

    # Strategy 1: exact + title aliases (靘?/靘?1/靘? 1...)
    for alias in _build_docx_title_aliases(title, include_bare_number=allow_bare_number):
        v = formula_blocks.get(alias)
        if v:
            return v
    v = formula_blocks.get(str(title))
    if v:
        return v

    # Strategy 2: strip section prefix ??exact label lookup
    m = _QUESTION_LABEL_RE.search(title_ns)
    label_key = m.group(1).replace(" ", "") if m else None
    if label_key:
        v = formula_blocks.get(label_key)
        if v:
            return v
    v = _find_docx_prefix_match(title, formula_blocks)
    if v:
        return v

    # Strategy 3: same-type label + number (only match within the same label type)
    m2 = re.search(r"(\d+)$", title_ns)
    if m2:
        num = m2.group(1)
        title_label_type = None
        for kw in _LABEL_TYPE_MAP:
            if kw in title_ns:
                title_label_type = kw
                break
        if title_label_type:
            same_type_key = title_label_type + num
            v = formula_blocks.get(same_type_key)
            if v:
                return v

        # Strategy 4: safe bare-number key (N<space> + question verb required)
        if allow_bare_number:
            v = formula_blocks.get(num)
            if v and _is_safe_exercise_block(v, num):
                return v

    # Strategy 5: exercise-section fallback (disabled for converted_docx_latex hydrate)
    if allow_section_fallback:
        t = _normalize_docx_question_title_key(title)
        if re.search(r"(?:\d+-\d+)?習題.*(?:基礎題|進階題)\d+$", t):
            sec_key = _find_exercise_section_key(formula_blocks)
            if sec_key:
                return str(formula_blocks.get(sec_key) or "")

    return ""


_COORD_GUARD_SECTION_RE = re.compile(r"1\s*-\s*2|平面坐標|坐標平面|坐標幾何")
_COORD_GUARD_TEXT_RE = re.compile(r"坐標|平面|距離|中點|象限|點")


def _is_b1_coordinate_context(volume: str, section_title: str, problem_text: str) -> bool:
    vol = str(volume or "")
    sec = str(section_title or "")
    text = str(problem_text or "")
    is_b1 = ("B1" in vol.upper()) or ("數學B1" in vol)
    if not is_b1:
        return False
    if _COORD_GUARD_SECTION_RE.search(sec):
        return True
    return bool(_COORD_GUARD_TEXT_RE.search(text))


def _normalize_coordinate_point_notation(text: str) -> str:
    """Normalize OCR/LaTeX-mangled coordinate point labels into A(x,y)-style forms."""
    out = str(text or "")
    label = r"([ABCPQR])"
    idx = r"([A-Za-z0-9+\-*/^()]+)"
    def _coord_sup_sub(lbl: str, sup: str, sub: str) -> str:
        # B1 regression: C^y_x should map to C(x,y); P^a_b should map to P(a,b).
        if str(lbl).upper() == "C":
            return f"{lbl}({sub},{sup})"
        return f"{lbl}({sup},{sub})"
    out = re.sub(
        rf"(?:\{{\s*\}}\s*)?\^\{{\s*{idx}\s*\}}\s*{label}\s*_\{{\s*{idx}\s*\}}",
        lambda m: f"{m.group(2)}({m.group(1)},{m.group(3)})",
        out,
    )
    out = re.sub(
        rf"{label}\s*\^\{{\s*{idx}\s*\}}\s*_\{{\s*{idx}\s*\}}",
        lambda m: _coord_sup_sub(m.group(1), m.group(2), m.group(3)),
        out,
    )
    out = re.sub(
        rf"{label}\s*_\{{\s*{idx}\s*\}}\s*\^\{{\s*{idx}\s*\}}",
        lambda m: f"{m.group(1)}({m.group(3)},{m.group(2)})",
        out,
    )
    out = re.sub(
        rf"{label}\s*\^\s*{idx}\s*_\s*{idx}",
        lambda m: _coord_sup_sub(m.group(1), m.group(2), m.group(3)),
        out,
    )
    out = re.sub(
        rf"{label}\s*_\s*{idx}\s*\^\s*{idx}",
        lambda m: f"{m.group(1)}({m.group(3)},{m.group(2)})",
        out,
    )
    return out


def _lookup_docx_question_assets(title: str, q_assets: dict) -> list:
    """High-confidence lookup for question assets by Gemini question title.

    Mirrors the same four high-confidence strategies as
    :func:`_lookup_docx_formula_block`.  Cross-type and unrestricted bare-number
    lookups are intentionally excluded.
    """
    if not title or not q_assets:
        return []
    title_ns = _normalize_docx_question_title_key(title)

    # Strategy 1: exact + title aliases
    for alias in _build_docx_title_aliases(title, include_bare_number=True):
        v = q_assets.get(alias)
        if v:
            return v
    v = q_assets.get(str(title))
    if v:
        return v

    # Strategy 2: strip section prefix
    m = _QUESTION_LABEL_RE.search(title_ns)
    label_key = m.group(1).replace(" ", "") if m else None
    if label_key:
        v = q_assets.get(label_key)
        if v:
            return v

    # Strategy 3: same-type label + number
    m2 = re.search(r"(\d+)$", title_ns)
    if m2:
        num = m2.group(1)
        title_label_type = None
        for kw in _LABEL_TYPE_MAP:
            if kw in title_ns:
                title_label_type = kw
                break
        if title_label_type:
            same_type_key = title_label_type + num
            v = q_assets.get(same_type_key)
            if v:
                return v

        # Strategy 4: safe bare-number
        v = q_assets.get(num)
        # q_assets keys are produced from detected question starts, so bare-number key
        # here is already a question-level anchor and can be safely used.
        if v:
            return v

    # Strategy 5: same-label prefix key (e.g. key="靘? ?貊?銝?.." for title "靘? 1")
    v = _find_docx_prefix_match(title, q_assets)
    if v:
        mapped = []
        for a in (v or []):
            aa = dict(a)
            aa.setdefault("mapping_status", "prefix_match")
            mapped.append(aa)
        return mapped

    # Strategy 6: exercise section fallback (e.g. only key="1-1蝧?")
    t = _normalize_docx_question_title_key(title)
    if re.search(r"(?:\d+-\d+)?習題.*(?:基礎題|進階題)\d+$", t):
        sec_key = _find_exercise_section_key(q_assets)
        if sec_key:
            fallback_assets = []
            for a in (q_assets.get(sec_key) or []):
                aa = dict(a)
                aa.setdefault("mapping_status", "exercise_section_fallback")
                fallback_assets.append(aa)
            if fallback_assets:
                return fallback_assets

    return []


def attach_docx_media_to_question_blocks(blocks):
    question_assets: dict[str, list[dict[str, Any]]] = {}
    orphan_images: list[dict[str, Any]] = []
    image_kw = [
        "圖", "表", "下圖", "上圖", "右圖", "左圖", "如圖",
        "題圖", "示意圖", "座標圖", "函數圖", "幾何圖",
    ]

    question_points: list[dict[str, Any]] = []
    image_blocks: list[dict[str, Any]] = []
    for b in blocks:
        if b.get("type") == "paragraph":
            txt = str(b.get("text", "") or "")
            if _is_question_start_text(txt):
                title = _extract_question_title_from_text(txt)
                img_dep = classify_image_dependency(txt)
                q = {
                    "title": title,
                    "block_index": int(b.get("block_index") or 0),
                    "text": txt,
                    "has_image_kw": any(k in txt for k in image_kw),
                    "needs_image_asset": bool(img_dep.get("needs_image")),
                    "image_role": str(img_dep.get("image_role") or "unknown"),
                    "has_formula_kw": _is_formula_question_text(txt),
                }
                question_points.append(q)
                question_assets.setdefault(title, [])
        elif b.get("type") == "image":
            image_blocks.append(dict(b))

    for img in image_blocks:
        img_idx = int(img.get("block_index") or 0)
        if not question_points:
            orphan_images.append(img)
            continue

        prev_q = None
        next_q = None
        for q in question_points:
            if q["block_index"] <= img_idx:
                prev_q = q
            elif q["block_index"] > img_idx and next_q is None:
                next_q = q
                break

        attached = False
        def _classify_asset(asset_obj, q):
            ext = os.path.splitext(str(asset_obj.get("path") or ""))[1].lower().lstrip(".")
            # Treat as formula_asset when:
            # (a) WMF/EMF image adjacent to a formula-keyword question, OR
            # (b) image originated from a paragraph that contained [FORMULA_IMAGE_N]
            is_ole_formula = asset_obj.get("is_formula_placeholder_source", False)
            if (ext in ("wmf", "emf") and q.get("has_formula_kw")) or is_ole_formula:
                asset_obj["media_kind"] = "formula_asset"
                asset_obj["asset_type"] = "word_formula_image"
                reason = "ole_formula_placeholder_source" if is_ole_formula else "formula_question_block"
                if has_app_context():
                    current_app.logger.info(
                        f"[DOCX MEDIA CLASSIFY] rid={asset_obj.get('rid')} kind=formula_asset reason={reason}"
                    )
                m = re.search(r"\[FORMULA_IMAGE_(\d+)\]", str(q.get("text", "") or ""))
                if m:
                    asset_obj["placeholder_index"] = int(m.group(1))
                    asset_obj["placeholder_token"] = f"[FORMULA_IMAGE_{m.group(1)}]"
                asset_obj["original_path"] = str(asset_obj.get("path") or "")
                asset_obj["original_format"] = _docx_image_original_format(
                    asset_obj.get("path") or "", str(asset_obj.get("content_type") or "")
                )
            else:
                asset_obj["media_kind"] = "image_asset"
                asset_obj["asset_type"] = "word_embedded_image"
                reason = (
                    "question_requires_diagram"
                    if q.get("needs_image_asset")
                    else "decorative_or_default_image_asset"
                )
                if has_app_context():
                    current_app.logger.info(
                        f"[DOCX MEDIA CLASSIFY] rid={asset_obj.get('rid')} kind=image_asset reason={reason}"
                    )

        # Case 1: image before first question; attach when image/formula kw OR OLE formula image.
        if prev_q is None and next_q is not None:
            img_ext = os.path.splitext(str(img.get("path") or ""))[1].lower().lstrip(".")
            is_ole_img = img.get("is_formula_placeholder_source", False)
            if (
                next_q.get("needs_image_asset")
                or next_q["has_formula_kw"]
                or is_ole_img
                or img_ext in ("wmf", "emf")
            ):
                asset = dict(img)
                asset["image_attach_reason"] = "near_next_question"
                asset["needs_image_review"] = True
                _classify_asset(asset, next_q)
                question_assets.setdefault(next_q["title"], []).append(asset)
                attached = True
            else:
                orphan_images.append(img)
                continue

        # Case 2: image after a known question and before next question.
        if not attached and prev_q is not None and next_q is not None:
            d_prev = abs(img_idx - prev_q["block_index"])
            d_next = abs(next_q["block_index"] - img_idx)
            if d_prev == d_next:
                if prev_q.get("needs_image_asset") or next_q.get("needs_image_asset"):
                    shared_prev = dict(img)
                    shared_next = dict(img)
                    shared_prev["image_attach_reason"] = "shared_nearby_image"
                    shared_prev["needs_image_review"] = True
                    shared_prev["shared_image"] = True
                    shared_next["image_attach_reason"] = "shared_nearby_image"
                    shared_next["needs_image_review"] = True
                    shared_next["shared_image"] = True
                    _classify_asset(shared_prev, prev_q)
                    _classify_asset(shared_next, next_q)
                    question_assets.setdefault(prev_q["title"], []).append(shared_prev)
                    question_assets.setdefault(next_q["title"], []).append(shared_next)
                    attached = True
                else:
                    orphan_images.append(img)
                    continue
            elif next_q.get("needs_image_asset") and d_next <= d_prev:
                asset = dict(img)
                asset["image_attach_reason"] = "near_next_question"
                asset["needs_image_review"] = True
                _classify_asset(asset, next_q)
                question_assets.setdefault(next_q["title"], []).append(asset)
                attached = True
            elif prev_q.get("needs_image_asset") and d_prev < d_next:
                asset = dict(img)
                asset["image_attach_reason"] = "near_prev_question"
                asset["needs_image_review"] = True
                _classify_asset(asset, prev_q)
                question_assets.setdefault(prev_q["title"], []).append(asset)
                attached = True
            elif prev_q.get("needs_image_asset"):
                asset = dict(img)
                asset["image_attach_reason"] = "image_inside_question_block"
                asset["needs_image_review"] = True
                _classify_asset(asset, prev_q)
                question_assets.setdefault(prev_q["title"], []).append(asset)
                attached = True
            else:
                orphan_images.append(img)
                continue

        # Case 3: image after last question -> attach only when diagram is required.
        if not attached and prev_q is not None and next_q is None:
            if prev_q.get("needs_image_asset"):
                asset = dict(img)
                asset["image_attach_reason"] = "image_inside_question_block"
                asset["needs_image_review"] = True
                _classify_asset(asset, prev_q)
                question_assets.setdefault(prev_q["title"], []).append(asset)
                attached = True
            else:
                orphan_images.append(img)
                continue

        if not attached:
            orphan_images.append(img)

    return question_assets, orphan_images


_SCAN_ZONE_HEADERS = ("基礎題", "進階題", "自我評量")
_SCAN_EXERCISE_BLOCK_HDR_RE = re.compile(r"^\s*(\d+-\d+)\s*習題\s*$")
_SCAN_ZONE_HDR_RE = re.compile(r"^\s*(基礎題|進階題|自我評量)\s*$")
_SCAN_EXAM_MARKER_RE = re.compile(
    r"[〔\[]?\s*(\d{2,3})\s*統測\s*([AB])\s*[〕\]]?",
    flags=re.IGNORECASE,
)
_SCAN_KEY_LINE_RE = re.compile(r"^\s*KEY\b", re.IGNORECASE)
_SCAN_CHAPTER_EX_NUM_RE = re.compile(r"^\s*(\d{1,2})(?:[\.、\)\t]|\s+)")
_SCAN_EXAMPLE_START_RE = re.compile(r"^\s*例(?:題)?\s*(\d{1,2})\b")
_SCAN_EXAMPLE_NUM_RE = re.compile(r"例(?:題)?\s*(\d{1,2})\b")
_EXAMPLE_BOUNDARY_CHARS = "。；;!?）)】]"


def _scan_find_example_inline_start(line: str) -> re.Match[str] | None:
    """Find 例/例題 N when it starts a new block (line head or after sentence boundary)."""
    for m in _SCAN_EXAMPLE_NUM_RE.finditer(str(line or "")):
        if m.start() == 0:
            return m
        prefix = str(line or "")[: m.start()].rstrip()
        if not prefix or prefix[-1] in _EXAMPLE_BOUNDARY_CHARS:
            return m
    return None
_SCAN_SUITANG_START_RE = re.compile(r"^\s*隨堂練習[\s\.…·]*(\d{1,2})\b")
_SCAN_SUITANG_PREFIX_RE = re.compile(r"^\s*隨堂練習")
_SCAN_SUITANG_NUM_INLINE_RE = re.compile(r"隨堂練習[\s\.…·]*(\d{1,2})\b")
_SCAN_SUBSECTION_HEADING_RE = re.compile(r"^\s*\d+\s*-\s*\d+(?:\.\d+)?\s+\S")
_SCAN_MC_OPTION_RE = re.compile(r"^\s*[\(（]\s*([A-DＡ-Ｄa-dａ-ｄ])\s*[\)）]")
_SCAN_SUBPART_RE = re.compile(r"^\s*[\(（]\s*\d+\s*[\)）]")


_SOLUTION_AND_ANSWER_CUES = re.compile(
    r"可化為|因式分解得|故不等式|的解為|恆大於0|所以不等式|解為|恆成立|當x.*時"
)
_CONCEPT_EXPOSITION_CUES = re.compile(
    r"當判別式|觀察上圖|觀察下圖|二次函數值的|我們稱此|落在x軸下方|落在x軸"
)
_SCAN_BUFFER_STOP_CUE_PATTERNS = (_SOLUTION_AND_ANSWER_CUES, _CONCEPT_EXPOSITION_CUES)


def _scan_truncate_line_at_solution_cue(line: str) -> tuple[str, bool]:
    """Return (line_prefix_before_cue, stop_buffering)."""
    raw = str(line or "")
    earliest: int | None = None
    for pat in _SCAN_BUFFER_STOP_CUE_PATTERNS:
        m = pat.search(raw)
        if m and (earliest is None or m.start() < earliest):
            earliest = m.start()
    if earliest is None:
        return raw, False
    head = raw[:earliest].rstrip()
    return head, True


def _scan_is_structure_only_line(line: str) -> bool:
    """True when the line is only a structural marker (not marker + question body)."""
    s = str(line or "").strip()
    if not s:
        return False
    if _SCAN_KEY_LINE_RE.match(s):
        return True
    if _SCAN_EXAM_MARKER_RE.search(s) and not re.search(r"[\(（][A-DＡ-Ｄ]", s):
        return True
    if _SCAN_SUBSECTION_HEADING_RE.match(s):
        return True
    if _SCAN_ZONE_HDR_RE.match(s):
        return True
    if _SCAN_EXERCISE_BLOCK_HDR_RE.match(s) or re.match(r"^\s*(\d+-\d+)\s*習題\s*$", s):
        return True
    ex_m = _SCAN_EXAMPLE_START_RE.match(s)
    if ex_m:
        rest = s[ex_m.end() :].strip()
        return not rest
    if _SCAN_SUITANG_PREFIX_RE.match(s):
        body = _SCAN_SUITANG_PREFIX_RE.sub("", s, count=1).strip()
        body = re.sub(r"^[\s\.…·]+", "", body).strip()
        if _SCAN_SUITANG_NUM_INLINE_RE.search(body):
            body = _SCAN_SUITANG_NUM_INLINE_RE.sub("", body, count=1).strip()
        return not body
    return False


def _scan_finalize_question_buffer(lines: list[str]) -> str:
    kept: list[str] = []
    for ln in lines:
        s = str(ln or "").strip()
        if _SCAN_KEY_LINE_RE.match(s):
            break
        if _scan_is_structure_only_line(s):
            continue
        truncated, stop = _scan_truncate_line_at_solution_cue(str(ln or ""))
        if truncated.strip():
            kept.append(truncated)
        if stop:
            break
    while kept and _scan_is_structure_only_line(str(kept[-1] or "")):
        kept.pop()
    return clean_problem_leading_title("\n".join(kept).strip())


def _scan_flush_question_block(blocks: dict[str, str], key: str | None, buf: list[str]) -> None:
    if not key:
        return
    text = _scan_finalize_question_buffer(buf)
    if not text:
        return
    blocks[str(key).strip()] = text


def _scan_line_flushes_current_block(
    line: str,
    *,
    in_chapter_exercise: bool,
    pending_suithang_header: bool = False,
) -> bool:
    """Return True when *line* starts a new structural block (must flush previous)."""
    if _SCAN_KEY_LINE_RE.match(line):
        return True
    if _SCAN_EXAM_MARKER_RE.search(line):
        return True
    if _SCAN_SUBSECTION_HEADING_RE.match(line):
        return True
    if _SCAN_SUITANG_PREFIX_RE.match(line):
        return True
    if _SCAN_ZONE_HDR_RE.match(line):
        return True
    if _SCAN_EXERCISE_BLOCK_HDR_RE.match(line) or re.match(r"^\s*(\d+-\d+)\s*習題\s*$", line):
        return True
    if _SCAN_EXAMPLE_START_RE.match(line):
        return True
    if _SCAN_SUITANG_NUM_INLINE_RE.search(line):
        return True
    if pending_suithang_header:
        return False
    if in_chapter_exercise and _SCAN_CHAPTER_EX_NUM_RE.match(line):
        return True
    return False


def _scan_try_start_suithang(line: str, pending_header: bool) -> tuple[str | None, str | None, bool]:
    """Return (key, first_line, still_pending_header)."""
    m_inline = _SCAN_SUITANG_NUM_INLINE_RE.search(line)
    if m_inline:
        first = clean_problem_leading_title(line[m_inline.start() :])
        return f"隨堂練習{int(m_inline.group(1))}", first or None, False
    if pending_header and _SCAN_CHAPTER_EX_NUM_RE.match(line):
        n = int(_SCAN_CHAPTER_EX_NUM_RE.match(line).group(1))
        return f"隨堂練習{n}", line, False
    if _SCAN_SUITANG_PREFIX_RE.match(line):
        return None, None, True
    return None, None, pending_header


def _scan_converted_docx_blocks_from_lines(
    lines: list[str],
    *,
    source_scope: str = "",
) -> dict[str, str]:
    """Build canonical question block keys aligned with scan_docx_title_inventory()."""
    scope = str(source_scope or "").strip()
    sa_ctx = None
    if scope != "section_textbook":
        sa_ctx = detect_chapter_self_assessment_context("\n".join(lines))
    if sa_ctx:
        return _scan_chapter_self_assessment_blocks(lines, sa_ctx)

    blocks: dict[str, str] = {}
    section_code: str | None = None
    current_zone = "其他"
    in_chapter_exercise = False
    current_key: str | None = None
    buffer: list[str] = []
    buffer_stop_extend = False
    pending_exam_lines: list[str] = []
    awaiting_exam = False
    pending_suithang_header = False

    def flush() -> None:
        nonlocal current_key, buffer, buffer_stop_extend
        _scan_flush_question_block(blocks, current_key, buffer)
        current_key = None
        buffer = []
        buffer_stop_extend = False

    def append_buffer_line(line: str) -> None:
        nonlocal buffer_stop_extend
        if buffer_stop_extend:
            return
        truncated, stop = _scan_truncate_line_at_solution_cue(line)
        if truncated.strip():
            buffer.append(truncated)
        if stop:
            buffer_stop_extend = True

    def start_key(key: str, first_line: str | None = None) -> None:
        nonlocal current_key, buffer, awaiting_exam, pending_exam_lines, pending_suithang_header, buffer_stop_extend
        flush()
        awaiting_exam = False
        pending_exam_lines = []
        pending_suithang_header = False
        buffer_stop_extend = False
        current_key = key
        if first_line:
            cleaned = clean_problem_leading_title(first_line)
            buffer = [cleaned] if cleaned else []
        else:
            buffer = []

    def begin_exam_staging(line: str) -> None:
        nonlocal awaiting_exam, pending_exam_lines, pending_suithang_header
        pending_suithang_header = False
        if awaiting_exam:
            pending_exam_lines.append(line)
            return
        flush()
        awaiting_exam = True
        pending_exam_lines = [line]

    for raw in lines:
        line = str(raw or "").strip()
        if not line:
            if current_key:
                buffer.append("")
            elif awaiting_exam:
                pending_exam_lines.append("")
            continue

        if _SCAN_KEY_LINE_RE.match(line):
            flush()
            pending_suithang_header = False
            continue

        boundary_flush = False
        if current_key and _scan_line_flushes_current_block(
            line,
            in_chapter_exercise=in_chapter_exercise,
            pending_suithang_header=pending_suithang_header,
        ):
            flush()
            pending_suithang_header = False
            boundary_flush = True

        exam_m = _SCAN_EXAM_MARKER_RE.search(line)
        if exam_m:
            exam_key = f"{int(exam_m.group(1))}統測{exam_m.group(2).upper()}"
            before = line[: exam_m.start()].strip()
            after = line[exam_m.end() :].strip()
            pending_suithang_header = False

            stem: list[str] = []
            if buffer:
                stem.extend(buffer)
            if pending_exam_lines:
                stem.extend(pending_exam_lines)
            if before:
                stem.append(before)

            prior_key = current_key
            buffer = []
            pending_exam_lines = []
            awaiting_exam = False
            buffer_stop_extend = False
            if prior_key and prior_key != exam_key:
                _scan_flush_question_block(blocks, prior_key, [])

            current_key = exam_key
            buffer = stem
            if after and not _SCAN_KEY_LINE_RE.match(after):
                append_buffer_line(after)
            continue

        if _SCAN_SUBSECTION_HEADING_RE.match(line):
            flush()
            pending_suithang_header = False
            in_chapter_exercise = False
            continue

        ex_hdr = _SCAN_EXERCISE_BLOCK_HDR_RE.match(line)
        if not ex_hdr:
            blk = re.match(r"^\s*(\d+-\d+)\s*習題\s*$", line)
            if blk:
                ex_hdr = blk
        if ex_hdr:
            flush()
            pending_suithang_header = False
            section_code = ex_hdr.group(1)
            in_chapter_exercise = True
            current_zone = "其他"
            awaiting_exam = False
            pending_exam_lines = []
            continue

        if _SCAN_ZONE_HDR_RE.match(line):
            flush()
            pending_suithang_header = False
            current_zone = _SCAN_ZONE_HDR_RE.match(line).group(1)
            in_chapter_exercise = bool(section_code)
            continue

        ex_m = _scan_find_example_inline_start(line)
        if ex_m:
            body = clean_problem_leading_title(line[ex_m.start() :])
            start_key(f"例題{int(ex_m.group(1))}", body or None)
            in_chapter_exercise = False
            continue

        st_key, st_line, pending_suithang_header = _scan_try_start_suithang(
            line, pending_suithang_header
        )
        if st_key:
            start_key(st_key, st_line)
            in_chapter_exercise = False
            continue
        if pending_suithang_header and not _SCAN_CHAPTER_EX_NUM_RE.match(line):
            continue
        if boundary_flush and _scan_is_structure_only_line(line):
            continue

        if in_chapter_exercise and section_code and current_zone in _SCAN_ZONE_HEADERS:
            mnum = _SCAN_CHAPTER_EX_NUM_RE.match(line)
            if mnum:
                n = int(mnum.group(1))
                start_key(f"{section_code}習題 {current_zone}{n}", line)
                continue

            if current_key:
                if _SCAN_SUBPART_RE.match(line):
                    append_buffer_line(line)
                    continue
                if _SCAN_MC_OPTION_RE.match(line) or awaiting_exam:
                    begin_exam_staging(line)
                    continue
                if re.match(r"^\s*設\s", line):
                    begin_exam_staging(line)
                    continue
                append_buffer_line(line)
                continue

        if current_key and not in_chapter_exercise:
            append_buffer_line(line)
            continue

        if awaiting_exam or pending_exam_lines or _SCAN_MC_OPTION_RE.match(line):
            if not awaiting_exam:
                begin_exam_staging(line)
            else:
                pending_exam_lines.append(line)
            continue

    flush()
    return blocks


def build_docx_question_formula_context(blocks, *, source_scope: str = ""):
    lines: list[str] = []
    for b in blocks or []:
        btype = b.get("type")
        if btype == "paragraph":
            txt = str(b.get("text", "") or "").strip()
            if txt:
                lines.append(txt)
        elif btype == "image":
            lines.append("[BLOCK_IMAGE]")
    return _scan_converted_docx_blocks_from_lines(lines, source_scope=source_scope)


def _safe_title_for_filename(title: str) -> str:
    t = re.sub(r"\s+", "", str(title or "").strip())
    t = re.sub(r"[\\/:*?\"<>|]", "_", t)
    return t[:40] or "untitled"


def _copy_docx_asset_to_question_assets(src_path: str, dst_dir: str, filename: str) -> str | None:
    try:
        import shutil
        if not src_path:
            return None
        if not os.path.isabs(src_path):
            src_abs = os.path.join(current_app.root_path, src_path)
        else:
            src_abs = src_path
        if not os.path.exists(src_abs):
            return None
        os.makedirs(dst_dir, exist_ok=True)
        dst_abs = os.path.join(dst_dir, filename)
        shutil.copy2(src_abs, dst_abs)
        return dst_abs
    except Exception:
        return None


def parse_volume(volume_str: str):
    """Parse vocational math volume label into (subject, volume_number). Never raises re.error."""
    text = str(volume_str or "").strip()
    if not text:
        return None, None

    subject: str | None = None
    vol_num: int | None = None

    try:
        # 數學B1、數學 B2、數學b4
        m = re.search(r"數學\s*([AB])\s*(\d+)", text, re.IGNORECASE)
        if m:
            return m.group(1).upper(), int(m.group(2))

        # B1、A2（單字邊界避免誤判過長字串中的片段）
        m = re.search(r"\b([AB])\s*(\d+)\b", text, re.IGNORECASE)
        if m:
            return m.group(1).upper(), int(m.group(2))

        zh_map = {"一": 1, "二": 2, "三": 3, "四": 4}
        for zh, num in zh_map.items():
            if zh + "冊" in text:
                vol_num = num
                break

        if "數學B" in text or "B" in text.upper():
            subject = "B"
        elif "數學A" in text or "A" in text.upper():
            subject = "A"

        return subject, vol_num
    except re.error:
        return None, None


def grade_for_vocational_math_volume(volume: str) -> int | None:
    """Authoritative grade for vocational Math B volumes.

    數學B1 / 數學B2 → 10
    數學B3 / 數學B4 → 11
    Other volumes → None (caller keeps existing grade).
    """
    subject, vol_num = parse_volume(volume)
    if subject != "B" or not isinstance(vol_num, int):
        return None
    if vol_num in (1, 2):
        return 10
    if vol_num in (3, 4):
        return 11
    return None


def normalize_json_text_before_parse(text):
    """Normalize JSON text before parsing."""
    if not text:
        return text

    normalized = str(text)
    # 已知問題：有些欄位在 Gemini 回傳時可能沒有包含完整的繁體中文字串
    # keep legacy cleanup no-op when source token is corrupted
    normalized = normalized.replace('"銝????拙予??銝????', "")
    return normalized


def sanitize_detailed_solution_text(text, max_chars=500):
    """Sanitize detailed solution text for storage."""
    if text is None:
        return ""

    cleaned = str(text).strip()
    if not cleaned:
        return ""

    banned_phrases = [
        "Let's trace",
        "Let's re-do",
        "This is not",
        "English chain-of-thought",
    ]
    for phrase in banned_phrases:
        cleaned = cleaned.replace(phrase, "")

    # 將會進行二次比對，減少人工審查負擔
    paragraph_parts = [p.strip() for p in re.split(r"\n{2,}", cleaned) if p.strip()]
    if paragraph_parts:
        cleaned = paragraph_parts[-1]

    if len(cleaned) > max_chars:
        cleaned = cleaned[-max_chars:]

    return cleaned.strip()


def process_textbook_file(
    file_path,
    curriculum_info,
    queue,
    skip_code_gen=False,
    outline_only=False,
    toc_pages=5,
    import_policy=None,
    optional_enrich_pdf_path=None,
):
    """Process textbook file and import parsed content."""

    try:
        import_policy = dict(import_policy or {})
        execution_arch = str(import_policy.get("execution_arch", "native") or "native").strip().lower()
        if execution_arch in ("x86", "x64"):
            runtime_arch = "x64" if "64" in str(platform.architecture()[0]) else "x86"
            if runtime_arch != execution_arch:
                message = f"Execution architecture mismatch: requested={execution_arch}, runtime={runtime_arch}"
                current_app.logger.error(message)
                if queue:
                    queue.put(f"ERROR: {message}")
                return {"status": "error", "message": message}
        if optional_enrich_pdf_path and queue:
            queue.put("INFO: optional enrich PDF path provided")
        # ======================================================
        # [NEW] 排除 Word 的暫存檔 (以 ~$ 開頭)
        # ======================================================
        filename = os.path.basename(file_path)
        if filename.startswith("~$"):
            message = f"Skip Word temporary file: {filename}"
            current_app.logger.warning(message)
            if queue:
                queue.put(f"WARN: {message}")
            return {"status": "skipped", "message": message}
        # ======================================================

        # 步驟 1: 從 Word 檔案中提取內容
        content_by_page = extract_content_from_file(
            file_path,
            queue,
            max_pages=toc_pages if outline_only else None,
            import_policy=import_policy,
        )

        # [V2.5] 建立暫時性的單元對照關係
        if outline_only:
            volume_val = str(curriculum_info.get('volume', ''))
            curr_val = str(curriculum_info.get('curriculum', ''))
            
            parsed_data = None
            structure_source = "pdf_toc"

            # 嘗試解析 AI TOC 結構
            if content_by_page:
                toc_json_string = call_gemini_for_toc(content_by_page, curriculum_info, queue)
                if toc_json_string:
                    toc_json_string = normalize_json_text_before_parse(toc_json_string)
                    parsed_data = parse_ai_response(toc_json_string, queue)
                    if parsed_data and parsed_data.get('chapters'):
                        message = "AI TOC extraction succeeded"
                        current_app.logger.info(message)
                        queue.put(f"SUCCESS: {message}")
                        
                        # [V2.6] 移除 OutlinePlaceholder 欄位以簡化結構
                        pass
            
            # 如果 AI 解析失敗，則使用 YAML Fallback
            if not parsed_data:
                struct_map = get_structure_map(curr_val, volume_val)
                if struct_map and struct_map.data:
                    message = f"AI 解析失敗或無內容，改用 YAML 結構匯入 ({volume_val})..."
                    current_app.logger.info(message)
                    queue.put(f"INFO: {message}")
                    structure_source = "yaml_fallback"
                    
                    # 將 YAML 轉換為 parsed_data 格式
                    yaml_chapters = struct_map.data.get('chapters', [])
                    parsed_chapters = []
                    for ch in yaml_chapters:
                        ch_title = f"{ch.get('index')} {ch.get('title')}"
                        sections = []
                        for sec in ch.get('sections', []):
                            sec_title = f"{sec.get('code')} {sec.get('title')}"
                            sections.append({
                                "section_title": sec_title,
                                "concepts": []  # [V2.6] 不需要 Placeholder concept
                            })
                        parsed_chapters.append({
                            "chapter_title": ch_title,
                            "sections": sections
                        })
                    parsed_data = {"chapters": parsed_chapters}
            
            if parsed_data:
                result = import_outline_structure_only(
                    parsed_data,
                    curriculum_info,
                    queue,
                    source_file_path=file_path
                )
                return {
                    "status": "success", 
                    "message": f"目錄結構建立完成 (來源: {structure_source})", 
                    "structure_source": structure_source,
                    "skipped_skills": True,
                    "skipped_examples": True,
                    "skipped_practices": True,
                    "skipped_code_generation": True,
                    **result
                }
            else:
                return {"status": "error", "message": "Failed to parse outline from PDF and YAML fallback."}

        if not content_by_page:
            message = "Failed to extract content from file."
            current_app.logger.error(message)
            queue.put(f"ERROR: {message}")
            return {"status": "error", "message": "Content extraction failed."}

        docx_formula_source_mode = str(
            ((_DOCX_IMPORT_CONTEXT or {}).get("docx_formula_source_mode") or "")
        ).strip()
        content_by_page = _normalize_extracted_content_math(
            content_by_page,
            queue,
            docx_formula_source_mode=docx_formula_source_mode,
        )
        page_analysis_payload = None

        # 步驟 2: 呼叫 AI 進行分析
        file_meta = parse_textbook_filename_metadata(file_path)
        import_policy["section_code"] = str(file_meta.get("section_code", "") or "").strip()
        ai_json_result_string = call_gemini_for_analysis(
            content_by_page,
            curriculum_info,
            queue,
            page_analysis_payload=page_analysis_payload,
            import_policy=import_policy,
        )
        # 步驟 3: 解析 AI 回傳的 JSON 結構
        if ai_json_result_string is None:
            return {"status": "error", "message": "AI analysis failed."}
        if not ai_json_result_string:
            return {"status": "error", "message": "AI returned empty response."}

        ai_json_result_string = normalize_json_text_before_parse(ai_json_result_string)
        parsed_data = parse_ai_response(ai_json_result_string, queue)
        if not parsed_data:
            return {"status": "error", "message": "Failed to parse AI JSON response."}

        # ==============================================================================
        # 已刪除過時的 _mark_needs_review_for_low_quality_pages 呼叫
        # ==============================================================================
        parsed_data = _normalize_parsed_textbook_math(
            parsed_data,
            queue,
            docx_formula_source_mode=docx_formula_source_mode,
        )

        # 步驟 4: 將解析出來的內容寫入資料庫
        if docx_formula_source_mode == "converted_docx_latex":
            extracted_text = "\n".join(str(v or "") for _k, v in sorted((content_by_page or {}).items()))
            section_code = str(import_policy.get("section_code", "") or "unknown").replace(" ", "")
            volume = str(curriculum_info.get("volume", "") or "unknown").replace(" ", "")
            sc_for_scan = None if section_code == "unknown" else section_code
            import_scope = str(file_meta.get("source_scope") or "section_textbook").strip()
            inventory_items = scan_docx_title_inventory(extracted_text, section_code=sc_for_scan)
            if isinstance(_DOCX_IMPORT_CONTEXT, dict):
                _DOCX_IMPORT_CONTEXT["title_inventory_items"] = inventory_items
                _DOCX_IMPORT_CONTEXT["source_scope"] = import_scope
                rescanned = scan_converted_docx_question_blocks(
                    extracted_text,
                    source_scope=import_scope,
                )
                _DOCX_IMPORT_CONTEXT["question_formula_blocks"] = rescanned
            ctx_blocks = (_DOCX_IMPORT_CONTEXT or {}).get("question_formula_blocks", {}) if isinstance(_DOCX_IMPORT_CONTEXT, dict) else {}
            parsed_data, hydrate_filled, hydrate_blocks = hydrate_converted_docx_latex_parsed_data(
                parsed_data,
                extracted_text=extracted_text,
                section_code=sc_for_scan,
                inventory_items=inventory_items,
                question_blocks=ctx_blocks if isinstance(ctx_blocks, dict) else None,
            )
            current_app.logger.info(
                f"[DOCX HYDRATE] filled_problem_text={hydrate_filled} scanned_blocks={hydrate_blocks}"
            )
            if queue is not None:
                queue.put(f"INFO: [DOCX HYDRATE] filled_problem_text={hydrate_filled} scanned_blocks={hydrate_blocks}")
            expected_titles = sorted({str(it.get("canonical_title", "")).strip() for it in inventory_items if it.get("canonical_title")})
            returned_titles = collect_returned_titles_from_parsed_data(parsed_data)
            inv = build_title_inventory(
                expected_titles,
                returned_titles,
                section_code=section_code,
                inventory_items=inventory_items,
            )
            allow_partial_import = bool(import_policy.get("allow_partial_import", False))
            write_aborted = False
            report_name = f"{volume}_{section_code}_title_inventory_report.md"
            report_path = os.path.join("reports", "import_debug", report_name)
            volume_raw = str(curriculum_info.get("volume", "") or "")
            curriculum_raw = str(curriculum_info.get("curriculum", "") or "")
            publisher_raw = str(parse_textbook_filename_metadata(file_path).get("publisher", "") or "")
            warn = detect_curriculum_volume_warning(curriculum_raw, volume_raw, publisher_raw)
            if warn:
                current_app.logger.warning(warn)
                if queue is not None:
                    queue.put(f"WARN: {warn}")
            write_title_inventory_report(
                report_path,
                volume=volume,
                section=section_code,
                allow_partial_import=allow_partial_import,
                write_aborted=write_aborted,
                inv=inv,
                warning=warn,
            )
            current_app.logger.info("[IMPORT INVENTORY GUARD] report_only=true")
            current_app.logger.info(f"[IMPORT INVENTORY GUARD] expected_titles_count={inv.get('expected_titles_count', 0)}")
            current_app.logger.info(f"[IMPORT INVENTORY GUARD] returned_titles_count={inv.get('returned_titles_count', 0)}")
            current_app.logger.info(f"[IMPORT INVENTORY GUARD] missing_titles_count={inv.get('missing_titles_count', 0)}")
            current_app.logger.info(f"[IMPORT INVENTORY GUARD] allow_partial_import={str(allow_partial_import).lower()}")
            if queue is not None:
                queue.put("INFO: [IMPORT INVENTORY GUARD] report_only=true")
                queue.put(f"INFO: [IMPORT INVENTORY GUARD] expected_titles_count={inv.get('expected_titles_count', 0)}")
                queue.put(f"INFO: [IMPORT INVENTORY GUARD] returned_titles_count={inv.get('returned_titles_count', 0)}")
                queue.put(f"INFO: [IMPORT INVENTORY GUARD] missing_titles_count={inv.get('missing_titles_count', 0)}")
                queue.put(f"INFO: [IMPORT INVENTORY GUARD] report_path={report_path}")

        result = save_to_database(
            parsed_data,
            curriculum_info,
            queue,
            source_file_path=file_path,
            content_by_page=content_by_page,
            outline_only=outline_only,
            import_policy=import_policy,
            optional_enrich_pdf_path=optional_enrich_pdf_path,
        )
        try:
            temp_dir = (_DOCX_IMPORT_CONTEXT or {}).get("temp_media_dir")
            if temp_dir and bool(current_app.config.get("CLEAN_ORPHAN_DOCX_MEDIA", True)):
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
                current_app.logger.info(f"[DOCX MEDIA CLEANUP] removed orphan temp dir={temp_dir}")
        except Exception:
            pass

        skills_count = result.get('skills_processed', 0)
        curriculums_count = result.get('curriculums_added', 0)
        examples_count = result.get('examples_added', 0)
        practice_count = result.get('practice_questions_imported', 0)
        in_class_practice_count = result.get('in_class_practices_imported', 0)
        chapter_exercises_count = result.get('chapter_exercises_imported', 0)
        self_assessments_count = result.get('self_assessments_imported', 0)
        exam_practices_count = result.get('exam_practices_imported', 0)
        other_practices_count = result.get('other_practices_imported', 0)
        practice_needs_review_count = result.get('practice_questions_needs_review', 0)
        practice_skipped_count = result.get('duplicates_skipped', result.get('practice_questions_skipped', 0))
        processed_skill_ids = result.get('processed_skill_ids', [])

        message = (
            f"Import complete: skills={skills_count}, curriculums={curriculums_count}, "
            f"examples={examples_count}, practices={practice_count}, in_class={in_class_practice_count}, "
            f"chapter_exercises={chapter_exercises_count}, self_assessment={self_assessments_count}, "
            f"exam={exam_practices_count}, other={other_practices_count}, "
            f"needs_review={practice_needs_review_count}, skipped={practice_skipped_count}"
        )
        current_app.logger.info(message)
        queue.put(f"INFO: {message}")

        # 步驟 5: 產生對應的技能程式碼 (選用)
        code_gen_status = "skipped"
        if skip_code_gen:
            message = "Skip code generation by request."
            current_app.logger.info(message)
            queue.put(f"INFO: {message}")
        elif processed_skill_ids:
            queue.put(f"INFO: start code generation for {len(processed_skill_ids)} skills")
            for idx, skill_id in enumerate(processed_skill_ids):
                queue.put(f"INFO: [{idx+1}/{len(processed_skill_ids)}] 正在寫入 {skill_id}.py ...")
                try:
                    # [修正] 強制 Architect 重新載入最新 Prompt
                    success, msg = auto_generate_skill_code(skill_id, queue, force_architect_refresh=True)
                    if success:
                        queue.put(f"INFO: {skill_id} code generated")
                    else:
                        queue.put(f"WARN: {skill_id} code generation failed")
                except Exception as e:
                    queue.put(f"ERROR: 技能 {skill_id} 程式碼寫入失敗: {e}")
                    current_app.logger.error(f"Generate Error {skill_id}: {e}")
                
                time.sleep(2) # Rate Limit
            code_gen_status = f"{len(processed_skill_ids)} generated"

        return {
            "status": "success", 
            "message": (f"教材分析與匯入成功。\n"
                        f"新增/更新技能 {skills_count} 個\n"
                        f"新增課程綱要 {curriculums_count} 筆\n"
                        f"新增課本例題 {examples_count} 筆\n"
                        f"新增練習題 {practice_count} 筆\n"
                        f"新增隨堂練習 {in_class_practice_count} 筆\n"
                        f"練習題需審核 {practice_needs_review_count} 筆\n"
                        f"練習題跳過 {practice_skipped_count} 筆\n"
                        f"自動產生程式碼 {code_gen_status}")
        }

    except Exception as e:
        current_app.logger.error(f"解析教材內容時發生錯誤: {e}")
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": f"處理失敗: {str(e)}"}

# --- 教材內容處理 ---
process_textbook_pdf = process_textbook_file

def extract_content_from_file(file_path, queue, max_pages=None, import_policy=None):
    """Extract text content from pre-converted MathType->LaTeX DOCX."""
    message = f"正在從 {file_path} 提取內容..."
    current_app.logger.info(message)
    queue.put(f"INFO: {message}")

    global _DOCX_IMPORT_CONTEXT
    _DOCX_IMPORT_CONTEXT = {}
    
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension not in ('.docx', '.doc'):
        message = f"不支援的檔案類型: {file_extension}。請使用 .docx 檔案。"
        current_app.logger.error(message)
        queue.put(f"ERROR: {message}")
        return {}

    try:
        content_by_page, doc_meta = extract_converted_latex_docx(file_path)
        extracted_text = str((content_by_page or {}).get(1, "") or "")
        detect_meta = detect_converted_latex_docx(extracted_text)
        file_meta = parse_textbook_filename_metadata(file_path)
        source_scope = str(file_meta.get("source_scope") or "section_textbook").strip()
        sa_ctx = detect_chapter_self_assessment_context(extracted_text)
        if source_scope == "section_textbook":
            sa_ctx = None
        question_blocks = scan_converted_docx_question_blocks(
            extracted_text,
            source_scope=source_scope,
        )
        _DOCX_IMPORT_CONTEXT = {
            "docx_formula_source_mode": "converted_docx_latex",
            "is_converted_latex_docx": True,
            "latex_signal_count": int(detect_meta.get("latex_signal_count", 0)),
            "formula_placeholder_count": int(detect_meta.get("formula_placeholder_count", 0)),
            "question_assets": {},
            "question_formula_blocks": question_blocks,
            "formula_assets_extraction_skipped": True,
            "ocr_skipped": True,
            "pix2tex_skipped": True,
            "doc_meta": doc_meta,
            "source_scope": source_scope,
            "chapter_self_assessment_mode": bool(sa_ctx),
            "chapter_self_assessment_prefix": str(sa_ctx.get("title_prefix", "") if sa_ctx else ""),
        }
        if sa_ctx:
            queue.put(
                f"INFO: [DOCX CHAPTER SA] mode=chapter_self_assessment "
                f"prefix={sa_ctx.get('title_prefix')} titles={len(question_blocks)}"
            )
        queue.put("INFO: docx_formula_source_mode=converted_docx_latex")
        queue.put("INFO: formula_assets_extraction_skipped=true")
        queue.put("INFO: ocr_skipped=true")
        queue.put("INFO: pix2tex_skipped=true")
        queue.put(f"INFO: is_converted_latex_docx={True}")
        queue.put(f"INFO: latex_signal_count={detect_meta.get('latex_signal_count', 0)}")
        queue.put(f"INFO: formula_placeholder_count={detect_meta.get('formula_placeholder_count', 0)}")
        queue.put(f"INFO: [DOCX BLOCK SCAN] question_blocks={len(question_blocks)}")
        return content_by_page
    except Exception as e:
        message = f"提取檔案內容時發生異常 (Exception): {e}"
        current_app.logger.error(message)
        queue.put(f"ERROR: {message}")
        return {}


def _sanitize_and_parse_json(s: str, queue=None):
    """Sanitize and parse AI JSON response text."""
    if not s:
        return None, "", s, []

    original = s
    
    # ===== 第 0 步：輸出 JSON 調試資訊 =====
    current_app.logger.debug(f"[JSON_DEBUG] 原始 JSON 長度: {len(s)} 字元")
    
    # ===== 第 1 步：清除 code fence wrapper =====
    s = re.sub(r'^```(?:json)?\s*|\s*```$', '', s, flags=re.MULTILINE).strip()
    
    # ===== 第 2 步：移除前後不合規的字元 =====
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)
    
    # ===== 第 3 步：過濾潛在的 BOM =====
    if s.startswith('\ufeff'): s = s[1:]
    
    # ===== 第 4 步：嘗試使用多種策略進行 JSON 解析 =====
    candidates = []
    
    # 蝑 0: ??嚗?蝘駁 control chars / fences嚗?
    candidates.append(("raw", s))
    
    # 蝑 1: 靽???escape - ?芸?敺銝?? JSON escape ???? escape
    escaped_conservative = re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', s)
    candidates.append(("保守 escape", escaped_conservative))
    
    # 蝑 2: 瞈??escape - ??迨蝡????賡???
    escaped_aggressive = re.sub(r'(?<!\\)\\(?!\\)', r'\\\\', s)
    candidates.append(("積極 escape", escaped_aggressive))
    
    # 蝑 3: ?敺?摨?- ??????賡???
    escaped_brutal = s.replace('\\', '\\\\')
    candidates.append(("暴力 escape", escaped_brutal))
    
    # 蝑 4: ?岫?曉蝚砌???{ ??敺???} ??銝?
    first_brace = s.find('{')
    last_brace = s.rfind('}')
    if first_brace >= 0 and last_brace > first_brace:
        substr = s[first_brace:last_brace + 1]
        candidates.append(("擷取 {} 字串", substr))
        candidates.append(("擷取 {} 字串 + 保守 escape", re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', substr)))

    attempts = []
    for strategy_name, cand in candidates:
        try:
            obj = json.loads(cand)
            current_app.logger.info(f"[JSON_SUCCESS] 策略 '{strategy_name}' 成功解析 JSON")
            return obj, cand, original, attempts
        except json.JSONDecodeError as e:
            snippet = (cand[:200] + '...') if len(cand) > 200 else cand
            error_detail = f"{e.msg} at line {e.lineno}, col {e.colno}"
            attempts.append((strategy_name, snippet, error_detail))
            current_app.logger.debug(f"[JSON_FAIL] 策略 '{strategy_name}' 失敗: {error_detail}")

    if queue is not None:
        queue.put(f"ERROR: JSON 解析失敗（已嘗試 {len(attempts)} 種策略，請檢查原始回應）")
    
    if candidates:
        return None, candidates[-1][1], original, [(s, e, d) for s, _, (_, _, d) in zip([c[0] for c in candidates], [], attempts)]
    else:
        return None, "", original, attempts


def _call_gemini_with_retry(model, analysis_prompt, queue=None, context_message='AI 分析', parse_json=False):
    max_retries = 3
    retry_delay = 2

    def _validate_json_completeness(text):
        cleaned_text = re.sub(r'^```(?:json)?\s*|\s*```$', '', str(text or ''), flags=re.MULTILINE).strip()
        if not (cleaned_text.startswith("{") or cleaned_text.startswith("[")):
            return False, "missing_opening_brace", cleaned_text
        if not (cleaned_text.endswith("}") or cleaned_text.endswith("]")):
            return False, "missing_closing_brace", cleaned_text
        # 放寬結尾防禦閘門，只要具備標準 JSON 結尾符號 (} 或 ]) 即可通過
        if re.search(r'(?:\}\s*\]|\s*\}|\]\s*\})\s*$', cleaned_text, flags=re.DOTALL) is None:
            return False, "missing_json_tail", cleaned_text
        fixed_text = sanitize_gemini_json_text(cleaned_text)
        try:
            json.loads(cleaned_text)
            if fixed_text != cleaned_text:
                json.loads(fixed_text)
                current_app.logger.info("[TEXTBOOK IMPORTER] Gemini JSON parsed after LaTeX escape sanitize.")
                return True, "ok_sanitized", fixed_text
            return True, "ok", cleaned_text
        except json.JSONDecodeError as first_error:
            current_app.logger.debug(
                "[TEXTBOOK IMPORTER] Gemini first json.loads failed at "
                f"line {first_error.lineno}, col {first_error.colno}, pos {first_error.pos}: "
                f"{first_error.msg}"
            )
            try:
                json.loads(fixed_text)
                current_app.logger.info("[TEXTBOOK IMPORTER] Gemini JSON parsed after LaTeX escape sanitize.")
                return True, "ok_sanitized", fixed_text
            except json.JSONDecodeError as second_error:
                _log_gemini_json_parse_failed_after_sanitize(first_error, second_error, cleaned_text)
                return False, "json_decode_failed", cleaned_text

    for attempt in range(1, max_retries + 1):
        try:
            if not hasattr(model, "generate_content"):
                current_app.logger.error(f"[_call_gemini_with_retry] invalid model type: {type(model).__name__}")
                return None

            if queue is not None:
                queue.put(f"INFO: {context_message} attempt {attempt}/{max_retries}")

            generation_config = {
                "temperature": 0.2,
                "max_output_tokens": 65536,
            }
            if parse_json:
                generation_config["response_mime_type"] = "application/json"

            response = model.generate_content(
                analysis_prompt,
                generation_config=generation_config,
            )

            raw_text = getattr(response, "text", "")
            result_text = str(raw_text or "").strip()
            if result_text:
                current_app.logger.info(f"Gemini response length = {len(result_text)}")
                current_app.logger.info(f"first 300 chars = {result_text[:300]}")
                current_app.logger.info(f"last 300 chars = {result_text[-300:]}")
                if parse_json:
                    is_valid_json, fail_reason, _ = _validate_json_completeness(result_text)
                    if is_valid_json:
                        return result_text
                    if queue is not None and fail_reason == "missing_closing_brace":
                        queue.put("WARNING: Gemini JSON missing closing brace")
                    if queue is not None:
                        queue.put("WARNING: Gemini JSON validation failed, retrying")
                    if attempt >= max_retries:
                        raise RuntimeError("Gemini JSON failed after retries")
                    time.sleep(retry_delay * attempt)
                    continue
                return result_text

            candidates = getattr(response, "candidates", None)
            if candidates:
                parts = []
                for cand in candidates:
                    content = getattr(cand, "content", None)
                    if not content:
                        continue
                    for p in getattr(content, "parts", []) or []:
                        t = getattr(p, "text", None)
                        if t:
                            parts.append(t)
                merged = "\n".join(parts).strip()
                if merged:
                    current_app.logger.info(f"Gemini response length = {len(merged)}")
                    current_app.logger.info(f"first 300 chars = {merged[:300]}")
                    current_app.logger.info(f"last 300 chars = {merged[-300:]}")
                    if parse_json:
                        is_valid_json, fail_reason, _ = _validate_json_completeness(merged)
                        if is_valid_json:
                            return merged
                        if queue is not None and fail_reason == "missing_closing_brace":
                            queue.put("WARNING: Gemini JSON missing closing brace")
                        if queue is not None:
                            queue.put("WARNING: Gemini JSON retry")
                        if attempt >= max_retries:
                            raise RuntimeError("Gemini JSON failed after retries")
                        time.sleep(retry_delay * attempt)
                        continue
                    return merged

            raise RuntimeError("Gemini output merge failed")

        except ResourceExhausted as e:
            if attempt >= max_retries:
                err_type = type(e).__name__
                err_msg = str(e) or repr(e)
                tb = traceback.format_exc()
                current_app.logger.error(f"_call_gemini_with_retry 發生錯誤: [{err_type}] {err_msg}\n{tb}")
                if queue is not None:
                    queue.put(f"ERROR: Gemini 呼叫失敗: [{err_type}] {err_msg}")
                raise
            time.sleep(retry_delay * attempt)

        except Exception as e:
            err_type = type(e).__name__
            err_msg = str(e) or repr(e)
            tb = traceback.format_exc()

            current_app.logger.error(f"_call_gemini_with_retry 發生錯誤: [{err_type}] {err_msg}\n{tb}")
            if queue is not None:
                queue.put(f"ERROR: Gemini 呼叫失敗: [{err_type}] {err_msg}")

            raise

def call_gemini_for_toc(content_by_page, curriculum_info, queue):
    """Call Gemini to extract chapter/section TOC."""
    message = "--- 開始 AI 目錄萃取流程 ---"
    current_app.logger.info(message)
    queue.put(f"INFO: {message}")

    prompt = "TOC extraction prompt"
    try:
        from core.ai_analyzer import get_model
        model = get_model()
        # ==============================================================================
        # 【型態解鎖】將 content_by_page 字典中的所有頁面純文字依序串接，防止 TypeError
        # ==============================================================================
        full_toc_text = "\n".join(str(v or "") for _k, v in sorted((content_by_page or {}).items()))
        response = _call_gemini_with_retry(model, prompt + "\n" + full_toc_text, queue=queue)
        # ==============================================================================
        return response
    except Exception as e:
        current_app.logger.error(f"call_gemini_for_toc failed: {e}")
        return None

def call_gemini_for_analysis(content_by_page, curriculum_info, queue, page_analysis_payload=None, import_policy=None):
    """Call Gemini to analyze extracted textbook content."""
    message = "--- 正在執行 AI 內容解析 ---"
    current_app.logger.info(message)
    queue.put(f"INFO: {message}")

    # ==========================
    # 1. 設定基礎 Prompt
    # ==========================
    prompt_jh_kangxuan = "Analyze textbook content and return JSON."

    # ==========================
    # 2. 龍騰/普高專用 Prompt（章節標題規則）
    # ==========================
    prompt_sh_longteng = "Analyze textbook content and return JSON."

    # ==========================
    # 3. 行內公式與結構 Prompt
    # ==========================

    prompt_generic = "Analyze textbook content and return JSON."

    # ==============================================================================
    # 【結構化合約】技術高中數學B1-B4 特化強型態 Prompt 與標準 JSON 範例
    # ==============================================================================
    prompt_vh_mathB4 = (
        "您是一位精通技術高中（高職）數學教材架構的專家。請分析輸入的課本標準文本內容，"
        "精確切分出內文中的觀念引導、獨立例題、隨堂練習與進階練習題。"
        "對於每一道題目，請完整保留其 LaTeX 公式（如 $...$ 或 $$...$$），切勿省略、截斷或自行發明題幹。"
        "您必須嚴格按照給定的 JSON 結構範例回傳資料，不允許自創任何外層或內層欄位。"
    )

    json_example_vh_mathB = """
{
  "chapters": [
    {
      "chapter_title": "第1章 坐標系與函數圖形",
      "sections": [
        {
          "section_title": "1-4 一元二次不等式",
          "concepts": [
            {
              "concept_name": "一元二次不等式的解法",
              "concept_en_id": "QuadraticInequalitiesSolution",
              "concept_paragraph": "探討 ax^2+bx+c > 0 形式的不等式解法與判別式 D 的關係。",
              "examples": [
                {
                  "id": "1",
                  "title": "例題1",
                  "source_description": "例題1",
                  "problem_text": "解不等式 $x^2 - 5x + 6 > 0$。",
                  "correct_answer": "$x > 3$ 或 $x < 2$",
                  "detailed_solution": "因式分解得 $(x-2)(x-3) > 0$，故解為 $x > 3$ 或 $x < 2$。"
                }
              ],
              "practice_questions": [
                {
                  "id": "1",
                  "title": "隨堂練習1",
                  "source_description": "隨堂練習1",
                  "problem_text": "解不等式 $x^2 - x - 2 \\\\le 0$。",
                  "correct_answer": "$-1 \\\\le x \\\\le 2$",
                  "detailed_solution": "因式分解得 $(x-2)(x+1) \\\\le 0$，故解為 $-1 \\\\le x \\\\le 2$。"
                }
              ]
            }
          ]
        }
      ]
    }
  ]
}
"""

    json_example_vh_mathB_metadata_only = """
{
  "chapters": [
    {
      "chapter_title": "第1章 坐標系與函數圖形",
      "sections": [
        {
          "section_title": "1-4 一元二次不等式",
          "concepts": [
            {
              "concept_name": "一元二次不等式的解法",
              "concept_en_id": "QuadraticInequalitiesSolution",
              "concept_paragraph": "",
              "examples": [
                {
                  "id": "1",
                  "title": "例題1",
                  "source_description": "例題1",
                  "problem_text": "例題1",
                  "correct_answer": "",
                  "detailed_solution": ""
                }
              ],
              "practice_questions": [
                {
                  "id": "1",
                  "title": "隨堂練習1",
                  "source_description": "隨堂練習1",
                  "problem_text": "隨堂練習1",
                  "correct_answer": "",
                  "detailed_solution": ""
                }
              ]
            }
          ]
        }
      ]
    }
  ]
}
"""
    # ==============================================================================

    curriculum = curriculum_info.get('curriculum', '').strip()
    publisher = curriculum_info.get('publisher', '').strip()
    volume = str(curriculum_info.get('volume', '')).strip()
    subject, vol_num = parse_volume(volume)
    is_vocational_mathb = curriculum == 'vocational' and subject == 'B'
    debug_message = (
        f"DEBUG: curriculum='{curriculum}', publisher='{publisher}', volume='{volume}', "
        f"parsed_subject='{subject}', parsed_volume={vol_num}"
    )
    current_app.logger.info(debug_message)
    queue.put(debug_message)

    import_policy = dict(import_policy or {})
    docx_formula_source_mode = str(import_policy.get("docx_formula_source_mode", "auto_detect") or "auto_detect").strip()
    section_code = str(import_policy.get("section_code", "") or "").strip() or None
    converted_metadata_only = docx_formula_source_mode == "converted_docx_latex"

    if converted_metadata_only and is_vocational_mathb:
        base_prompt = (
            "您是一位技高數學B教材結構分析專家。本批為 converted_docx_latex 匯入："
            "題幹與 LaTeX 已由 DOCX 決定性掃描補回。您只需輸出章節結構與每題 metadata，"
            "禁止重寫完整題幹、答案或詳解。"
        )
        queue.put("INFO: use vocational mathB converted_docx_latex metadata-only prompt")
    elif curriculum == 'junior_high' and publisher == 'kangxuan':
        base_prompt = prompt_jh_kangxuan
        queue.put("INFO: use junior_high kangxuan prompt")
    elif is_vocational_mathb:
        base_prompt = prompt_vh_mathB4
        queue.put(f"INFO: 已載入技高數學{subject}{vol_num} 專用提示詞範本")
    elif curriculum == 'sh_longteng' or (curriculum == 'general' and publisher == 'longteng'):
        base_prompt = prompt_sh_longteng
        queue.put("INFO: use longteng/general prompt")
    else:
        base_prompt = prompt_generic
        queue.put("INFO: use generic prompt")

    try:
        model = get_model("architect")
    except Exception as e:
        err_type = type(e).__name__
        err_msg = str(e) or repr(e)
        tb = traceback.format_exc()

        current_app.logger.error(f"AI 分析失敗: [{err_type}] {err_msg}\n{tb}")
        if "Gemini API Key" in err_msg or "API_KEY" in err_msg:
            queue.put("ERROR: Missing Gemini API Key.")
        else:
            queue.put(f"ERROR: AI 分析失敗: [{err_type}] {err_msg}")
        return None
    if page_analysis_payload:
        blocks = []
        for k in sorted(page_analysis_payload.keys(), key=lambda x: int(x)):
            p = page_analysis_payload[k]
            block = (
                f"--- Page {k} ---\n"
                f"[RAW PDF TEXT]\n{p.get('raw_text','')}\n\n"
                f"[NORMALIZED TEXT]\n{p.get('normalized_text','')}\n\n"
                f"[VISION OCR TEXT]\n{p.get('vision_ocr_text') or ''}\n\n"
                f"[FORMULA WARNINGS]\n{', '.join(p.get('formula_warnings', [])) or 'none'}\n"
            )
            blocks.append(block)
        full_content = "\n".join(blocks)
    else:
        raw_extracted = "\n".join(str(v or "") for _k, v in sorted((content_by_page or {}).items()))
        if converted_metadata_only:
            full_content = build_converted_docx_latex_gemini_outline_payload(
                raw_extracted,
                section_code=section_code,
            )
            queue.put("INFO: converted_docx_latex metadata-only Gemini payload (title inventory)")
        else:
            full_content = "\n".join([f"--- Page {k} ---\n{v}" for k, v in content_by_page.items()])

    if converted_metadata_only:
        json_example = json_example_vh_mathB_metadata_only if is_vocational_mathb else json_example_vh_mathB_metadata_only
    else:
        json_example = json_example_vh_mathB if is_vocational_mathb else "{}"

    converted_latex_prompt_rules = ""
    if converted_metadata_only:
        title_rules = (
            "【題目標題對齊規則】\n"
            "1. 每一筆 examples / practice_questions 必須同時填 title 與 source_description，且兩者完全相同。\n"
            "2. 例題：例題1、例題2…；隨堂練習：隨堂練習1…；章節習題：如 1-4習題 基礎題1。\n"
            "3. 統測題逐題列出（如 105統測A），禁止合併成 bucket。\n"
            "4. problem_text 僅填與 title 相同之短字串；correct_answer、detailed_solution 填 \"\"。\n"
            "5. 禁止多行 problem_text；題幹由系統自 DOCX 補回。\n"
        )
        if is_vocational_mathb:
            title_rules += (
                "6. 小節碼須與 section_title 一致；自我評量、基礎題、進階題區域須依原文標題。\n"
                "7. 課文說明放 concept_paragraph，勿當成獨立題目。\n"
            )
        sa_outline = detect_chapter_self_assessment_context(raw_extracted)
        if sa_outline:
            prefix = sa_outline.get("title_prefix", "CH1自我評量")
            title_rules += (
                f"8. 本章為章末自我評量題庫：每題 title/source_description 必須為 inventory canonical_title"
                f"（格式：{prefix} 題N）。\n"
                "9. 每題僅能放在 examples，source_type 必須為 self_assessment；practice_questions 留空。\n"
                "10. section_title 依題目所屬小節（如 1-1 數線與絕對值）；勿自創 concept_en_id，"
                "應使用該小節既有技能英文名。\n"
            )
        converted_latex_prompt_rules = CONVERTED_DOCX_LATEX_JSON_RULES + "\n" + title_rules

    # ==============================================================================
    # 【安全字元防線】修復全形符號、羅馬數字與畸形空白換行，防止 JSON 轉義爆炸
    # ==============================================================================
    if 'full_content' in locals() and full_content:
        # 1. 修正全形減號、特殊長橫線、破折號，統一替換為標準半形減號 '-'，防範 LaTeX 轉義語法打架
        full_content = re.sub(r'[－—―─–]', '-', full_content)
        
        # 2. 修正全形引號與全形逗號等高位元組字元，降低 Gemini 封裝 JSON 字串時的轉義負擔
        full_content = re.sub(r'[“”]', '"', full_content)
        full_content = re.sub(r'[‘’]', "'", full_content)
        full_content = re.sub(r'[，]', ',', full_content)
        
        # 3. 將常見的羅馬數字字元（Ⅰ, Ⅱ, Ⅲ...）替換為標準英文字母，防止高位元組字元破壞 JSON 閉合
        roma_map = {
            'Ⅰ': 'I', 'Ⅱ': 'II', 'Ⅲ': 'III', 'Ⅳ': 'IV', 'Ⅴ': 'V',
            'Ⅵ': 'VI', 'Ⅶ': 'VII', 'Ⅷ': 'VIII', 'Ⅸ': 'IX', 'Ⅹ': 'X'
        }
        for roma, eng in roma_map.items():
            full_content = full_content.replace(roma, eng)
            
        # 4. 修正帶圈數字/序號，替換為帶括號的標準數字 (1), (2), (3)...，確保 Gemini 在生成 keys 時名稱一致
        circle_map = {
            '①': '(1)', '②': '(2)', '③': '(3)', '④': '(4)', '⑤': '(5)',
            '⑥': '(6)', '⑦': '(7)', '⑧': '(8)', '⑨': '(9)', '⑩': '(10)'
        }
        for cir, num in circle_map.items():
            full_content = full_content.replace(cir, num)
            
        # 5. 強效清理：消除段落中多餘且畸形的空白與大段空行，壓縮為單一空白，防止 whitespace 溢出造成格式毀損
        # 這裡我們只壓縮三個以上的換行，保留單/雙換行
        full_content = re.sub(r'\n{3,}', '\n\n', full_content)

    # 動態組裝：必須帶入 base_prompt、JSON 範例、LaTeX 規則與全文，避免僅送死字串造成幻覺目錄
    analysis_prompt = (
        f"{base_prompt}\n\n"
        f"【請嚴格依照以下 JSON 範例格式結構輸出，嚴禁自行發明論文或無關的目錄結構】\n"
        f"{json_example}\n\n"
        f"{converted_latex_prompt_rules}\n\n"
        f"{'【以下為題目標題清單與章節提示（metadata-only，勿重寫題幹）】' if converted_metadata_only else '【以下是需要您切分、LaTeX化並結構化解析的課本標準文本內容】'}\n"
        f"{full_content}"
    )

    try:
        ai_response = _call_gemini_with_retry(
            model, 
            analysis_prompt, 
            queue, 
            context_message="執行教材內容結構化解析作業",
            parse_json=True
        )
        return ai_response
    except Exception as e:
        err_type = type(e).__name__
        err_msg = str(e) or repr(e)
        tb = traceback.format_exc()

        current_app.logger.error(f"AI 分析失敗: [{err_type}] {err_msg}\n{tb}")
        queue.put(f"ERROR: AI 分析失敗: [{err_type}] {err_msg}")
        return None

def parse_ai_response(ai_data_or_string, queue):
    """Parse AI response payload into dict."""
    if isinstance(ai_data_or_string, dict):
        return ai_data_or_string
    text = str(ai_data_or_string or "").strip()
    if not text:
        return None
    try:
        return safe_load_gemini_json(text)
    except Exception:
        if queue is not None:
            queue.put("ERROR: parse_ai_response failed")
        return None


def snake_to_pascal_case(snake_case_string):
    """Convert snake_case or kebab-case to PascalCase."""
    if not snake_case_string:
        return ""
    return ''.join(word.capitalize() for word in re.split('_|-', str(snake_case_string)))

def clean_skill_en_name(raw_en_name, queue=None):
    """Clean skill english name and keep PascalCase suffix if present."""
    if not raw_en_name:
        return ""
    match = re.search(r'[A-Z]', raw_en_name)
    if match:
        start_index = match.start()
        return raw_en_name[start_index:]
    return raw_en_name


def _formula_context_label(path):
    return " / ".join(str(p) for p in path if p is not None and str(p).strip())


def _normalize_extracted_content_math(content_by_page, queue=None, docx_formula_source_mode: str = ""):
    """Normalize extracted page text before Gemini sees OCR/PDF math artifacts."""
    if not isinstance(content_by_page, dict):
        return content_by_page

    normalized_pages = {}
    for page_no, page_text in content_by_page.items():
        if not isinstance(page_text, str):
            normalized_pages[page_no] = page_text
            continue

        if docx_formula_source_mode == "converted_docx_latex":
            normalized_pages[page_no] = page_text
            continue
        check = detect_suspicious_formula(page_text)
        normalized_text = normalize_math_text(page_text)
        if check.get("is_suspicious"):
            reasons = ",".join(check.get("reasons", []))
            log_msg = f"[FORMULA CHECK] suspicious formula detected in extracted page={page_no} reasons={reasons}"
            current_app.logger.warning(log_msg)
            if queue is not None:
                queue.put(f"WARN: {log_msg}")

        if normalized_text != page_text:
            current_app.logger.info(
                f"[FORMULA NORMALIZE] extracted_page={page_no} before={page_text[:120]!r} after={normalized_text[:120]!r}"
            )

        normalized_pages[page_no] = normalized_text

    return normalized_pages


def _normalize_imported_math_value(
    value,
    *,
    section_title="",
    source_description="",
    field_name="",
    queue=None,
    docx_formula_source_mode: str = "",
):
    if value is None or not isinstance(value, str):
        return value, None

    suspicious = detect_suspicious_formula(value)
    if docx_formula_source_mode == "converted_docx_latex":
        if has_app_context():
            current_app.logger.info(
                f"[FORMULA NORMALIZE SKIP] converted_docx_latex_preserve_latex=true field={field_name}"
            )
        return value, {
            "is_suspicious": bool(suspicious.get("reasons")),
            "reasons": list(dict.fromkeys(suspicious.get("reasons", []))),
            "suggestions": suspicious.get("suggestions", []),
            "normalized_preview": value,
        }

    normalized = normalize_math_text(value)
    suspicious_after = detect_suspicious_formula(normalized)
    reasons = list(dict.fromkeys(suspicious.get("reasons", []) + suspicious_after.get("reasons", [])))

    if reasons:
        label = _formula_context_label([section_title, source_description, field_name])
        log_msg = f"[FORMULA CHECK] suspicious formula detected in {label} reasons={reasons}"
        current_app.logger.warning(log_msg)
        if queue is not None:
            queue.put(f"WARN: {log_msg}")

    if normalized != value:
        current_app.logger.info(
            f"[FORMULA NORMALIZE] field={field_name} before={value[:120]!r} after={normalized[:120]!r}"
        )

    return normalized, {
        "is_suspicious": bool(reasons),
        "reasons": reasons,
        "suggestions": suspicious.get("suggestions", []) + suspicious_after.get("suggestions", []),
        "normalized_preview": normalized,
    }


def _normalize_parsed_textbook_math(parsed_data, queue=None, docx_formula_source_mode: str = ""):
    """Normalize known textbook JSON text fields before DB persistence."""
    if not isinstance(parsed_data, dict):
        return parsed_data

    for chapter in parsed_data.get("chapters", []) or []:
        if not isinstance(chapter, dict):
            continue
        for section in chapter.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            section_title = section.get("section_title", "")
            for concept in section.get("concepts", []) or []:
                if not isinstance(concept, dict):
                    continue

                for key in ("concept_description", "concept_paragraph"):
                    if isinstance(concept.get(key), str):
                        concept[key], check = _normalize_imported_math_value(
                            concept[key],
                            section_title=section_title,
                            source_description=concept.get("concept_name", ""),
                            field_name=key,
                            queue=queue,
                            docx_formula_source_mode=docx_formula_source_mode,
                        )
                        if check and check.get("is_suspicious"):
                            concept["needs_review"] = True
                            concept["parse_warning"] = ";".join(check.get("reasons", []))

                for ex in concept.get("examples", []) or []:
                    if not isinstance(ex, dict):
                        continue
                    source_description = ex.get("source_description", "example")
                    for key in (
                        "problem_text",
                        "problem",
                        "correct_answer",
                        "answer",
                        "detailed_solution",
                        "solution",
                        "hint",
                        "hints",
                    ):
                        if isinstance(ex.get(key), str):
                            ex[key], check = _normalize_imported_math_value(
                                ex[key],
                                section_title=section_title,
                                source_description=source_description,
                                field_name=key,
                                queue=queue,
                                docx_formula_source_mode=docx_formula_source_mode,
                            )
                            if check and check.get("is_suspicious"):
                                ex["needs_review"] = True
                                existing_warning = str(ex.get("parse_warning", "") or "").strip()
                                reasons = ";".join(check.get("reasons", []))
                                ex["parse_warning"] = ";".join(filter(None, [existing_warning, reasons]))
                    for sq in ex.get("sub_questions", []) or []:
                        if not isinstance(sq, dict):
                            continue
                        for key in ("problem", "answer", "solution"):
                            if isinstance(sq.get(key), str):
                                sq[key], _ = _normalize_imported_math_value(
                                    sq[key],
                                    section_title=section_title,
                                    source_description=source_description,
                                    field_name=f"sub_questions.{key}",
                                    queue=queue,
                                    docx_formula_source_mode=docx_formula_source_mode,
                                )

                for practice in concept.get("practice_questions", []) or []:
                    if not isinstance(practice, dict):
                        continue
                    for key in ("question", "problem_text", "solution", "answer", "hint", "hints"):
                        if isinstance(practice.get(key), str):
                            practice[key], check = _normalize_imported_math_value(
                                practice[key],
                                section_title=section_title,
                                source_description="practice",
                                field_name=key,
                                queue=queue,
                                docx_formula_source_mode=docx_formula_source_mode,
                            )
                            if check and check.get("is_suspicious"):
                                practice["needs_review"] = True
                                practice["parse_warning"] = ";".join(check.get("reasons", []))
                    for sq in practice.get("sub_questions", []) or []:
                        if not isinstance(sq, dict):
                            continue
                        for key in ("problem", "answer", "solution"):
                            if isinstance(sq.get(key), str):
                                sq[key], _ = _normalize_imported_math_value(
                                    sq[key],
                                    section_title=section_title,
                                    source_description="practice",
                                    field_name=f"sub_questions.{key}",
                                    queue=queue,
                                    docx_formula_source_mode=docx_formula_source_mode,
                                )

    return parsed_data


def _first_non_empty_str(mapping, keys):
    for key in keys:
        value = mapping.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


ALLOWED_SOURCE_TYPES = {
    "textbook_example",
    "in_class_practice",
    "chapter_exercise",
    "basic_exercise",
    "advanced_exercise",
    "self_assessment",
    "exam_practice",
    "generated_question",
    "student_uploaded",
    "textbook_practice",
}

CHAPTER_EXERCISE_TYPES = {
    "chapter_exercise",
    "basic_exercise",
    "advanced_exercise",
}

_SUPERSCRIPT_TRANS = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻", "0123456789+-")
_SUBSCRIPT_TRANS = str.maketrans("₀₁₂₃₄₅₆₇₈₉₊₋", "0123456789+-")


def normalize_fill_blank_artifacts(text: str) -> tuple[str, dict]:
    original = str(text or "")
    out = original
    log = {"changed": False, "reasons": []}

    # Instruction semantic normalization first.
    before = out
    out = re.sub(r"(求|解|計算)\s*(?:\(\s*\)|_{2,})\s*", r"\1 ", out)
    if out != before:
        log["changed"] = True
        log["reasons"].append("normalized fill-blank instruction text")

    # Generic blank symbols/slots normalization.
    blank_patterns = [
        r"[?﹦＿]",
        r"\(\s*\)",
        r"_+",
        r"_{2,}",
    ]
    before = out
    for pat in blank_patterns:
        out = re.sub(pat, "[BLANK]", out)
    out = re.sub(r"(?:\s*\[BLANK\]\s*){2,}", " [BLANK] ", out)
    out = re.sub(r"\s+", " ", out).strip()
    if out != before:
        log["changed"] = True
        log["reasons"].append("normalized fill blank symbol to [BLANK]")

    return out, log


def _contains_perm_comb_formula(text: str) -> bool:
    t = str(text or "")
    return bool(
        re.search(
            r"(?:"
            r"\{\s*\}\s*\^\s*\{?\s*\d+\s*\}?\s*[PC]\s*_\s*\{?\s*\d+\s*\}?|"
            r"[PC]\s*\^\s*\{?\s*\d+\s*\}?\s*_\s*\{?\s*\d+\s*\}?|"
            r"[PC]\s*_\s*\{?\s*\d+\s*\}?\s*\^\s*\{?\s*\d+\s*\}?|"
            r"[⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻]+\s*[PC]\s*[₀₁₂₃₄₅₆₇₈₉₊₋]+|"
            r"[PC]\s*[₀₁₂₃₄₅₆₇₈₉₊₋]+\s*[⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻]+"
            r")",
            t,
        )
    )

def is_answer_blank_placeholder_context(raw_block: str, problem_text: str) -> bool:
    block = str(raw_block or "")
    text = str(problem_text or "")
    combined = f"{block}\n{text}"
    instruction_hit = bool(re.search(r"(求|解|計算|試求)", combined))
    formula_hit = _contains_perm_comb_formula(combined)
    placeholder_count = len(re.findall(r"\[FORMULA_IMAGE_\d+\]|\[WORD_EQUATION_UNPARSED\]", block))
    subq_count = len(re.findall(r"\(\s*\d+\s*\)", combined))
    formula_with_placeholder = bool(
        re.search(
            r"(?:[PC].{0,24}(?:\[\s*FORMULA_IMAGE_\d+\]|\[\s*WORD_EQUATION_UNPARSED\])|"
            r"(?:\[\s*FORMULA_IMAGE_\d+\]|\[\s*WORD_EQUATION_UNPARSED\]).{0,24}[PC])",
            block,
        )
    )
    placeholder_close_to_subq = bool(subq_count and placeholder_count and abs(placeholder_count - subq_count) <= 1)
    return bool(instruction_hit and formula_hit and (formula_with_placeholder or placeholder_close_to_subq))


def normalize_permutation_combination_notation(
    text: str, *, volume: str = "", section_title: str = ""
) -> tuple[str, dict]:
    original = str(text or "")
    out = original
    log = {"changed": False, "reasons": []}
    coord_guard = _is_b1_coordinate_context(volume=volume, section_title=section_title, problem_text=original)
    if coord_guard:
        guarded = _normalize_coordinate_point_notation(out)
        if guarded != out:
            log["changed"] = True
            log["reasons"].append("normalized coordinate point notation in B1 coordinate context")
            out = guarded

    # e.g., ⁷P₃ / ⁷C₃
    def _replace_pre(match):
        n = match.group(1).translate(_SUPERSCRIPT_TRANS)
        op = match.group(2)
        r = match.group(3).translate(_SUBSCRIPT_TRANS)
        return f"{op}^{{{n}}}_{{{r}}}"

    out = re.sub(r"([⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻]+)\s*([PC])\s*([₀₁₂₃₄₅₆₇₈₉₊₋]+)", _replace_pre, out)

    # e.g., P₃⁷ / C₃⁷
    def _replace_post(match):
        op = match.group(1)
        r = match.group(2).translate(_SUBSCRIPT_TRANS)
        n = match.group(3).translate(_SUPERSCRIPT_TRANS)
        return f"{op}^{{{n}}}_{{{r}}}"

    out = re.sub(r"([PC])\s*([₀₁₂₃₄₅₆₇₈₉₊₋]+)\s*([⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻]+)", _replace_post, out)

    # e.g., P^7_3 / C^7_3
    out = re.sub(r"\b([PC])\s*\^\s*\{?\s*([0-9]+)\s*\}?\s*_\s*\{?\s*([0-9]+)\s*\}?", r"\1^{\2}_{\3}", out)

    # e.g., P_3^7 / C_3^7
    out = re.sub(r"\b([PC])\s*_\s*\{?\s*([0-9]+)\s*\}?\s*\^\s*\{?\s*([0-9]+)\s*\}?", r"\1^{\3}_{\2}", out)

    # e.g., {}^{7}P_{3} / {}^7P_3 / { }^{7} P_{3}
    out = re.sub(
        r"(?:\{\s*\}\s*)?\^\s*\{?\s*([0-9]+)\s*\}?\s*([PC])\s*_\s*\{?\s*([0-9]+)\s*\}?",
        r"\2^{\1}_{\3}",
        out,
    )

    # e.g., P(5,3) / C(8,2) (numeric params only; keep probability forms like P(A), P(A|B))
    # In B1 coordinate context, C(1,2)/P(a,b) should remain point notation.
    if not coord_guard:
        out = re.sub(r"\b([PC])\s*\(\s*([0-9]+)\s*,\s*([0-9]+)\s*\)", r"\1^{\2}_{\3}", out)

    if out != original:
        log["changed"] = True
        log["reasons"].append("normalized permutation notation to P^{n}_{r}/C^{n}_{r}")
    return out, log


def normalize_probability_event_notation(text: str) -> tuple[str, dict]:
    original = str(text or "")
    out = original
    log = {"changed": False, "reasons": []}

    token = r"[A-Za-z](?:'|\\prime)?"
    inner_re = re.compile(
        rf"^\s*({token})(?:\s*(\\cup|\\cap|\\setminus|\||-|或|且)\s*({token}))?\s*$"
    )
    outer_re = re.compile(r"P\s*\(\s*([^()]+?)\s*\)")

    def _canon_token(tok: str) -> str:
        t = str(tok or "").strip()
        t = re.sub(r"\s+", "", t)
        t = t.replace("\\prime", "'")
        return t

    def _already_inline_math(src: str, start: int, end: int) -> bool:
        left = src[max(0, start - 8):start]
        right = src[end:end + 8]
        return bool(re.search(r"\\\(\s*$", left) and re.search(r"^\s*\\\)", right))

    def _replace(match: re.Match) -> str:
        nonlocal out
        start, end = match.span()
        if _already_inline_math(out, start, end):
            return match.group(0)
        inner = match.group(1)
        m = inner_re.match(inner)
        if not m:
            return match.group(0)
        lhs = _canon_token(m.group(1))
        op = m.group(2)
        rhs = _canon_token(m.group(3)) if m.group(3) else None
        if not op:
            return rf"\(P({lhs})\)"
        op_map = {
            "∪": r"\cup",
            "∩": r"\cap",
            r"\cup": r"\cup",
            r"\\cup": r"\cup",
            r"\cap": r"\cap",
            r"\\cap": r"\cap",
            r"\setminus": r"\setminus",
            r"\\setminus": r"\setminus",
        }
        op_canon = op_map.get(op, op)
        return rf"\(P({lhs} {op_canon} {rhs})\)"

    out = outer_re.sub(_replace, out)

    if out != original:
        log["changed"] = True
        log["reasons"].append("normalized probability event notation to inline LaTeX")
    return out, log


def _is_subsection_heading_line(line: str) -> bool:
    t = str(line or "").strip()
    return bool(re.search(r"^\s*\d+\s*-\s*\d+\s*\.\s*\d+\s*[^\s].*$", t))


def repair_missing_single_variable_text(problem_text: str) -> tuple[str, dict]:
    text = str(problem_text or "")
    repair = {"applied": False, "symbol": None, "reason": ""}
    if not text.strip():
        repair["reason"] = "empty_text"
        return text, repair

    missing_slots = []
    if re.search(r"設\s*\?", text):
        missing_slots.append("subject")
    if re.search(r"=\s*\?", text):
        missing_slots.append("value")
    if not missing_slots:
        repair["reason"] = "no_missing_variable_slot"
        return text, repair

    symbols = set(re.findall(r"(?<![A-Za-z])[A-Za-z](?![A-Za-z])", text))
    candidates = sorted(sym for sym in symbols if sym.isalpha() and sym.lower() in "nmxyab")
    if len(candidates) != 1:
        repair["reason"] = "non_unique_candidate_variable"
        return text, repair

    sym = candidates[0]
    fixed = text
    if "subject" in missing_slots:
        fixed = re.sub(r"設\s*\?", f"設 {sym} ", fixed, count=1)
    if "value" in missing_slots:
        fixed = re.sub(r"=\s*\?", f"= {sym} ", fixed, count=1)

    if fixed != text:
        repair["applied"] = True
        repair["symbol"] = sym
        repair["reason"] = f"filled missing variable with unique symbol {sym}"
    else:
        repair["reason"] = "pattern_not_replaced"
    return fixed, repair


def validate_problem_block_purity(problem: dict) -> dict:
    if not isinstance(problem, dict):
        return problem
    text = str(problem.get("problem_text", "") or problem.get("problem", "") or "").strip()
    if not text:
        return problem

    lines = [ln.strip() for ln in text.splitlines() if str(ln or "").strip()]
    if any(_is_subsection_heading_line(ln) for ln in lines):
        problem["needs_review"] = True
        problem["block_boundary_error"] = True

    explanation_cues = _NON_QUESTION_EXPLANATION_CUES + ("補充",)
    problem_verbs = _QUESTION_VERBS
    explanation_hits = sum(1 for cue in explanation_cues if cue in text)
    has_problem_verb = any(v in text for v in problem_verbs)
    if explanation_hits >= 2 and not has_problem_verb:
        problem["needs_review"] = True
        problem["likely_concept_explanation"] = True

    formula_placeholders = (
        r"\[FORMULA_MISSING\]",
        r"\[FORMULA_IMAGE_\d+\]",
        r"\[WORD_EQUATION_UNPARSED\]",
        r"\[BLOCK_IMAGE\]",
    )
    has_formula_gap = any(re.search(p, text) for p in formula_placeholders)
    readable_text = re.sub(r"\[[A-Z_0-9]+\]", " ", text)
    readable_text = re.sub(r"\s+", " ", readable_text).strip()
    if has_formula_gap and (len(readable_text) < 8 or not has_problem_verb):
        problem["needs_review"] = True
        problem["formula_missing"] = True
        logs = problem.get("repair_log", [])
        if not isinstance(logs, list):
            logs = [str(logs)]
        logs.append("marked formula_missing due to formula placeholder")
        problem["repair_log"] = logs

    # ??憿憿凳/?賊?畾撩嚗????蝻箏撘?銝?葫?批捆
    set_question_hint = bool(re.search(r"(集合|子集|交集|聯集)", text))
    has_choice_labels = len(re.findall(r"\([A-Da-d]\)", text)) >= 2
    has_set_notation = bool(re.search(r"[\{\}U∩∪]|A\s*[∩∪\\]\s*B|A\s*-\s*B", text))
    if set_question_hint and has_choice_labels and not has_set_notation:
        problem["needs_review"] = True
        problem["formula_missing"] = True
        logs = problem.get("repair_log", [])
        if not isinstance(logs, list):
            logs = [str(logs)]
        logs.append("marked formula_missing due to missing set notation/options")
        problem["repair_log"] = logs

    if "[BLANK]" in text:
        problem["has_answer_blank"] = True
        problem["question_format"] = "fill_blank"
        logs = problem.get("repair_log", [])
        if not isinstance(logs, list):
            logs = [str(logs)]
        logs.append("preserved [BLANK] as fill blank, not formula missing")
        problem["repair_log"] = logs

    if problem.get("skill_id") and problem.get("block_boundary_error"):
        problem["needs_review"] = True
        problem["skill_boundary_mismatch"] = True

    if re.search(r"(表格|統計表|列表|table)", text, flags=re.IGNORECASE):
        has_table_payload = bool(re.search(r"\d", text) and re.search(r"[|｜、，,;]", text))
        if not has_table_payload and not re.search(r"\[[A-Z_]*TABLE[A-Z_0-9]*\]", text):
            problem["needs_review"] = True
            problem["needs_table_review"] = True
            logs = problem.get("repair_log", [])
            if not isinstance(logs, list):
                logs = [str(logs)]
            logs.append("table-dependent question without enough table content")
            problem["repair_log"] = logs

    return problem


def classify_practice_source_bucket(source_type: str) -> str:
    st = str(source_type or "").strip().lower()
    if st == "in_class_practice":
        return "in_class_practice"
    if st in CHAPTER_EXERCISE_TYPES:
        return "chapter_exercise"
    if st == "self_assessment":
        return "self_assessment"
    if st == "exam_practice":
        return "exam_practice"
    return "other_practice"


def get_question_title(item: dict) -> str:
    if not isinstance(item, dict):
        return ""
    return _first_non_empty_str(
        item,
        ("practice_title", "example_title", "title", "display_name", "name", "source_title", "source_description"),
    )


_SECTION_EXPOSITION_TITLE_EXACT = frozenset({
    "觀念整理", "觀念", "整理", "說明", "重點", "補充",
    "觀念說明", "整理說明", "章節說明", "概念說明",
    "課文內容", "課文說明", "課文",
})

_SECTION_EXPOSITION_TITLE_RE = re.compile(
    r"^(?:觀念整理|觀念說明|整理說明|章節說明|概念說明|課文內容|課文說明|課文|觀念\s*\d*|說明\s*\d*)$"
)


def _is_section_exposition_title(title: str) -> bool:
    """Return True when *title* looks like a section-exposition label (not a real question).

    Catches Gemini-generated source_description values like '課文內容' that represent
    narrative textbook passages rather than examples, exercises, or practices.
    These must NOT be saved as textbook_example records.
    """
    t = str(title or "").strip()
    if not t:
        return False
    # Exact known exposition labels
    if t in _SECTION_EXPOSITION_TITLE_EXACT:
        return True
    # Regex: starts with exposition keywords, no digits or question-type suffix
    if _SECTION_EXPOSITION_TITLE_RE.match(t):
        return True
    return False


_NON_SKILL_BUCKET_NAMES = frozenset({
    "習題",
    "章節介紹",
    "排列組合概論",
    "統測歷屆試題",
    "歷屆試題",
    "章節習題",
    "基礎題",
    "進階題",
    "自我評量",
    "綜合練習",
    "綜合題",
    "課文內容",
    "觀念整理",
})

_NON_SKILL_BUCKET_EN_IDS = frozenset({
    "chapterintroduction",
    "exercises",
    "practice",
    "review",
    "defaultsectionconcept",
    "unknown",
    "sectionexercises",
    "chapterexercises",
    "basictopic",
    "advancedtopic",
    "selfassessment",
    "comprehensivepractice",
    "textbookcontent",
    "conceptsummary",
    "unifiedexam",
    "pastexam",
})


def is_non_skill_bucket(concept_name: str, clean_en_id: str) -> bool:
    """判斷 AI 回傳的 concept 是否為非技能分類容器（不建立 SkillInfo）。"""
    name = (concept_name or "").strip()
    name_compact = re.sub(r"\s+", "", name)
    en_id = re.sub(r"[^a-zA-Z0-9]", "", (clean_en_id or "")).lower()

    if name in _NON_SKILL_BUCKET_NAMES or name_compact in _NON_SKILL_BUCKET_NAMES:
        return True
    return en_id in _NON_SKILL_BUCKET_EN_IDS


def normalize_source_type_by_title(item: dict, default_source_type: str = "textbook_example") -> str:
    title = get_question_title(item)
    raw_source_type = str(item.get("source_type", "") or "").strip().lower() if isinstance(item, dict) else ""
    reason = ""

    # Section exposition titles are not importable question records.
    if _is_section_exposition_title(title):
        normalized = "section_exposition"
        reason = "title_is_section_exposition"
        if isinstance(item, dict):
            item["source_type"] = normalized
            try:
                current_app.logger.info(
                    f"[SOURCE TYPE] title={title!r} normalized_source_type={normalized} reason={reason}"
                )
            except Exception:
                pass
        return normalized

    if re.search(r"^\s*例(?:題)?\s*\d+", title):
        normalized = "textbook_example"
        reason = "title_matches_example"
    elif "隨堂練習" in title:
        normalized = "in_class_practice"
        reason = "title_contains_practice"
    elif ("自我評量" in title) or ("評量" in title) or bool(re.search(r"自我評量\s*\d*", title)):
        normalized = "self_assessment"
        reason = "title_contains_self_assessment"
    elif any(k in title for k in ("統測", "學測")):
        normalized = "exam_practice"
        reason = "title_contains_exam"
    elif "基礎題" in title:
        normalized = "basic_exercise"
        reason = "title_contains_basic_exercise"
    elif "進階題" in title:
        normalized = "advanced_exercise"
        reason = "title_contains_advanced_exercise"
    elif "習題" in title:
        normalized = "textbook_practice"
        reason = "title_contains_practice"
    elif "例題" in title:
        normalized = "textbook_practice"
        reason = "title_contains_example"
    elif "綜合題" in title:
        normalized = "chapter_exercise"
        reason = "title_contains_chapter_exercise"
    elif raw_source_type in ALLOWED_SOURCE_TYPES:
        normalized = raw_source_type
        reason = "item_source_type_allowed"
    else:
        fallback = str(default_source_type or "textbook_example").strip().lower() or "textbook_example"
        if raw_source_type and raw_source_type not in ALLOWED_SOURCE_TYPES:
            normalized = "textbook_practice"
            reason = "invalid_source_type_fallback_textbook_practice"
            if isinstance(item, dict):
                item["needs_review"] = True
        else:
            normalized = fallback
            reason = "default_source_type"

    if isinstance(item, dict):
        item["source_type"] = normalized
        try:
            current_app.logger.info(
                f"[SOURCE TYPE] title={title or '<empty>'} raw_source_type={raw_source_type or 'None'} "
                f"normalized_source_type={normalized} reason={reason}"
            )
        except Exception:
            pass
    return normalized


_SOURCE_TYPE_QUALITY_RANK = {
    "basic_exercise": 60,
    "advanced_exercise": 58,
    "chapter_exercise": 56,
    "in_class_practice": 54,
    "exam_practice": 52,
    "self_assessment": 50,
    "textbook_example": 48,
    "textbook_practice": 46,
}


def _normalize_title_for_dedupe(title: str) -> str:
    t = re.sub(r"\s+", "", str(title or "").strip())
    t = re.sub(r"^靘??!憿?(\d+)$", r"靘?\1", t)
    return t


def _build_intra_import_dedupe_key(section_title: str, source_type: str, title: str) -> str:
    sec = _normalize_title_for_dedupe(section_title)
    st = str(source_type or "").strip().lower()
    tt = _normalize_title_for_dedupe(title)
    return f"{sec}|{st}|{tt}"


def _score_intra_import_item(item: dict, source_type: str = "", title: str = "", q_assets: dict | None = None) -> int:
    st = str(source_type or "").strip().lower()
    score = _SOURCE_TYPE_QUALITY_RANK.get(st, 0)
    text = str(item.get("problem_text", "") or item.get("problem", "") or "")
    if re.search(r"\[FORMULA_IMAGE_\d+\]", text):
        score += 100
    if "[FORMULA_MISSING]" in text:
        score += 10
    if item.get("needs_formula_review") is True:
        score += 8
    score += min(len(text), 400) // 4
    if isinstance(item.get("formula_assets"), list) and item.get("formula_assets"):
        score += 80
    if q_assets:
        title_assets = _lookup_docx_question_assets(title, q_assets) or []
        formula_assets = [a for a in title_assets if str(a.get("media_kind", "")) == "formula_asset"]
        if formula_assets:
            score += 70 + min(len(formula_assets), 20)
    return score


def _dedupe_intra_import_section_items(section_data: dict, q_assets: dict | None = None) -> int:
    """Deduplicate repeated questions within one section before DB write.

    Returns merged duplicate count.
    """
    concepts = section_data.get("concepts", []) if isinstance(section_data, dict) else []
    if not isinstance(concepts, list) or not concepts:
        return 0

    # Gather all candidates with location pointers.
    all_entries = []
    for ci, concept in enumerate(concepts):
        if not isinstance(concept, dict):
            continue
        for bucket_name, default_st in (("examples", "textbook_example"), ("practice_questions", "in_class_practice")):
            arr = concept.get(bucket_name, [])
            if not isinstance(arr, list):
                continue
            for ii, item in enumerate(arr):
                if not isinstance(item, dict):
                    continue
                title = get_question_title(item) or ""
                source_type = normalize_source_type_by_title(item, default_source_type=default_st)
                key = _build_intra_import_dedupe_key(
                    section_data.get("section_title", ""),
                    source_type,
                    title,
                )
                score = _score_intra_import_item(item, source_type=source_type, title=title, q_assets=q_assets)
                all_entries.append(
                    {
                        "concept_idx": ci,
                        "bucket": bucket_name,
                        "item_idx": ii,
                        "item": item,
                        "title": title,
                        "source_type": source_type,
                        "key": key,
                        "score": score,
                    }
                )

    if not all_entries:
        return 0

    # Pick winner per dedupe key.
    best_by_key = {}
    for e in all_entries:
        prev = best_by_key.get(e["key"])
        if prev is None or e["score"] > prev["score"]:
            best_by_key[e["key"]] = e

    keep_locs = {
        (e["concept_idx"], e["bucket"], e["item_idx"])
        for e in best_by_key.values()
    }
    removed = 0
    for ci, concept in enumerate(concepts):
        if not isinstance(concept, dict):
            continue
        for bucket_name in ("examples", "practice_questions"):
            arr = concept.get(bucket_name, [])
            if not isinstance(arr, list):
                continue
            new_arr = []
            for ii, item in enumerate(arr):
                if (ci, bucket_name, ii) in keep_locs:
                    new_arr.append(item)
                else:
                    removed += 1
            concept[bucket_name] = new_arr
    return removed


def dedupe_intra_import_parsed_data(parsed_data: dict, q_assets: dict | None = None) -> tuple[dict, int]:
    """Deduplicate duplicate questions within the same AI JSON payload."""
    if not isinstance(parsed_data, dict):
        return parsed_data, 0
    total_removed = 0
    for chapter in parsed_data.get("chapters", []) or []:
        if not isinstance(chapter, dict):
            continue
        for section in chapter.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            total_removed += _dedupe_intra_import_section_items(section, q_assets=q_assets)
    return parsed_data, total_removed


def _normalize_sub_questions(raw_sub_questions):
    normalized = []
    if not isinstance(raw_sub_questions, list):
        return normalized
    for idx, item in enumerate(raw_sub_questions, start=1):
        if not isinstance(item, dict):
            continue
        sq_problem_raw = _first_non_empty_str(item, ("problem_text", "problem", "question"))
        sq_problem, _ = standardize_problem_latex(sq_problem_raw)
        sq_problem, _ = normalize_probability_event_notation(sq_problem)
        normalized.append(
            {
                "label": _first_non_empty_str(item, ("label", "index", "no", "number")) or str(idx),
                "problem": sq_problem,
                "answer": _first_non_empty_str(item, ("correct_answer", "answer")),
                "solution": _first_non_empty_str(item, ("detailed_solution", "solution")),
            }
        )
    return normalized


def _render_sub_questions_problem(problem_text, sub_questions):
    if not sub_questions:
        return (problem_text or "").strip()
    lines = [str(problem_text or "").strip()] if str(problem_text or "").strip() else []
    for sq in sub_questions:
        label = str(sq.get("label", "") or "").strip()
        p = str(sq.get("problem", "") or "").strip()
        if p:
            lines.append(f"({label}) {p}" if label else p)
    return "\n".join(lines).strip()


def _render_sub_questions_answer(answer_text, sub_questions):
    if not sub_questions:
        return (answer_text or "").strip()
    parts = []
    for sq in sub_questions:
        label = str(sq.get("label", "") or "").strip()
        ans = str(sq.get("answer", "") or "").strip()
        if ans:
            parts.append(f"({label}) {ans}" if label else ans)
    return "\n".join(parts) if parts else (answer_text or "").strip()


def _render_sub_questions_solution(solution_text, sub_questions):
    if not sub_questions:
        return (solution_text or "").strip()
    lines = [str(solution_text or "").strip()] if str(solution_text or "").strip() else []
    for sq in sub_questions:
        label = str(sq.get("label", "") or "").strip()
        s = str(sq.get("solution", "") or "").strip()
        p = str(sq.get("problem", "") or "").strip()
        a = str(sq.get("answer", "") or "").strip()
        if s:
            lines.append(f"({label}) {s}" if label else s)
        elif p or a:
            lines.append(f"({label}) {p} = {a}".strip() if label else f"{p} = {a}".strip())
    return "\n".join(lines).strip()


def _normalize_textbook_question_structure(parsed_data, queue=None):
    """
    Normalize AI output structure so examples/practice_questions are consistently separable.
    - examples: independent textbook examples
    - practice_questions: independent in-class practices / exercises
    - backward compatibility: example.followup_practices -> concept.practice_questions
    """
    # ==============================================================================
    # 【完全體原子級型態對齊閘門】相容純文字字串題目的硬核降維重組
    # ==============================================================================
    if isinstance(parsed_data, list):
        parsed_data = {"chapters": parsed_data}
    elif isinstance(parsed_data, dict):
        if "chapters" not in parsed_data:
            if "subsections" in parsed_data and "sections" not in parsed_data:
                parsed_data["sections"] = parsed_data["subsections"]
            parsed_data = {
                "chapters": [
                    {
                        "chapter_title": parsed_data.get("chapter_title") or parsed_data.get("section_title") or "預設第一章",
                        "sections": [parsed_data],
                    }
                ]
            }

    if isinstance(parsed_data, dict):
        for ch in parsed_data.get("chapters", []) or []:
            if not isinstance(ch, dict):
                continue

            ch_title_raw = _first_non_empty_str(ch, ("chapter_title", "chapter", "name", "title"))
            if ch_title_raw:
                ch["chapter_title"] = ch_title_raw

            for sec in ch.get("sections", []) or []:
                if not isinstance(sec, dict):
                    continue

                sec_title_raw = _first_non_empty_str(sec, ("section_title", "section", "title", "name"))
                if sec_title_raw:
                    sec["section_title"] = sec_title_raw

                # 收集所有可能藏在單元內部的子單元塊或內容大雜燴
                sub_blocks = []
                for sub_key in ("subsections", "sub_sections", "concepts", "concept_list", "content"):
                    if sub_key in sec and isinstance(sec[sub_key], list):
                        sub_blocks.extend(sec[sub_key])

                if not sub_blocks:
                    sub_blocks.append(sec)

                # 清空並重新構建 concepts 陣列
                sec["concepts"] = []

                for block in sub_blocks:
                    if not isinstance(block, dict):
                        continue

                    c_name = _first_non_empty_str(block, ("concept_name", "title", "concept", "name")) or sec_title_raw or "未命名概念"
                    c_en_id = _first_non_empty_str(block, ("concept_en_id", "concept_id", "id", "en_id")) or "Unknown"

                    concept_node = {
                        "concept_name": c_name,
                        "concept_en_id": c_en_id,
                        "concept_paragraph": _first_non_empty_str(block, ("concept_paragraph", "paragraph", "description")) or c_name,
                        "examples": [],
                        "practice_questions": [],
                    }

                    # 收集該子區塊中所有潛在的題目來源陣列
                    raw_examples = []
                    raw_practices = []

                    for ex_key in ("examples", "example_list", "example_problems", "example"):
                        if ex_key in block and isinstance(block[ex_key], list):
                            raw_examples.extend(block[ex_key])

                    for pr_key in (
                        "practice_questions", "practice_list", "practice_problems", "practice",
                        "advanced_problems", "advanced_questions", "basic_problems", "basic_questions",
                        "exercises", "exercise_list", "self_assessment_questions", "self_assessment",
                        "advanced", "basic", "questions", "question_list",
                    ):
                        if pr_key in block and isinstance(block[pr_key], list):
                            raw_practices.extend(block[pr_key])

                    # exercise_group 僅在無 questions 子陣列時，才把整塊 dict 當單題處理（避免與字串題重複）
                    if block.get("type") == "exercise_group" and "text" in block:
                        qlist = block.get("questions")
                        if not isinstance(qlist, list) or not qlist:
                            raw_practices.append(block)

                    # ==============================================================================
                    # 【全局外層題目無差別回收防線】確保丟在最外層的習題、例題 100% 歸流
                    # ==============================================================================
                    if block is sub_blocks[-1]:
                        # 回收外層練習題/進階題/基礎題/習題
                        for pr_key in (
                            "advanced", "basic", "exercises", "practice_questions",
                            "advanced_problems", "advanced_questions", "basic_problems",
                            "basic_questions", "questions", "question_list", "subsections", "sub_sections",
                        ):
                            if pr_key in sec and isinstance(sec[pr_key], list):
                                # 排除已經變成節點的 sections 本身，其餘純資料陣列全數倒入
                                if pr_key not in ("sections", "subsections", "sub_sections") or not any(
                                    isinstance(x, dict) and "section_title" in x for x in sec[pr_key]
                                ):
                                    raw_practices.extend(sec[pr_key])

                        # 回收外層漏網的例題變體
                        for ex_key in ("examples", "example_list", "example_problems", "example"):
                            if ex_key in sec and isinstance(sec[ex_key], list):
                                raw_examples.extend(sec[ex_key])
                    # ==============================================================================

                    # 5. 原子級型態校正：將純字串題目與變體字典統一重組為標準結構
                    for src_arr, target_bucket in ((raw_examples, "examples"), (raw_practices, "practice_questions")):
                        for item in src_arr:
                            standard_item = {}

                            if isinstance(item, str):
                                # 型態死鎖突破：純文字字串題型處理 (例如 "9. 成本...")
                                txt_stripped = item.strip()
                                # 利用正則匹配開頭的題號
                                match_num = re.match(r"^(\d+)\s*[\.\、\)\t\s]\s*(.*)$", txt_stripped)
                                if match_num:
                                    q_id = match_num.group(1)
                                    q_text = match_num.group(2).strip()
                                else:
                                    q_id = "1"
                                    q_text = txt_stripped

                                label_prefix = "例題" if target_bucket == "examples" else "習題"
                                standard_item = {
                                    "problem_text": q_text,
                                    "title": f"{label_prefix}{q_id}",
                                    "source_description": f"{label_prefix}{q_id}",
                                }
                            elif isinstance(item, dict):
                                # 常規字典型態校正
                                prob_raw = _first_non_empty_str(item, ("problem_text", "problem", "question", "question_text", "stem", "text"))
                                if not prob_raw and "questions" in item and isinstance(item["questions"], list):
                                    # 處理特殊巢狀結構，如果字典裡又包了一層純字串陣列
                                    for sub_str in item["questions"]:
                                        if isinstance(sub_str, str):
                                            src_arr.append(sub_str)
                                    continue

                                if not prob_raw:
                                    continue

                                standard_item = dict(item)
                                standard_item["problem_text"] = prob_raw
                                ans_raw = _first_non_empty_str(item, ("correct_answer", "answer", "ans", "right_answer"))
                                if ans_raw:
                                    standard_item["correct_answer"] = ans_raw
                                sol_raw = _first_non_empty_str(item, ("detailed_solution", "solution", "sol", "explanation", "analysis"))
                                if sol_raw:
                                    standard_item["detailed_solution"] = sol_raw

                                title_raw = _first_non_empty_str(item, ("title", "source_description", "source_title", "name"))
                                if not title_raw and "id" in item:
                                    prefix_label = "例題" if target_bucket == "examples" else "習題"
                                    title_raw = f"{prefix_label}{item['id']}"
                                if title_raw:
                                    standard_item["title"] = title_raw
                                    standard_item["source_description"] = title_raw
                            else:
                                continue

                            concept_node[target_bucket].append(standard_item)

                    if concept_node["examples"] or concept_node["practice_questions"]:
                        sec["concepts"].append(concept_node)

    # ==============================================================================

    if not isinstance(parsed_data, dict):
        return parsed_data

    for chapter in parsed_data.get("chapters", []) or []:
        if not isinstance(chapter, dict):
            continue
        for section in chapter.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            for concept in section.get("concepts", []) or []:
                if not isinstance(concept, dict):
                    continue

                normalized_examples = []
                extracted_practices = []

                for ex in concept.get("examples", []) or []:
                    if not isinstance(ex, dict):
                        continue

                    example_title = _first_non_empty_str(ex, ("example_title", "source_description", "title"))
                    problem_text = _first_non_empty_str(ex, ("problem_text", "problem", "question"))
                    answer = _first_non_empty_str(ex, ("correct_answer", "answer"))
                    solution = _first_non_empty_str(ex, ("detailed_solution", "solution"))
                    source_type = normalize_source_type_by_title(ex, default_source_type="textbook_example")
                    source_page = ex.get("source_page", ex.get("page"))
                    page_index = ex.get("page_index")
                    sub_questions = _normalize_sub_questions(ex.get("sub_questions", []))

                    ex_normalized = dict(ex)
                    ex_normalized["source_description"] = example_title or ex_normalized.get("source_description", "例題")
                    if problem_text:
                        ex_normalized["problem_text"] = problem_text
                    if answer:
                        ex_normalized["correct_answer"] = answer
                    if solution:
                        ex_normalized["detailed_solution"] = solution
                    ex_normalized["source_type"] = source_type if source_type else "textbook_example"
                    ex_normalized["source_page"] = source_page if source_page is not None else None
                    ex_normalized["page_index"] = page_index if page_index is not None else None
                    if sub_questions:
                        ex_normalized["sub_questions"] = sub_questions
                        ex_normalized["problem_text"] = _render_sub_questions_problem(problem_text, sub_questions)
                        ex_normalized["correct_answer"] = _render_sub_questions_answer(answer, sub_questions)
                        ex_normalized["detailed_solution"] = _render_sub_questions_solution(solution, sub_questions)
                    if source_type == "textbook_example":
                        normalized_examples.append(ex_normalized)
                    else:
                        extracted_practices.append(ex_normalized)

                    for fp in ex.get("followup_practices", []) or []:
                        if not isinstance(fp, dict):
                            continue
                        p_title = _first_non_empty_str(fp, ("practice_title", "title", "source_description"))
                        p_problem = _first_non_empty_str(fp, ("problem_text", "problem", "question"))
                        p_answer = _first_non_empty_str(fp, ("correct_answer", "answer"))
                        p_solution = _first_non_empty_str(fp, ("detailed_solution", "solution"))
                        p_source_type = normalize_source_type_by_title(fp, default_source_type="in_class_practice")
                        linked_example_title = _first_non_empty_str(fp, ("linked_example_title",)) or ex_normalized["source_description"]
                        p_source_page = fp.get("source_page", fp.get("page"))
                        p_page_index = fp.get("page_index")

                        practice_item = dict(fp)
                        practice_item["source_description"] = p_title or "隨堂練習"
                        if p_problem:
                            practice_item["problem_text"] = p_problem
                        if p_answer:
                            practice_item["correct_answer"] = p_answer
                        if p_solution:
                            practice_item["detailed_solution"] = p_solution
                        practice_item["source_type"] = p_source_type
                        practice_item["linked_example_title"] = linked_example_title
                        practice_item["source_page"] = p_source_page if p_source_page is not None else None
                        practice_item["page_index"] = p_page_index if p_page_index is not None else None
                        if not practice_item.get("skill_id") and ex_normalized.get("skill_id"):
                            practice_item["skill_id"] = ex_normalized.get("skill_id")
                        extracted_practices.append(practice_item)

                normalized_practices = []
                for practice in (concept.get("practice_questions", []) or []) + extracted_practices:
                    if not isinstance(practice, dict):
                        continue
                    p_title = _first_non_empty_str(practice, ("practice_title", "source_description", "title"))
                    p_problem = _first_non_empty_str(practice, ("problem_text", "problem", "question"))
                    p_answer = _first_non_empty_str(practice, ("correct_answer", "answer"))
                    p_solution = _first_non_empty_str(practice, ("detailed_solution", "solution"))
                    p_source_type = normalize_source_type_by_title(practice, default_source_type="in_class_practice")
                    linked_example_title = _first_non_empty_str(practice, ("linked_example_title",))
                    p_source_page = practice.get("source_page", practice.get("page"))
                    p_page_index = practice.get("page_index")
                    sub_questions = _normalize_sub_questions(practice.get("sub_questions", []))

                    normalized_practice = dict(practice)
                    normalized_practice["source_description"] = p_title or normalized_practice.get("source_description", "隨堂練習")
                    if p_problem:
                        normalized_practice["problem_text"] = p_problem
                    if p_answer:
                        normalized_practice["correct_answer"] = p_answer
                    if p_solution:
                        normalized_practice["detailed_solution"] = p_solution
                    normalized_practice["source_type"] = p_source_type
                    if linked_example_title:
                        normalized_practice["linked_example_title"] = linked_example_title
                    normalized_practice["source_page"] = p_source_page if p_source_page is not None else None
                    normalized_practice["page_index"] = p_page_index if p_page_index is not None else None
                    if sub_questions:
                        normalized_practice["sub_questions"] = sub_questions
                        normalized_practice["problem_text"] = _render_sub_questions_problem(p_problem, sub_questions)
                        normalized_practice["correct_answer"] = _render_sub_questions_answer(p_answer, sub_questions)
                        normalized_practice["detailed_solution"] = _render_sub_questions_solution(p_solution, sub_questions)

                    normalized_practices.append(normalized_practice)

                concept["examples"] = normalized_examples
                concept["practice_questions"] = normalized_practices
                if isinstance(concept.get("self_assessment_questions"), list):
                    normalized_sa = []
                    for q in concept.get("self_assessment_questions", []) or []:
                        if not isinstance(q, dict):
                            continue
                        qn = dict(q)
                        qn["source_type"] = normalize_source_type_by_title(qn, default_source_type="self_assessment")
                        qn["source_page"] = q.get("source_page", q.get("page", None))
                        qn["page_index"] = q.get("page_index", None)
                        normalized_sa.append(qn)
                    concept["self_assessment_questions"] = normalized_sa
                    concept["practice_questions"].extend(normalized_sa)
                if isinstance(concept.get("exercises"), list):
                    normalized_exercises = []
                    for q in concept.get("exercises", []) or []:
                        if not isinstance(q, dict):
                            continue
                        qn = dict(q)
                        qn["source_type"] = normalize_source_type_by_title(qn, default_source_type="chapter_exercise")
                        normalized_exercises.append(qn)
                    concept["exercises"] = normalized_exercises
                    concept["practice_questions"].extend(normalized_exercises)

    return parsed_data


def import_outline_structure_only(parsed_data, curriculum_info, queue, source_file_path=None):
    """Import outline structure only into SkillCurriculum."""
    try:
        from models import db, SkillCurriculum
        current_app.logger.info(" -> [OutlineOnly] 開始建立單元大綱...")
        
        # 撘瑕皜 session 銝剔?隞颱? pending 霈嚗??autoflush 閫貊????SkillInfo ?航炊
        db.session.rollback()
        
        filename_meta = parse_textbook_filename_metadata(source_file_path) if source_file_path else {}
        volume_val = str(curriculum_info.get('volume', ''))
        curr_val = str(curriculum_info.get('curriculum', ''))
        is_vocational_mathb = (curr_val == 'vocational' and 'B' in volume_val)
        
        # 將結構對照關係儲存
        struct_map_obj = get_structure_map(curr_val, volume_val)

        chapters_created = 0
        sections_created = 0
        sections_updated = 0
        
        # ?冽餈質馱撌脰???蝡?嚗??銴?蝞?
        processed_chapters = set()

        chapters = parsed_data.get('chapters', [])
        for ch_data in chapters:
            raw_ch_title = ch_data.get('chapter_title', '未知章節').strip()
            
            # 單元大綱解析
            chapter_title = raw_ch_title
            
            sections = ch_data.get('sections', [])
            for sec_data in sections:
                sec_title = sec_data.get('section_title', '').strip()

                # Skip review / 自我評量 / 複習等非正式教學單元
                # sections and should not appear in the curriculum outline.
                if _is_review_section_title(sec_title):
                    current_app.logger.info(
                        f"[OutlineOnly] skip review section: {sec_title!r}"
                    )
                    continue

                # 對應到結構大綱
                sec_code = ""
                # ?岫敺?蝭璅??? 1-1, 1-2 蝑誨蝣?
                match_code = re.search(r'(\d+-\d+)', sec_title)
                if match_code:
                    sec_code = match_code.group(1)
                
                structure_meta = None
                if struct_map_obj and sec_code:
                    structure_meta = struct_map_obj.get_metadata(sec_code)
                
                # 瘙箏??蝯?蝭璅?
                final_ch_title = chapter_title
                if structure_meta and structure_meta.get('chapter_title'):
                    final_ch_title = structure_meta['chapter_title']
                
                # 瘙箏??蝯?蝭璅?
                final_sec_title = sec_title
                if structure_meta and structure_meta.get('section_title'):
                    final_sec_title = structure_meta['section_title']

                # 瘙箏? skill_id (???箇?揣撘?銝遣蝡?SkillInfo)
                # 雿輻 SectionTitle ?Ｙ???ID嚗???outline ?韌隞亥霅
                clean_sec_title = re.sub(r'[^a-zA-Z0-9]', '', final_sec_title)
                if not clean_sec_title:
                    clean_sec_title = "UnknownSection"
                temp_skill_id = f"outline_{curr_val}_{volume_val}_{clean_sec_title}"
                
                # 檢查該技能是否已存在
                existing_curr = SkillCurriculum.query.filter_by(
                    curriculum=curr_val,
                    volume=volume_val,
                    chapter=final_ch_title,
                    section=final_sec_title
                ).first()
                
                if not existing_curr:
                    new_curr = SkillCurriculum(
                        skill_id=temp_skill_id,
                        curriculum=curr_val,
                        grade=int(curriculum_info.get('grade', 10)),
                        volume=volume_val,
                        chapter=final_ch_title,
                        section=final_sec_title,
                        display_order=0 
                    )
                    db.session.add(new_curr)
                    sections_created += 1
                    if final_ch_title not in processed_chapters:
                        chapters_created += 1
                        processed_chapters.add(final_ch_title)
                else:
                    sections_updated += 1
        
        db.session.commit()
        return {
            "chapters_created": chapters_created,
            "sections_created": sections_created,
            "sections_updated": sections_updated
        }
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"[import_outline_structure_only] 失敗: {e}")
        raise e


def save_to_database(
    parsed_data,
    curriculum_info,
    queue,
    source_file_path=None,
    content_by_page=None,
    outline_only=False,
    import_policy=None,
    optional_enrich_pdf_path=None,
):
    """Save parsed AI content into database."""
    message = "正在將解析結果寫入資料庫..."
    current_app.logger.info(message)
    queue.put(f"INFO: {message}")
    import_policy = dict(import_policy or {})
    q_assets: dict = {}
    docx_formula_blocks: dict = {}
    skills_processed = 0
    curriculums_added = 0
    chapters_created = 0
    chapters_updated = 0
    sections_created = 0
    sections_updated = 0
    examples_added = 0
    practice_questions_imported = 0
    in_class_practices_imported = 0
    chapter_exercises_imported = 0
    self_assessments_imported = 0
    exam_practices_imported = 0
    other_practices_imported = 0
    practice_questions_needs_review = 0
    practice_questions_skipped = 0
    duplicates_skipped_count = 0
    updated_duplicates = 0
    processed_skill_ids = []
    detected_titles = []
    in_class_nums = []
    missing_image_questions = []
    docx_attached_count = 0
    docx_direct_display_images = 0
    docx_vector_images = 0
    docx_conversion_success = 0
    docx_conversion_failed = 0
    docx_conversion_pending = 0
    docx_formula_assets_count = 0
    docx_formula_needs_review_count = 0
    merge_guard_kept_existing = 0
    merge_guard_updated_incoming = 0
    docx_copied_to_question_assets = 0
    formula_asset_persist_cache: dict[str, str] = {}
    is_pdf_source = str(source_file_path or "").lower().endswith(".pdf")
    is_docx_source = str(source_file_path or "").lower().endswith((".docx", ".doc"))
    converted_latex_import = (
        str(import_policy.get("docx_formula_source_mode", "") or "").strip() == "converted_docx_latex"
        or bool((_DOCX_IMPORT_CONTEXT or {}).get("formula_assets_extraction_skipped"))
    )
    if is_docx_source:
        ctx = _DOCX_IMPORT_CONTEXT or {}
        q_assets = ctx.get("question_assets", {}) if isinstance(ctx, dict) else {}
        docx_formula_blocks = ctx.get("question_formula_blocks", {}) if isinstance(ctx, dict) else {}
        attached_asset_count = sum(len(v or []) for v in q_assets.values()) if isinstance(q_assets, dict) else 0
        current_app.logger.info(f"[DOCX IMAGE DEBUG] attached_asset_count={attached_asset_count}")
        if isinstance(q_assets, dict):
            for t, assets in q_assets.items():
                for a in (assets or []):
                    current_app.logger.info(
                        f"[DOCX IMAGE DEBUG] attached title={t} source_type={a.get('media_kind', 'unknown')} "
                        f"asset_type={a.get('asset_type', 'unknown')} path={a.get('path')}"
                    )
    parsed_data, intra_import_duplicates_merged = dedupe_intra_import_parsed_data(
        parsed_data,
        q_assets=q_assets if is_docx_source else None,
    )
    if intra_import_duplicates_merged > 0:
        current_app.logger.info(
            f"[INTRA IMPORT DEDUPE] merged_duplicates={intra_import_duplicates_merged}"
        )
        queue.put(f"INFO: [INTRA IMPORT DEDUPE] merged_duplicates={intra_import_duplicates_merged}")
    page_image_cache = {}
    auto_fill_threshold = float(import_policy.get("auto_fill_confidence_threshold", 0.85) or 0.85)
    auto_fill_threshold = max(0.0, min(1.0, auto_fill_threshold))
    def _count_latex_and_placeholder_records(data):
        latex_re = re.compile(r"\\\(|\\\)|\\\[|\\\]|\\(?:frac|sqrt|left|right)\b|[\^_]")
        placeholder_re = re.compile(r"\[FORMULA_IMAGE_\d+\]|\[FORMULA_MISSING\]")
        latex_count = 0
        placeholder_count = 0
        def _record_text_has_latex(*parts):
            joined = "\n".join(str(p or "") for p in parts)
            return bool(latex_re.search(joined))
        for ch in (data or {}).get("chapters", []) or []:
            for sec in (ch or {}).get("sections", []) or []:
                for concept in (sec or {}).get("concepts", []) or []:
                    concept_has_latex = _record_text_has_latex(
                        concept.get("concept_description", ""),
                        concept.get("concept_paragraph", ""),
                    )
                    if concept_has_latex:
                        latex_count += 1
                    for ex in (concept or {}).get("examples", []) or []:
                        txt = str((ex or {}).get("problem_text", "") or "")
                        if _record_text_has_latex(
                            ex.get("problem_text", ""),
                            ex.get("problem", ""),
                            ex.get("correct_answer", ""),
                            ex.get("answer", ""),
                            ex.get("detailed_solution", ""),
                            ex.get("solution", ""),
                        ):
                            latex_count += 1
                        if placeholder_re.search(txt):
                            placeholder_count += 1
                    for pq in (concept or {}).get("practice_questions", []) or []:
                        txt = str((pq or {}).get("problem_text", "") or "")
                        if _record_text_has_latex(
                            pq.get("question", ""),
                            pq.get("problem_text", ""),
                            pq.get("answer", ""),
                            pq.get("solution", ""),
                        ):
                            latex_count += 1
                        if placeholder_re.search(txt):
                            placeholder_count += 1
        return latex_count, placeholder_count

    records_with_latex, records_with_placeholder = _count_latex_and_placeholder_records(parsed_data)
    
    # [NEW] 檢驗章節數量與頁數限制
    parse_name = resolve_parse_filename_for_import(source_file_path or "", curriculum_info)
    filename_meta = parse_textbook_filename_metadata(parse_name) if parse_name else {}
    
    # [NEW] 教材結構大綱對齊
    structure_meta = None
    structure_alignment_failed = False
    volume_val = str(curriculum_info.get('volume', ''))
    curr_val = str(curriculum_info.get('curriculum', ''))
    if filename_meta.get('section_code') and curr_val == 'vocational' and 'B' in volume_val:
        struct_map = get_structure_map(curr_val, volume_val)
        if struct_map:
            structure_meta = struct_map.get_metadata(filename_meta['section_code'])
            if structure_meta:
                current_app.logger.info(f"[STRUCTURE_MAP] Aligned to: {structure_meta}")
            else:
                current_app.logger.warning(f"[STRUCTURE_MAP] No match for {filename_meta['section_code']}")
                structure_alignment_failed = True
    
    if filename_meta:
        current_app.logger.info(f"[FILENAME_META] parsed: {filename_meta}")

    import_scope = str((filename_meta or {}).get("source_scope") or "section_textbook").strip()
    if import_scope in ("chapter_self_assessment", "chapter_review"):
        import_scope = "chapter_self_assessment"
    elif (
        import_scope != "section_textbook"
        and (_DOCX_IMPORT_CONTEXT or {}).get("chapter_self_assessment_mode")
    ):
        import_scope = "chapter_self_assessment"
    current_app.logger.info(f"[IMPORT SCOPE] scope={import_scope}")
    if queue is not None:
        queue.put(f"INFO: [IMPORT SCOPE] scope={import_scope}")

    prefix_map = {
        'junior_high': 'jh_',
        'general': 'gh_',
        'vocational': 'vh_'
    }
    curriculum = curriculum_info.get('curriculum', '')
    volume = str(curriculum_info.get('volume', '')).strip()
    subject, vol_num = parse_volume(volume)
    is_vocational_math = curriculum == 'vocational' and subject is not None and vol_num is not None
    is_vocational_mathb = is_vocational_math and subject == 'B'
    prefix = prefix_map.get(curriculum, '')

    # [V2.6] ?擃摮?B 蝟餃?嚗?澆?朣???暺?
    preferred_anchor_title = ""
    if isinstance(structure_meta, dict):
        preferred_anchor_title = _norm_section_label(str(structure_meta.get("section_title") or ""))
    if not preferred_anchor_title and filename_meta:
        sc_meta = str(filename_meta.get("section_code") or "").strip()
        st_meta = str(filename_meta.get("section_title") or "").strip()
        if sc_meta and st_meta and sc_meta not in st_meta:
            preferred_anchor_title = _norm_section_label(f"{sc_meta} {st_meta}")
        else:
            preferred_anchor_title = _norm_section_label(st_meta or sc_meta)

    existing_anchor_section = None
    if is_vocational_mathb and filename_meta.get('section_code') and not outline_only:
        s_code = filename_meta['section_code']
        # 撠?Ｘ??桅?蝭暺?(靘? 1-1 xxx)
        anchor_rows = SkillCurriculum.query.filter(
            SkillCurriculum.curriculum == curriculum,
            SkillCurriculum.volume == volume,
            SkillCurriculum.section.like(f"{s_code} %"),
        ).all()
        existing_anchor_section = None
        outline_anchor_section = None
        picked, _pick_reason = _pick_anchor_curriculum_row(
            anchor_rows,
            preferred_section_title=preferred_anchor_title,
            section_code=s_code,
        )
        if picked is not None:
            existing_anchor_section = picked
        else:
            for row in anchor_rows or []:
                sid = str(getattr(row, "skill_id", "") or "").strip()
                if is_outline_skill_id(sid) and outline_anchor_section is None:
                    outline_anchor_section = row
            existing_anchor_section = outline_anchor_section

        if existing_anchor_section:
            anchor_sid = str(getattr(existing_anchor_section, "skill_id", "") or "")
            if is_outline_skill_id(anchor_sid) and is_section_textbook_import_scope(
                _DOCX_IMPORT_CONTEXT, filename_meta
            ):
                resolved_sid, resolved_src = resolve_section_textbook_skill_id(
                    curriculum,
                    volume,
                    s_code,
                    preferred_anchor_title,
                    existing_anchor_section=existing_anchor_section,
                    structure_meta=structure_meta,
                )
                if resolved_sid and is_official_vh_skill_id(resolved_sid):
                    anchor_sid = resolved_sid
                    current_app.logger.info(
                        f"[ALIGNMENT] Replaced outline anchor with official skill_id={anchor_sid} "
                        f"source={resolved_src}"
                    )
            elif (
                is_outline_skill_id(anchor_sid)
                and is_question_bank_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta)
                and not is_section_textbook_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta)
            ):
                resolved_sid, resolved_src = resolve_question_bank_section_skill_id(
                    curriculum,
                    volume,
                    s_code,
                    preferred_anchor_title,
                )
                if resolved_sid and is_official_vh_skill_id(resolved_sid):
                    anchor_sid = resolved_sid
                    current_app.logger.info(
                        f"[ALIGNMENT] Replaced outline anchor with official skill_id={anchor_sid} "
                        f"source={resolved_src}"
                    )
                else:
                    current_app.logger.warning(
                        f"[QUESTION BANK SKILL ALIGN WARNING] section_code={s_code} "
                        f"reason=fallback_outline existing_anchor_is_outline=true anchor_skill_id={anchor_sid}"
                    )
            elif is_outline_skill_id(anchor_sid) and is_section_textbook_import_scope(
                _DOCX_IMPORT_CONTEXT, filename_meta
            ):
                current_app.logger.warning(
                    "[SECTION TEXTBOOK SKILL ALIGN WARNING] section_code=%s section_title=%s "
                    "reason=outline_anchor_unresolved anchor_skill_id=%s",
                    s_code,
                    preferred_anchor_title,
                    anchor_sid,
                )
            current_app.logger.info(
                f"[ALIGNMENT] Found existing anchor section: {existing_anchor_section.section} "
                f"skill_id={anchor_sid}"
            )
            if not structure_meta:
                structure_meta = {}
            # 撘瑕撠?structure_meta 鋆??Ｘ?璅?鞈?嚗??蝥?銴遣蝡?璅?銝???
            structure_meta['chapter_title'] = existing_anchor_section.chapter
            structure_meta['section_title'] = existing_anchor_section.section
            # 璅?撠???嚗??蝥粥??filename_meta fallback
            structure_alignment_failed = False
        else:
            # 若無任何對應單元，則對齊失敗
            message = f"Missing aligned section for code {s_code}; fallback may be used."
            current_app.logger.warning(f"[ALIGNMENT] {message}")
            queue.put(f"WARNING: {message}")
            structure_alignment_failed = True

    def _extract_title_number(text):
        if not text:
            return None
        match = re.search(r'(\d+)', str(text))
        return int(match.group(1)) if match else None

    def _sq_field_for_hash(value):
        raw = str(value or "")
        if converted_latex_import:
            text = raw
        else:
            text = normalize_math_text(raw)
        return re.sub(r"\s+", " ", text.replace("[FORMULA_MISSING]", "[FORMULA_TOKEN]")).strip()

    def _normalize_problem_hash(problem_text, sub_questions=None, source_type="", title=""):
        if converted_latex_import:
            normalized = re.sub(r"\s+", " ", str(problem_text or "")).strip()
        else:
            normalized = normalize_math_text(str(problem_text or ""))
        # Treat [FORMULA_IMAGE_N] and [FORMULA_MISSING] as same placeholder class
        # to prevent duplicate imports caused only by placeholder token style.
        normalized = re.sub(r"\[FORMULA_IMAGE_\d+\]", "[FORMULA_TOKEN]", normalized)
        normalized = normalized.replace("[FORMULA_MISSING]", "[FORMULA_TOKEN]")
        normalized = normalized.replace("[WORD_EQUATION_UNPARSED]", "[FORMULA_TOKEN]")
        normalized = re.sub(r"\s+", " ", normalized).strip()
        sq_norm = []
        for sq in sub_questions or []:
            if not isinstance(sq, dict):
                continue
            sq_norm.append(
                {
                    "label": str(sq.get("label", "") or "").strip(),
                    "problem": _sq_field_for_hash(sq.get("problem", "")),
                    "answer": _sq_field_for_hash(sq.get("answer", "")),
                    "solution": _sq_field_for_hash(sq.get("solution", "")),
                }
            )
        payload = {
            "source_type": str(source_type or "").strip().lower(),
            "title": str(title or "").strip(),
            "problem": normalized,
            "sub_questions": sq_norm,
        }
        return hashlib.sha1(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest()[:16]

    def validate_problem_formula_not_hallucinated(item_title: str, item: dict, problem_text: str, raw_block: str):
        text = str(problem_text or "")
        block = str(raw_block or "")
        has_placeholder = bool(
            re.search(r"\[FORMULA_MISSING\]|\[FORMULA_IMAGE_\d+\]|\[WORD_EQUATION_UNPARSED\]|\[UNREADABLE_FORMULA\]", block)
        )
        if is_docx_source:
            current_app.logger.info(f"[DOCX FORMULA SOURCE] title={item_title} raw_block={block[:240]}")
        has_problem_placeholder = bool(
            re.search(r"\[FORMULA_MISSING\]|\[FORMULA_IMAGE_\d+\]|\[WORD_EQUATION_UNPARSED\]|\[UNREADABLE_FORMULA\]", text)
        )
        if has_placeholder:
            current_app.logger.warning(f"[DOCX FORMULA WARNING] formula placeholder found title={item_title}")
        if has_placeholder or has_problem_placeholder:
            item["needs_review"] = True
            item["needs_formula_review"] = True
            item["formula_missing"] = True
        if has_placeholder and _contains_perm_comb_formula(text):
            has_ocr_source = bool(item.get("formula_ocr_source"))
            if not has_ocr_source:
                if is_answer_blank_placeholder_context(block, text):
                    normalized_text, blank_meta = normalize_fill_blank_artifacts(text)
                    item["has_answer_blank"] = True
                    item["question_format"] = "fill_blank"
                    item["needs_review"] = True
                    item["needs_formula_review"] = False
                    item["formula_missing"] = False
                    item["formula_hallucination_risk"] = False
                    logs = item.get("repair_log", [])
                    if not isinstance(logs, list):
                        logs = [str(logs)]
                    logs.append("formula placeholders treated as answer blanks")
                    logs.extend(blank_meta.get("reasons", []))
                    item["repair_log"] = logs
                    current_app.logger.info(
                        f"[DOCX FORMULA BLANK] title={item_title} treat placeholders as answer blank"
                    )
                    return normalized_text
                current_app.logger.warning(f"[DOCX FORMULA WARNING] formula missing before AI title={item_title}")
                fallback_text = text
                fallback_text = re.sub(
                    r"(?:\{\s*\}\s*\^\s*\{?\s*\d+\s*\}?\s*[PC]\s*_\s*\{?\s*\d+\s*\}?|"
                    r"[PC]\s*\^\s*\{?\s*\d+\s*\}?\s*_\s*\{?\s*\d+\s*\}?|"
                    r"[PC]\s*_\s*\{?\s*\d+\s*\}?\s*\^\s*\{?\s*\d+\s*\}?|"
                    r"\b[PC]\s*\(\s*\d+\s*,\s*\d+\s*\))",
                    "[FORMULA_MISSING]",
                    fallback_text,
                )
                item["needs_review"] = True
                item["needs_formula_review"] = True
                item["formula_missing"] = True
                item["formula_hallucination_risk"] = True
                item["parse_warning"] = "formula generated by AI without source"
                return fallback_text
        if has_placeholder and not _contains_perm_comb_formula(text):
            fallback_text = text
            if "[FORMULA_MISSING]" not in fallback_text:
                fallback_text = f"{fallback_text} [FORMULA_MISSING]".strip()
            item["needs_review"] = True
            item["needs_formula_review"] = True
            item["formula_missing"] = True
            return fallback_text
        return text

    def extract_formula_images_for_question_block(item_title: str):
        """Legacy hook kept for compatibility; formula OCR is metadata-only now.

        OCR fallback is handled in _build_docx_formula_assets_metadata after the
        formula asset has a readable png/jpeg display_path or converted_path.
        This helper intentionally does not return OCR text for problem_text
        replacement.
        """
        _ = item_title
        return []

    def _build_source_description(title, source_type, linked_example_title=None, needs_review=False, dedupe_hash="", section_context=None):
        title_text = str(title or "").strip() or "untitled"
        parts = [f"source_type={source_type}"]
        if section_context:
            parts.append(f"section={section_context}")
        if linked_example_title:
            parts.append(f"linked_example={linked_example_title}")
        if needs_review:
            parts.append("needs_review=true")
        if dedupe_hash:
            parts.append(f"dedupe={dedupe_hash}")
        return f"{title_text} [{' | '.join(parts)}]"

    def _infer_linked_example_title(practice_title, linked_example_title, saved_titles, needs_review):
        linked = str(linked_example_title or "").strip() or None
        review = bool(needs_review)
        if linked:
            return linked, review
        practice_num = _extract_title_number(practice_title)
        if practice_num is not None:
            inferred = f"例題{practice_num}"
            if inferred in saved_titles:
                return inferred, review
            if saved_titles:
                return saved_titles[-1], True
            return None, True
        if saved_titles:
            return saved_titles[-1], True
        return None, True

    def _build_image_metadata(
        question_title,
        question_text,
        chapter_title,
        section_title,
        source_type,
        question_code,
        force_has_image=False,
        image_description="",
        source_page=None,
        page_index=None,
        item_payload=None,
        concept_name="",
    ):
        has_formula_image_placeholder = bool(re.search(r"\[FORMULA_IMAGE_\d+\]", str(question_text or "")))
        img_dep = classify_image_dependency(
            question_text,
            image_hint=str(image_description or ""),
        )
        if not has_formula_image_placeholder and not img_dep.get("needs_image"):
            return _image_dependency_notes_metadata(img_dep)
        reason = image_description or detect_image_reason(question_text)
        current_app.logger.info(f"[QUESTION IMAGE] needs image title={question_title} source_page={source_page}")
        if queue is not None:
            queue.put(f"INFO: [QUESTION IMAGE] needs image title={question_title} source_page={source_page}")

        metadata = {
            "has_image": True,
            "needs_image_review": True,
            "image_assets": [],
        }
        if has_formula_image_placeholder:
            metadata["needs_formula_review"] = True
            metadata["formula_asset_type"] = "image_formula"
        infer_item = dict(item_payload or {})
        infer_item.setdefault("source_description", question_title)
        infer_item.setdefault("problem_text", question_text)
        infer_item.setdefault("has_image", bool(force_has_image))
        infer_item.setdefault("image_description", image_description or "")
        infer_item["source_page"] = source_page
        infer_item["page_index"] = page_index
        inferred_source_page, infer_reason = infer_source_page_for_question(
            infer_item,
            extracted_pages=content_by_page or {},
            section_title=section_title,
            concept_name=concept_name,
        )
        if inferred_source_page is not None and infer_reason != "explicit_source_page":
            current_app.logger.info(
                f"[QUESTION IMAGE] inferred source_page title={question_title} source_page={inferred_source_page} reason={infer_reason}"
            )
            if queue is not None:
                queue.put(
                    f"INFO: [QUESTION IMAGE] inferred source_page title={question_title} source_page={inferred_source_page} reason={infer_reason}"
                )

        if inferred_source_page is None:
            metadata["image_warning"] = "missing_source_page"
            current_app.logger.info(f"[QUESTION IMAGE] missing image asset title={question_title} reason=missing_source_page")
            if queue is not None:
                queue.put(f"INFO: [QUESTION IMAGE] missing image asset title={question_title} reason=missing_source_page")
            return metadata

        if not is_pdf_source:
            return metadata

        page_number = int(inferred_source_page)

        rel_dir, abs_dir, chapter_id, section_id = build_question_assets_dir(
            curriculum_info, chapter_title, section_title
        )
        filename = f"page_{int(page_number):03d}.png"
        abs_path = os.path.join(abs_dir, filename)
        rel_path = os.path.join(rel_dir, filename)
        cache_key = (os.path.abspath(str(source_file_path or "")), int(page_number), os.path.abspath(abs_dir))
        try:
            if cache_key in page_image_cache and os.path.exists(page_image_cache[cache_key]):
                reused_abs = page_image_cache[cache_key]
                rel_path = os.path.relpath(reused_abs, current_app.root_path)
                current_app.logger.info(f"[QUESTION IMAGE] reused page image path={rel_path}")
                if queue is not None:
                    queue.put(f"INFO: [QUESTION IMAGE] reused page image path={rel_path}")
            else:
                render_pdf_page_to_image(source_file_path, int(page_number) - 1, abs_path, dpi=200)
                page_image_cache[cache_key] = abs_path
                current_app.logger.info(f"[QUESTION IMAGE] rendered page image path={rel_path}")
                if queue is not None:
                    queue.put(f"INFO: [QUESTION IMAGE] rendered page image path={rel_path}")
            metadata["image_assets"].append(
                make_page_image_asset(
                    reason=reason,
                    rel_image_path=rel_path,
                    page_index=int(page_number) - 1,
                )
            )
            metadata["image_assets"][0]["source_page"] = int(page_number)
            metadata["image_assets"][0]["image_description"] = image_description or ""
            if infer_reason != "explicit_source_page":
                metadata["image_assets"][0]["source_page_inferred"] = True
                metadata["image_assets"][0]["source_page_infer_reason"] = infer_reason
                metadata["needs_image_review"] = True
        except Exception as e:
            current_app.logger.warning(f"[QUESTION IMAGE] render failed question={question_title} err={e}")
        return metadata

    def _build_docx_assets_metadata(
        question_title,
        chapter_title,
        section_title,
        source_type,
        question_text="",
        *,
        nearby_caption="",
        image_hint="",
    ):
        ctx = _DOCX_IMPORT_CONTEXT or {}
        q_assets = ctx.get("question_assets", {}) if isinstance(ctx, dict) else {}
        img_dep = classify_image_dependency(
            question_text,
            nearby_caption=nearby_caption,
            image_hint=image_hint,
        )
        all_candidates = _lookup_docx_question_assets(str(question_title or ""), q_assets)
        candidates = [a for a in (all_candidates or []) if str(a.get("media_kind", "image_asset")) == "image_asset"]
        if not img_dep.get("needs_image"):
            return _image_dependency_notes_metadata(img_dep)
        if not candidates:
            return {
                "has_image": True,
                "image_assets": [],
                "image_warning": "missing_docx_image_asset",
                "needs_review": True,
                "image_role": img_dep.get("image_role"),
                "image_dependency_reason": img_dep.get("reason"),
            }
        rel_dir = build_question_asset_dir(
            curriculum=curriculum_info.get("curriculum", "unknown"),
            publisher=curriculum_info.get("publisher", "unknown"),
            volume=curriculum_info.get("volume", "unknown"),
            chapter_title=chapter_title,
            section_title=section_title,
            source_filename=os.path.basename(str(source_file_path or "")),
        )
        abs_dir = os.path.join(current_app.root_path, rel_dir)
        os.makedirs(abs_dir, exist_ok=True)
        current_app.logger.info(f"[QUESTION ASSET] dir={rel_dir}")
        image_assets = []
        for idx, asset in enumerate(candidates, start=1):
            src_rel = str(asset.get("path") or "")
            ext = (os.path.splitext(src_rel)[1].lower() or ".bin").lstrip(".")
            problem_key = str(question_title or "") + "|" + str(source_type or "") + "|" + str(asset.get("block_index") or "")
            qhash = hashlib.sha1(problem_key.encode("utf-8")).hexdigest()[:8]
            filename = build_question_asset_filename(
                source_type=source_type,
                question_title=question_title,
                question_id_or_dedupe=qhash,
                fig_index=idx,
                ext=ext,
            )
            copied_abs = _copy_docx_asset_to_question_assets(src_rel, abs_dir, filename)
            if not copied_abs:
                continue
            rel_path = os.path.join(rel_dir, filename).replace("\\", "/")
            current_app.logger.info(
                f"[DOCX IMAGE] attached title={question_title} original={src_rel} question_asset={rel_path}"
            )
            needs_conv = ext in ("wmf", "emf")
            display_path = rel_path
            converted_path = None
            conv_error = None
            needs_review = False
            if needs_conv:
                png_filename = os.path.splitext(filename)[0] + ".png"
                png_abs = os.path.join(abs_dir, png_filename)
                png_rel = os.path.join(rel_dir, png_filename).replace("\\", "/")
                current_app.logger.info(f"[DOCX IMAGE] convert start input={copied_abs} output={png_abs}")
                ok, err = convert_vector_image_to_png(copied_abs, png_abs)
                if ok:
                    current_app.logger.info(f"[DOCX IMAGE] convert success output={png_abs}")
                    display_path = png_rel
                    converted_path = png_rel
                    needs_conv = False
                else:
                    current_app.logger.warning(f"[DOCX IMAGE WARNING] convert failed input={copied_abs} error={err}")
                    display_path = None
                    conv_error = err
                    needs_review = True
            image_assets.append(
                {
                    "asset_type": "word_embedded_image",
                    "source": "docx",
                    "path": rel_path,
                    "display_path": display_path,
                    "converted_path": converted_path,
                    "original_path": src_rel.replace("\\", "/"),
                    "content_type": asset.get("content_type", _guess_image_content_type(filename)),
                    "original_format": ext if ext in ("wmf", "emf") else None,
                    "needs_image_conversion": needs_conv,
                    "needs_image_review": needs_review,
                    "conversion_error": conv_error,
                    "image_attach_reason": asset.get("image_attach_reason", "image_inside_question_block"),
                }
            )
        if image_assets:
            return {
                "has_image": True,
                "image_assets": image_assets,
                "image_role": img_dep.get("image_role"),
                "image_dependency_reason": img_dep.get("reason"),
            }
        return {
            "has_image": True,
            "image_assets": [],
            "image_warning": "missing_docx_image_asset",
            "needs_review": True,
            "image_role": img_dep.get("image_role"),
            "image_dependency_reason": img_dep.get("reason"),
        }

    def _build_docx_formula_assets_metadata(question_title, question_text=""):
        ctx = _DOCX_IMPORT_CONTEXT or {}
        q_assets = ctx.get("question_assets", {}) if isinstance(ctx, dict) else {}
        all_candidates = _lookup_docx_question_assets(str(question_title or ""), q_assets)
        formula_candidates = [a for a in (all_candidates or []) if str(a.get("media_kind", "")) == "formula_asset"]
        has_formula_placeholder = bool(re.search(r"\[FORMULA_IMAGE_\d+\]|\[FORMULA_MISSING\]", str(question_text or "")))
        if not formula_candidates and has_formula_placeholder:
            # Global conservative fallback: if this question contains formula placeholder and
            # DOCX extracted any formula assets, do not leave formula_assets empty.
            global_formula_candidates = []
            if isinstance(q_assets, dict):
                for _k, _arr in q_assets.items():
                    for _a in (_arr or []):
                        if str(_a.get("media_kind", "")) == "formula_asset":
                            aa = dict(_a)
                            aa.setdefault("mapping_status", "global_formula_fallback")
                            global_formula_candidates.append(aa)
                            # keep fallback small and deterministic
                            if len(global_formula_candidates) >= 3:
                                break
                    if len(global_formula_candidates) >= 3:
                        break
            formula_candidates = global_formula_candidates
        if not formula_candidates:
            return None
        current_app.logger.info(f"[DOCX FORMULA ASSET] attached title={question_title} count={len(formula_candidates)}")
        rel_dir_base = build_question_asset_dir(
            curriculum=curriculum_info.get("curriculum", "unknown"),
            publisher=curriculum_info.get("publisher", "unknown"),
            volume=curriculum_info.get("volume", "unknown"),
            chapter_title=chapter_title,
            section_title=section_title,
            source_filename=os.path.basename(str(source_file_path or "")),
        )
        rel_dir = os.path.join(rel_dir_base, "formula_assets").replace("\\", "/")
        abs_dir = os.path.join(current_app.root_path, rel_dir)
        os.makedirs(abs_dir, exist_ok=True)
        placeholder_tokens = re.findall(r"\[FORMULA_IMAGE_(\d+)\]", str(question_text or ""))
        raw_block = _lookup_docx_formula_block(str(question_title or ""), docx_formula_blocks)
        if not placeholder_tokens:
            placeholder_tokens = re.findall(r"\[FORMULA_IMAGE_(\d+)\]", str(raw_block or ""))
        assets = []
        ocr_enabled = bool(current_app.config.get("ENABLE_DOCX_FORMULA_OCR_FALLBACK", False))
        ocr_model = None
        for idx, a in enumerate(formula_candidates, start=1):
            src_rel = str(a.get("path") or "")
            content_type = str(a.get("content_type") or _guess_image_content_type(src_rel))
            original_format = _docx_image_original_format(src_rel, content_type)
            ext = os.path.splitext(src_rel)[1].lower().lstrip(".") or original_format or "bin"
            if ext == "jpeg":
                ext = "jpg"
            src_abs = src_rel if os.path.isabs(src_rel) else os.path.join(current_app.root_path, src_rel)
            asset_hash = ""
            if src_abs and os.path.exists(src_abs):
                try:
                    with open(src_abs, "rb") as fsrc:
                        asset_hash = hashlib.sha1(fsrc.read()).hexdigest()
                except Exception:
                    asset_hash = ""
            if not asset_hash:
                fallback_key = f"{src_rel}|{a.get('rid')}|{a.get('block_index') or idx}"
                asset_hash = hashlib.sha1(fallback_key.encode("utf-8")).hexdigest()

            token_index = int(placeholder_tokens[idx - 1]) if idx - 1 < len(placeholder_tokens) else idx
            section_code = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "", str(section_title or "").strip())[:24] or "section"
            title_key = _safe_title_for_filename(str(question_title or ""))[:24]
            original_name = os.path.basename(src_rel) or f"asset_{idx}.{ext}"
            filename = f"{section_code}_{title_key}_{token_index}_{asset_hash[:8]}_{original_name}"
            rel_path = formula_asset_persist_cache.get(asset_hash)
            copied_abs = None
            if not rel_path:
                copied_abs = _copy_docx_asset_to_question_assets(src_rel, abs_dir, filename)
                if copied_abs:
                    rel_path = os.path.join(rel_dir, filename).replace("\\", "/")
                    formula_asset_persist_cache[asset_hash] = rel_path

            display_path = None
            converted_path = None
            conversion_status = "persisted" if rel_path else "failed"
            conversion_error = None
            if original_format in ("png", "jpeg"):
                display_path = rel_path
                if conversion_status != "failed":
                    conversion_status = "not_required"
            elif original_format in ("wmf", "emf") and rel_path:
                copied_abs = copied_abs or os.path.join(current_app.root_path, rel_path)
                png_filename = os.path.splitext(filename)[0] + ".png"
                png_abs = os.path.join(abs_dir, png_filename)
                png_rel = os.path.join(rel_dir, png_filename).replace("\\", "/")
                try:
                    ok, err = convert_vector_image_to_png(copied_abs, png_abs)
                except Exception as exc:
                    ok, err = False, str(exc)
                if ok:
                    display_path = png_rel
                    converted_path = png_rel
                    conversion_status = "success"
                else:
                    conversion_status = "failed"
                    conversion_error = err or "conversion_failed"
            elif not rel_path:
                conversion_status = "failed"
                conversion_error = "source_asset_not_found"

            token = f"[FORMULA_IMAGE_{token_index}]"
            asset_meta = {
                "source": "docx",
                "asset_type": str(a.get("asset_type") or "word_formula_image"),
                "original_path": rel_path,
                "path": rel_path,
                "display_path": display_path,
                "content_type": content_type,
                "original_format": original_format,
                "placeholder_token": token,
                "placeholder_index": token_index,
                "persist_status": "persisted" if rel_path else "failed",
                "old_temp_path": src_rel.replace("\\", "/"),
                "asset_hash": asset_hash,
                "conversion_status": conversion_status,
                "converted_path": converted_path,
                "rid": a.get("rid"),
                "image_attach_reason": a.get("image_attach_reason"),
                "is_formula_placeholder_source": bool(a.get("is_formula_placeholder_source", False)),
            }
            if a.get("mapping_status"):
                asset_meta["mapping_status"] = a.get("mapping_status")
            if conversion_error:
                asset_meta["conversion_error"] = conversion_error
            if ocr_enabled:
                ocr_rel = converted_path or display_path
                if ocr_rel and _docx_image_original_format(ocr_rel, content_type) in ("png", "jpeg"):
                    try:
                        from PIL import Image
                        if ocr_model is None:
                            ocr_model = get_model("vision_analyzer")
                        ocr_abs = ocr_rel if os.path.isabs(ocr_rel) else os.path.join(current_app.root_path, ocr_rel)
                        with Image.open(ocr_abs) as img:
                            prompt = (
                                "Transcribe only the math formula visible in this image. "
                                "Do not solve the problem, do not add explanation, and do not infer or invent missing numbers. "
                                "If the formula is unreadable, return [UNREADABLE_FORMULA]."
                            )
                            resp = ocr_model.generate_content(
                                [prompt, img],
                                generation_config={"temperature": 0.0, "max_output_tokens": 512},
                            )
                        ocr_text = str(getattr(resp, "text", "") or "").strip()
                        if ocr_text:
                            asset_meta["formula_ocr_text"] = ocr_text
                            asset_meta["formula_ocr_source"] = ocr_rel
                            asset_meta["formula_ocr_status"] = "success"
                        else:
                            asset_meta["formula_ocr_status"] = "failed"
                    except Exception as exc:
                        asset_meta["formula_ocr_status"] = "failed"
                        asset_meta["formula_ocr_error"] = str(exc)
            assets.append(asset_meta)
        placeholders = [a["placeholder_token"] for a in assets]
        meta = {
            "formula_assets": assets,
            "formula_placeholders": placeholders,
            "needs_formula_review": True,
            "formula_missing": True,
            "needs_review": True,
        }
        if any(a.get("formula_ocr_status") for a in assets):
            meta["formula_ocr_status"] = "success" if any(a.get("formula_ocr_status") == "success" for a in assets) else "failed"
            meta["formula_ocr_text"] = [a.get("formula_ocr_text") for a in assets if a.get("formula_ocr_text")]
            meta["formula_ocr_source"] = [a.get("formula_ocr_source") for a in assets if a.get("formula_ocr_source")]
        current_app.logger.warning(f"[DOCX FORMULA WARNING] formula_asset_requires_review title={question_title}")
        return meta

    def _build_math_metadata(raw_text, standardized_meta, needs_review=False):
        confidence = 0.9
        if needs_review or standardized_meta.get("needs_review", False):
            confidence = 0.4
        auto_fill_applied = confidence >= auto_fill_threshold
        review_required = confidence < auto_fill_threshold
        meta = {
            "math_format": "standard_latex",
            "raw_problem_backup": str(raw_text or ""),
            "math_warnings": list(standardized_meta.get("warnings", [])),
            "needs_review": bool(needs_review or standardized_meta.get("needs_review", False) or review_required),
            "import_confidence": confidence,
            "auto_fill_threshold": auto_fill_threshold,
            "auto_fill_applied": bool(auto_fill_applied),
            "review_required_by_confidence": bool(review_required),
            "import_source_priority": "docx_primary",
            "optional_enrich_pdf_path": optional_enrich_pdf_path or "",
            "rollback_metadata": {
                "preserved": bool(import_policy.get("preserve_rollback_metadata", True)),
                "policy": {
                    "execution_arch": import_policy.get("execution_arch", "native"),
                    "docx_primary": bool(import_policy.get("docx_primary", True)),
                    "pdf_optional_enrich": bool(import_policy.get("pdf_optional_enrich", True)),
                    "auto_post_ocr_ai": bool(import_policy.get("auto_post_ocr_ai", True)),
                    "auto_backfill_high_confidence": bool(import_policy.get("auto_backfill_high_confidence", True)),
                    "review_low_confidence": bool(import_policy.get("review_low_confidence", True)),
                },
            },
        }
        return meta

    def _load_record_metadata(record_obj):
        raw = getattr(record_obj, "notes", None)
        if isinstance(raw, str) and raw.strip():
            try:
                parsed = json.loads(raw)
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                return {}
        return {}

    def _dedupe_formula_assets(existing_assets, incoming_assets):
        merged = []
        seen = set()
        for arr in (existing_assets or [], incoming_assets or []):
            for a in arr or []:
                if not isinstance(a, dict):
                    continue
                key = (
                    str(a.get("original_path") or ""),
                    str(a.get("placeholder_token") or ""),
                    str(a.get("path") or ""),
                )
                if key in seen:
                    continue
                seen.add(key)
                merged.append(dict(a))
        return merged

    def _merge_duplicate_existing_record(
        existing_record,
        *,
        incoming_problem_text="",
        incoming_meta=None,
        incoming_correct_answer="",
        incoming_detailed_solution="",
        title="",
    ):
        """Merge richer duplicate payload into existing textbook_examples row.

        Returns (changed: bool, reason: str).
        """
        nonlocal merge_guard_kept_existing, merge_guard_updated_incoming
        changed = False
        reasons = []
        incoming_meta = dict(incoming_meta or {})
        existing_meta = _load_record_metadata(existing_record)

        existing_formula_assets = existing_meta.get("formula_assets", []) if isinstance(existing_meta, dict) else []
        incoming_formula_assets = incoming_meta.get("formula_assets", []) if isinstance(incoming_meta, dict) else []
        merged_formula_assets = _dedupe_formula_assets(existing_formula_assets, incoming_formula_assets)
        if len(merged_formula_assets) > len(existing_formula_assets):
            existing_meta["formula_assets"] = merged_formula_assets
            placeholders = [a.get("placeholder_token") for a in merged_formula_assets if a.get("placeholder_token")]
            if placeholders:
                existing_meta["formula_placeholders"] = sorted({str(p) for p in placeholders})
            changed = True
            reasons.append("formula_assets")

        for flag in ("needs_review", "needs_formula_review", "formula_missing"):
            if incoming_meta.get(flag) is True and existing_meta.get(flag) is not True:
                existing_meta[flag] = True
                changed = True
                reasons.append(flag)

        existing_problem_text = str(getattr(existing_record, "problem_text", "") or "")
        new_problem_text = str(incoming_problem_text or "")
        replace_problem_text, existing_quality, incoming_quality = should_replace_problem_text(
            existing_problem_text,
            new_problem_text,
        )
        title_for_log = str(title or getattr(existing_record, "source_description", "") or "").strip()
        if replace_problem_text and new_problem_text:
            existing_record.problem_text = new_problem_text
            changed = True
            reasons.append("problem_text_incoming_better")
            current_app.logger.info(
                f"[MERGE GUARD] updated_incoming_better_problem_text title={title_for_log} "
                f"existing_score={existing_quality['score']} incoming_score={incoming_quality['score']}"
            )
            merge_guard_updated_incoming += 1
        elif new_problem_text and existing_problem_text != new_problem_text:
            current_app.logger.info(
                f"[MERGE GUARD] kept_existing_better_problem_text title={title_for_log} "
                f"existing_score={existing_quality['score']} incoming_score={incoming_quality['score']}"
            )
            merge_guard_kept_existing += 1

        existing_answer = str(getattr(existing_record, "correct_answer", "") or "").strip()
        incoming_answer = str(incoming_correct_answer or "").strip()
        if _is_low_value_import_field(existing_answer) and not _is_low_value_import_field(incoming_answer):
            existing_record.correct_answer = incoming_answer
            changed = True
            reasons.append("correct_answer_incoming_nonblank")

        existing_solution = str(getattr(existing_record, "detailed_solution", "") or "").strip()
        incoming_solution = sanitize_detailed_solution_text(str(incoming_detailed_solution or ""), max_chars=500)
        if _is_low_value_import_field(existing_solution) and not _is_low_value_import_field(incoming_solution):
            existing_record.detailed_solution = incoming_solution
            changed = True
            reasons.append("detailed_solution_incoming_nonblank")

        if changed and existing_meta:
            attach_image_metadata(existing_record, existing_meta)
        return changed, ("+".join(sorted(set(reasons))) if reasons else "metadata_merge")

    def _extract_dedupe_from_source_description(source_description_text: str) -> str:
        m = re.search(r"dedupe=([0-9a-fA-F]+)", str(source_description_text or ""))
        return str(m.group(1)).lower() if m else ""

    def _extract_title_from_source_description(source_description_text: str) -> str:
        text = str(source_description_text or "")
        return text.split(" [", 1)[0].strip()

    def _iter_scope_existing_rows(curriculum, volume, chapter, section):
        q = TextbookExample.query.filter_by(
            source_curriculum=curriculum,
            source_volume=str(volume),
            source_chapter=chapter,
            source_section=section,
        )
        if hasattr(q, "all"):
            try:
                return q.all()
            except Exception:
                pass
        rows = getattr(q, "rows", None)
        filters = getattr(q, "filters", {})
        if isinstance(rows, list):
            out = []
            for row in rows:
                if all(getattr(row, k, None) == v for k, v in (filters or {}).items()):
                    out.append(row)
            return out
        return []

    def _find_existing_duplicate_by_dedupe(curriculum, volume, chapter, section, title, dedupe_hash):
        target_title = str(title or "").strip()
        target_hash = str(dedupe_hash or "").strip().lower()
        if not target_hash:
            for row in _iter_scope_existing_rows(curriculum, volume, chapter, section):
                sd = str(getattr(row, "source_description", "") or "")
                if _extract_title_from_source_description(sd) == target_title:
                    return row
            return None
        for row in _iter_scope_existing_rows(curriculum, volume, chapter, section):
            sd = str(getattr(row, "source_description", "") or "")
            if _extract_dedupe_from_source_description(sd) != target_hash:
                continue
            row_title = _extract_title_from_source_description(sd)
            if row_title == target_title:
                return row
        for row in _iter_scope_existing_rows(curriculum, volume, chapter, section):
            sd = str(getattr(row, "source_description", "") or "")
            if _extract_dedupe_from_source_description(sd) == target_hash:
                return row
        for row in _iter_scope_existing_rows(curriculum, volume, chapter, section):
            sd = str(getattr(row, "source_description", "") or "")
            if _extract_title_from_source_description(sd) == target_title:
                return row
        return None

    section_textbook_skill_id: str | None = None
    if is_section_textbook_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta):
        st_sid, st_src = resolve_section_textbook_skill_id(
            curriculum,
            volume,
            str(filename_meta.get("section_code") or ""),
            preferred_anchor_title,
            existing_anchor_section=existing_anchor_section,
            structure_meta=structure_meta,
        )
        if st_sid:
            section_textbook_skill_id = st_sid

    def _align_question_bank_skill(section_title: str, item: dict | None = None) -> dict[str, Any]:
        if not is_question_bank_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta):
            return {"skill_id": "", "needs_review": False, "source": "scope_not_question_bank"}
        sec_code = extract_section_code_from_title(section_title)
        inv_items = (
            (_DOCX_IMPORT_CONTEXT or {}).get("title_inventory_items")
            if isinstance(_DOCX_IMPORT_CONTEXT, dict)
            else None
        )
        anchor_title = inventory_section_title_for_code(inv_items, sec_code) or section_title
        align = resolve_question_bank_skill_for_section(
            str(curriculum_info.get("curriculum", "") or ""),
            str(curriculum_info.get("volume", "") or ""),
            anchor_title,
            section_code=sec_code,
        )
        if isinstance(item, dict) and align.get("skill_id"):
            if align.get("needs_review"):
                item["needs_review"] = True
                logs = item.get("repair_log", [])
                if not isinstance(logs, list):
                    logs = [str(logs)] if logs else []
                tag = f"question_bank_skill_align:{align.get('source')}"
                if tag not in logs:
                    logs.append(tag)
                item["repair_log"] = logs
        return align

    def _lookup_curriculum_skill_for_section_code(section_title: str, item: dict | None = None) -> str:
        return str(_align_question_bank_skill(section_title, item).get("skill_id") or "")

    def _determine_target_skill_id(base_clean_en_id, section_title, concept_name, example_obj):
        if section_textbook_skill_id and is_section_textbook_import_scope(
            _DOCX_IMPORT_CONTEXT, filename_meta
        ):
            return section_textbook_skill_id

        explicit_skill_id = str(example_obj.get("skill_id", "") or "").strip()
        if explicit_skill_id and not (
            is_question_bank_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta)
            and is_outline_skill_id(explicit_skill_id)
        ):
            return explicit_skill_id

        if is_question_bank_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta):
            align = _align_question_bank_skill(section_title, example_obj)
            mapped_skill = str(align.get("skill_id") or "")
            if mapped_skill:
                return mapped_skill

        target_clean_en_id = base_clean_en_id
        if str(target_clean_en_id or "") == "DispersionAndLinearTransformation":
            target_clean_en_id = "DispersionMeasures"
        if str(target_clean_en_id or "") == "ProbabilityOperations":
            target_clean_en_id = "ProbabilityProperties"

        if is_vocational_math:
            return normalize_vocational_math_skill_id(subject, vol_num, target_clean_en_id)
        return f"{prefix}{target_clean_en_id}"

    try:
        current_app.logger.info(" -> 開始寫入資料庫...")
        queue.put("INFO: -> 開始寫入資料庫...")
        chapters = parsed_data.get('chapters', [])
        
        for chapter_data in chapters:
            raw_chapter = chapter_data.get('chapter_title', '未知章節').strip()
            
            # === ?靽格迤 1嚗??摮蒂璅???蝭?迂 ===
            match = re.search(r'(\d+)', raw_chapter)
            if match:
                chapter_num = int(match.group(1))
            else:
                chapter_num = 999 

            # [V2.2] 引入結構大綱映射
            if structure_meta and structure_meta.get('chapter_index'):
                chapter_num = structure_meta['chapter_index']
                chapter_title = structure_meta['chapter_title']
            elif is_vocational_mathb and filename_meta.get('chapter_index'):
                chapter_num = filename_meta['chapter_index']
                chapter_title = filename_meta.get('chapter_title', raw_chapter)
            else:
                chapter_title = raw_chapter

            if is_vocational_mathb:
                # ?擃摮睬蝟餃?嚗?蝙?冽???瑽??
                if structure_meta and structure_meta.get('chapter_title'):
                    chapter_title = structure_meta['chapter_title']
                else:
                    chapter_title = raw_chapter
            elif match:
                clean_title = re.sub(r'^(\u55ae\u5143|Unit|\u7b2c)?\s*\d+\s*(\u55ae\u5143|\u7ae0)?\s*', '', raw_chapter).strip()
                chapter_title = f"第{chapter_num}章 {clean_title}" if clean_title else f"第{chapter_num}章"
            else:
                chapter_title = raw_chapter

            sections = chapter_data.get('sections', [])
            
            # 教材對齊結構 (在此處進行技能綁定)
            if curriculum_info.get('curriculum') == 'junior_high':
                chapter_title = chapter_title.replace('\n', ' ').strip()
                chapter_title = re.sub(r'^(?:Chapter|Unit|第?\s*(\d+)(?:\s*章)?)\s*', r'\1 ', chapter_title).strip()
                if chapter_title.isdigit():
                    try:
                        existing_chapter = SkillCurriculum.query.filter_by(
                            curriculum=curriculum_info['curriculum'],
                            grade=int(curriculum_info['grade']),
                            volume=curriculum_info['volume']
                        ).filter(SkillCurriculum.chapter.like(f"{chapter_title} %")).first()
                        if existing_chapter:
                            chapter_title = existing_chapter.chapter
                    except Exception:
                        pass
            
            for section_data in sections:
                section_title = section_data.get('section_title', '') or ''
                
                # [V2.2] 引入結構大綱映射
                if structure_meta and structure_meta.get('section_title'):
                    section_title = structure_meta['section_title']
                
                # [NEW] 餈質馱蝡???蝭???
                section_exists = SkillCurriculum.query.filter_by(
                    curriculum=curriculum_info.get('curriculum'),
                    volume=str(curriculum_info.get('volume', 1)),
                    chapter=chapter_title,
                    section=section_title
                ).first()
                if not section_exists:
                    sections_created += 1
                else:
                    sections_updated += 1
                
                # 蝪∪餈質馱蝡? (甇方?蝪∪???嚗?蝭?賣?閫貊雿????望擃?
                chapter_exists_any = SkillCurriculum.query.filter_by(
                    curriculum=curriculum_info.get('curriculum'),
                    volume=str(curriculum_info.get('volume', 1)),
                    chapter=chapter_title
                ).first()
                if not chapter_exists_any:
                    chapters_created += 1 # ?ㄐ?嗅祕?臬遣蝡?銝???急迨蝡??蝝??
                else:
                    chapters_updated += 1
                if is_vocational_mathb and (
                    ("3-1" in str(section_title or ""))
                    or ("3-2" in str(section_title or ""))
                    or ("3-3" in str(section_title or ""))
                ):
                    chapter_title = chapter_title.replace("第4章", "第3章")
                concepts = section_data.get('concepts', [])
                
                for concept_order, concept in enumerate(concepts, start=1):
                    concept_name = concept.get('concept_name', '未命名概念').strip()
                    concept_en_id = concept.get('concept_en_id', 'Unknown')
                    concept_paragraph = concept.get('concept_paragraph', '未提供內容').strip()
                    
                    clean_en_id = re.sub(r'[^a-zA-Z0-9]', '', concept_en_id)
                    if clean_en_id == "DispersionAndLinearTransformation":
                        clean_en_id = "DispersionMeasures"
                    order_index = concept_order
                    skip_skill_creation = is_non_skill_bucket(concept_name, clean_en_id)

                    if is_vocational_mathb and clean_en_id == "NumberOfPositiveDivisors":
                        clean_en_id = "MultiplicationPrinciple"
                        concept_name = "乘法原理"
                        concept_paragraph = "乘法原理"

                    if is_vocational_mathb:
                        mathb_1_1_order_map = {
                            "TreeDiagramCounting": 1,
                            "AdditionPrinciple": 2,
                            "MultiplicationPrinciple": 3,
                            "FactorialNotation": 4,
                        }
                        order_index = mathb_1_1_order_map.get(clean_en_id, concept_order)
                    if is_vocational_math and not is_section_textbook_import_scope(
                        _DOCX_IMPORT_CONTEXT, filename_meta
                    ):
                        final_skill_id = normalize_vocational_math_skill_id(subject, vol_num, clean_en_id)
                        skill_id_msg = f"INFO: vocational math skill_id = {final_skill_id}"
                        current_app.logger.info(skill_id_msg)
                        queue.put(skill_id_msg)
                    elif is_vocational_math:
                        final_skill_id = normalize_vocational_math_skill_id(subject, vol_num, clean_en_id)
                    else:
                        final_skill_id = f"{prefix}{clean_en_id}"
                    
                    # [V2.5] ?遣蝡?瑽芋撘?撘瑕雿輻 PlaceholderSkill 銝西歲???株???賜敦蝭
                    if outline_only:
                        final_skill_id = f"{prefix}OutlinePlaceholder"
                        # 蝣箔? Placeholder SkillInfo 摮
                        if not SkillInfo.query.get(final_skill_id):
                            db.session.add(SkillInfo(
                                skill_id=final_skill_id,
                                skill_ch_name="[Outline] Placeholder",
                                skill_en_name="OutlinePlaceholder",
                                is_active=False
                            ))
                            skills_processed += 1
                        
                        if final_skill_id not in processed_skill_ids:
                            processed_skill_ids.append(final_skill_id)
                        
                        # 撱箇? SkillCurriculum
                        existing_curr = SkillCurriculum.query.filter_by(
                            skill_id=final_skill_id,
                            chapter=chapter_title,
                            section=section_title
                        ).first()
                        if not existing_curr:
                            new_curr = SkillCurriculum(
                                skill_id=final_skill_id,
                                curriculum=curriculum_info.get('curriculum'),
                                grade=int(curriculum_info.get('grade', 10)),
                                volume=str(curriculum_info.get('volume', 1)),
                                chapter=chapter_title,
                                section=section_title,
                                paragraph=concept_paragraph,
                                display_order=(
                                    structure_meta.get('display_order_base', chapter_num * 10000)
                                    if structure_meta
                                    else (chapter_num * 10000 + (int(filename_meta.get('section_index') or 0) * 100))
                                ) + concept_order
                            )
                            db.session.add(new_curr)
                            curriculums_added += 1
                        continue # 頝喲?敺?????桀?亥???賣??

                    if not skip_skill_creation:
                        # === SkillInfo 新增/更新 ===
                        existing_skill = SkillInfo.query.get(final_skill_id)
                        if not existing_skill:
                            new_skill = SkillInfo(
                                skill_id=final_skill_id,
                                skill_en_name=clean_en_id,
                                skill_ch_name=concept_name,
                                category = section_title,
                                description=concept.get('concept_description', ''),
                                input_type='text',
                                gemini_prompt=f"Generate math problems about {concept_name}.",
                                is_active=True,
                                order_index=order_index
                            )
                            db.session.add(new_skill)
                            skills_processed += 1
                            processed_skill_ids.append(final_skill_id)
                        else:
                            existing_skill.skill_en_name = clean_en_id
                            existing_skill.skill_ch_name = concept_name
                            existing_skill.category = section_title
                            existing_skill.description = concept.get('concept_description', existing_skill.description)
                            existing_skill.order_index = order_index
                            if not existing_skill.gemini_prompt:
                                existing_skill.gemini_prompt = f"Generate math problems about {concept_name}."
                        
                        # === SkillCurriculum 新增 ===
                        existing_curr = SkillCurriculum.query.filter_by(
                            skill_id=final_skill_id,
                            chapter=chapter_title,
                            section=section_title
                        ).first()
                        
                        if not existing_curr:
                            # [V2.6] ?交?擃摮睬銝?朣仃??蝳迫撱箇??啁??暺?(?踹? 1_- ??1-1_-)
                            if is_vocational_mathb and structure_alignment_failed:
                                current_app.logger.warning(f"[ALIGNMENT] Skip creating new curriculum node for {section_title}")
                            else:
                                new_curr = SkillCurriculum(
                                    skill_id=final_skill_id,
                                    curriculum=curriculum_info.get('curriculum'),
                                    grade=int(curriculum_info.get('grade', 10)),
                                    volume=str(curriculum_info.get('volume', 1)),
                                    chapter=chapter_title,
                                    section=section_title,
                                    paragraph=concept_paragraph,
                                    display_order=(
                                        structure_meta.get('display_order_base', chapter_num * 10000)
                                        if structure_meta
                                        else (chapter_num * 10000 + (int(filename_meta.get('section_index') or 0) * 100))
                                    ) + skills_processed
                                )
                                db.session.add(new_curr)
                                curriculums_added += 1
                    else:
                        queue.put(
                            f"INFO: skip skill creation for concept='{concept_name}' ({clean_en_id}); keep examples import."
                        )

                    # === 題目寫入：依據 source_type 進行資料分類 ===
                    saved_example_skill_map = {}
                    saved_example_order = []
                    saved_example_titles = []
                    concept_known_pages = []
                    for _bucket in ("examples", "practice_questions"):
                        for _q in concept.get(_bucket, []) or []:
                            if not isinstance(_q, dict):
                                continue
                            sp = _q.get("source_page")
                            pi = _q.get("page_index")
                            try:
                                if sp is not None:
                                    concept_known_pages.append(int(sp))
                                elif pi is not None:
                                    concept_known_pages.append(int(pi) + 1)
                            except Exception:
                                continue
                    for ex_idx, ex in enumerate(concept.get('examples', []), start=1):
                        problem_text = ex.get('problem_text')
                        if not problem_text:
                            continue

                        example_title = get_question_title(ex) or "例題"
                        detected_titles.append(example_title)
                        source_type = normalize_source_type_by_title(ex, default_source_type="textbook_example")

                        # Skip section exposition entries (e.g. title='課文內容').
                        if source_type == "section_exposition":
                            current_app.logger.info(
                                f"[SKIP] section_exposition title={example_title!r} not saved as textbook_example"
                            )
                            continue

                        target_skill_id = _determine_target_skill_id(clean_en_id, section_title, concept_name, ex)
                        if (
                            (_DOCX_IMPORT_CONTEXT or {}).get("chapter_self_assessment_mode")
                            and source_type == "self_assessment"
                        ):
                            align = _align_question_bank_skill(section_title, ex)
                            sec_skill = str(align.get("skill_id") or "")
                            if sec_skill:
                                target_skill_id = sec_skill

                        sub_questions = ex.get("sub_questions", []) if isinstance(ex.get("sub_questions", []), list) else []
                        db_problem_text_raw = _render_sub_questions_problem(problem_text, sub_questions)
                        segmented_text, seg_meta = segment_question_block_text(db_problem_text_raw, question_title=example_title)
                        if seg_meta.get("changed"):
                            logs = ex.get("repair_log", [])
                            if not isinstance(logs, list):
                                logs = [str(logs)]
                            if seg_meta.get("reason"):
                                logs.append(seg_meta.get("reason"))
                            ex["repair_log"] = logs
                            db_problem_text_raw = segmented_text
                        block_kind = classify_non_question_block(db_problem_text_raw)
                        if block_kind in ("concept_explanation", "figure_caption", "narration"):
                            logs = ex.get("repair_log", [])
                            if not isinstance(logs, list):
                                logs = [str(logs)]
                            logs.append(f"detected {block_kind}, not imported as textbook_example")
                            ex["repair_log"] = logs
                            current_app.logger.info(
                                f"[DOCX BLOCK FILTER] skip example title={example_title} kind={block_kind}"
                            )
                            continue
                        converted_latex_mode = (
                            str((_DOCX_IMPORT_CONTEXT or {}).get("docx_formula_source_mode", "") or "").strip()
                            == "converted_docx_latex"
                        )
                        if not converted_latex_mode:
                            blank_norm_text, blank_meta = normalize_fill_blank_artifacts(db_problem_text_raw)
                            perm_norm_text, perm_meta = normalize_permutation_combination_notation(
                                blank_norm_text,
                                volume=str(curriculum_info.get("volume", "") or ""),
                                section_title=section_title,
                            )
                            db_problem_text_raw = perm_norm_text
                            if blank_meta.get("changed") or perm_meta.get("changed"):
                                logs = ex.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.extend(blank_meta.get("reasons", []))
                                logs.extend(perm_meta.get("reasons", []))
                                ex["repair_log"] = logs
                        raw_formula_block = _lookup_docx_formula_block(str(example_title), docx_formula_blocks)
                        if raw_formula_block and re.search(r"\[FORMULA_IMAGE_\d+\]|\[WORD_EQUATION_UNPARSED\]", raw_formula_block):
                            ex["needs_review"] = True
                            ex["needs_formula_review"] = True
                            ex["formula_missing"] = True
                        db_problem_text_raw = validate_problem_formula_not_hallucinated(
                            example_title, ex, db_problem_text_raw, raw_formula_block
                        )
                        if not converted_latex_mode:
                            repaired_text, repair_meta = repair_missing_single_variable_text(db_problem_text_raw)
                            if repair_meta.get("applied"):
                                db_problem_text_raw = repaired_text
                                logs = ex.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.append(repair_meta.get("reason"))
                                ex["repair_log"] = logs
                            elif repair_meta.get("reason") == "non_unique_candidate_variable":
                                ex["needs_review"] = True
                        if converted_latex_mode:
                            latex_fix = normalize_converted_docx_latex_text(db_problem_text_raw)
                            db_problem_text = str(latex_fix.get("text", db_problem_text_raw) or db_problem_text_raw)
                            db_problem_text_norm = str(db_problem_text or "").strip()
                            ex_math_meta = {}
                            current_app.logger.info(
                                "[FORMULA NORMALIZE SKIP] converted_docx_latex_preserve_latex=true field=problem_text"
                            )
                            if latex_fix.get("changes"):
                                current_app.logger.info(
                                    f"[LATEX INLINE NORMALIZE] title={example_title} changes={len(latex_fix.get('changes', []))}"
                                )
                        else:
                            db_problem_text_norm = str(normalize_math_text(db_problem_text_raw) or "").strip()
                            db_problem_text_before_standardize = db_problem_text_norm
                            db_problem_text, ex_math_meta = standardize_problem_latex(db_problem_text_norm)
                            if (
                                db_problem_text != db_problem_text_before_standardize
                                and LATEX_DEBUG_SIGNAL_RE.search(str(db_problem_text_raw or ""))
                            ):
                                current_app.logger.info(
                                    f"[LATEX STANDARDIZE] title={example_title} "
                                    f"before={db_problem_text_before_standardize} after={db_problem_text}"
                                )
                        if not converted_latex_mode:
                            db_problem_text_post, post_perm_meta = normalize_permutation_combination_notation(
                                db_problem_text,
                                volume=str(curriculum_info.get("volume", "") or ""),
                                section_title=section_title,
                            )
                            if post_perm_meta.get("changed"):
                                logs = ex.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.extend(post_perm_meta.get("reasons", []))
                                ex["repair_log"] = logs
                                current_app.logger.info(
                                    f"[PERM COMB POST NORMALIZE] title={example_title} before={db_problem_text} after={db_problem_text_post}"
                                )
                            db_problem_text, post_blank_meta = normalize_fill_blank_artifacts(db_problem_text_post)
                            if post_blank_meta.get("changed"):
                                logs = ex.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.extend(post_blank_meta.get("reasons", []))
                                ex["repair_log"] = logs
                        db_problem_text, prob_meta = normalize_probability_event_notation(db_problem_text)
                        if prob_meta.get("changed"):
                            logs = ex.get("repair_log", [])
                            if not isinstance(logs, list):
                                logs = [str(logs)]
                            logs.extend(prob_meta.get("reasons", []))
                            ex["repair_log"] = logs
                        db_answer = _render_sub_questions_answer(ex.get('correct_answer', ''), sub_questions)
                        db_solution = _render_sub_questions_solution(ex.get('detailed_solution', ''), sub_questions)
                        if is_section_textbook_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta):
                            needs_review = section_textbook_allowed_needs_review(ex)
                        else:
                            needs_review = bool(ex.get("needs_review", False)) or structure_alignment_failed
                        ex["problem_text"] = db_problem_text
                        ex = validate_problem_block_purity(ex)
                        if is_section_textbook_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta):
                            needs_review = section_textbook_allowed_needs_review(ex)
                            ex["needs_review"] = needs_review
                        else:
                            needs_review = bool(ex.get("needs_review", False)) or structure_alignment_failed
                        linked_example_title = None
                        if source_type == "in_class_practice":
                            linked_example_title, needs_review = _infer_linked_example_title(
                                example_title,
                                ex.get("linked_example_title"),
                                saved_example_titles,
                                needs_review,
                            )
                            ex["linked_example_title"] = linked_example_title
                        dedupe_hash = _normalize_problem_hash(
                            db_problem_text, sub_questions=sub_questions, source_type=source_type, title=example_title
                        )
                        source_description = _build_source_description(
                            example_title,
                            source_type=source_type,
                            linked_example_title=linked_example_title,
                            needs_review=needs_review,
                            dedupe_hash=dedupe_hash
                        )

                        existing_ex = TextbookExample.query.filter_by(
                            skill_id=target_skill_id,
                            source_curriculum=curriculum_info.get('curriculum'),
                            source_volume=str(curriculum_info.get('volume')),
                            source_chapter=chapter_title,
                            source_section=section_title,
                            source_description=source_description
                        ).first()
                        if not existing_ex:
                            existing_ex = TextbookExample.query.filter_by(
                                source_curriculum=curriculum_info.get('curriculum'),
                                source_volume=str(curriculum_info.get('volume')),
                                source_chapter=chapter_title,
                                source_section=section_title,
                                source_description=source_description
                            ).first()
                        if not existing_ex:
                            existing_ex = _find_existing_duplicate_by_dedupe(
                                curriculum_info.get('curriculum'),
                                curriculum_info.get('volume'),
                                chapter_title,
                                section_title,
                                example_title,
                                dedupe_hash,
                            )

                        if source_type == "textbook_example":
                            title_num = _extract_title_number(example_title)
                            if title_num is not None:
                                saved_example_skill_map[title_num] = target_skill_id
                            saved_example_order.append((example_title, target_skill_id))
                            saved_example_titles.append(example_title)

                        if existing_ex:
                            duplicate_meta = {}
                            duplicate_formula_meta = None
                            if is_docx_source:
                                duplicate_formula_meta = _build_docx_formula_assets_metadata(
                                    example_title, question_text=db_problem_text
                                )
                            if duplicate_formula_meta:
                                duplicate_meta.update(duplicate_formula_meta)
                            for k in ("needs_review", "needs_formula_review", "formula_missing"):
                                if ex.get(k) is True:
                                    duplicate_meta[k] = True

                            skill_migrated = False
                            if (
                                (_DOCX_IMPORT_CONTEXT or {}).get("chapter_self_assessment_mode")
                                and source_type == "self_assessment"
                                and target_skill_id
                                and is_official_vh_skill_id(target_skill_id)
                                and str(getattr(existing_ex, "skill_id", "") or "") != target_skill_id
                            ):
                                old_skill = str(getattr(existing_ex, "skill_id", "") or "")
                                existing_ex.skill_id = target_skill_id
                                skill_migrated = True
                                current_app.logger.info(
                                    "[QUESTION BANK SKILL ALIGN] section_code=%s title=%s "
                                    "migrated_skill_id %s -> %s source=duplicate_realign",
                                    extract_section_code_from_title(section_title),
                                    example_title,
                                    old_skill,
                                    target_skill_id,
                                )

                            changed, _reason = _merge_duplicate_existing_record(
                                existing_ex,
                                incoming_problem_text=db_problem_text,
                                incoming_meta=duplicate_meta,
                                incoming_correct_answer=db_answer,
                                incoming_detailed_solution=db_solution,
                                title=example_title,
                            )
                            if changed or skill_migrated:
                                updated_duplicates += 1
                                current_app.logger.info(
                                    f"[PRACTICE IMPORT] updated duplicate title={example_title} "
                                    f"reason={'skill_realign' if skill_migrated and not changed else 'metadata_merge'}"
                                )
                            else:
                                duplicates_skipped_count += 1
                                current_app.logger.info(
                                    f"[PRACTICE IMPORT] skipped duplicate title={example_title} reason=dedupe_match"
                                )
                            continue

                        try:
                            difficulty_level = int(ex.get('difficulty_level', 1))
                        except Exception:
                            difficulty_level = 1

                        new_ex = TextbookExample(
                            skill_id=target_skill_id,
                            source_curriculum=curriculum_info.get('curriculum'),
                            source_volume=str(curriculum_info.get('volume')),
                            source_chapter=chapter_title,
                            source_section=section_title,
                            source_paragraph=concept_name,
                            source_description=source_description,
                            problem_text=db_problem_text,
                            problem_type=(structure_meta.get('type') if structure_meta and structure_meta.get('type') else ex.get('problem_type', source_type or 'calculation')),
                            correct_answer=db_answer,
                            detailed_solution=sanitize_detailed_solution_text(db_solution, max_chars=500),
                            difficulty_level=difficulty_level
                        )
                        chapter_rel_dir, _, chapter_id, section_id = build_question_assets_dir(
                            curriculum_info, chapter_title, section_title
                        )
                        _ = chapter_rel_dir  # keep naming consistent, not used here
                        example_code = build_question_code(chapter_id, section_id, "example", ex_idx)
                        image_meta = None
                        if is_pdf_source:
                            image_meta = _build_image_metadata(
                                question_title=example_title,
                                question_text=db_problem_text,
                                chapter_title=chapter_title,
                                section_title=section_title,
                                source_type=source_type,
                                question_code=example_code,
                                force_has_image=bool(ex.get("has_image", False)),
                                image_description=str(ex.get("image_description", "") or ""),
                                source_page=ex.get("source_page"),
                                page_index=ex.get("page_index"),
                                item_payload={**ex, "_neighbor_source_pages": concept_known_pages},
                                concept_name=concept_name,
                            )
                        if is_docx_source:
                            docx_meta = _build_docx_assets_metadata(
                                example_title,
                                chapter_title,
                                section_title,
                                source_type,
                                question_text=db_problem_text,
                                image_hint=str(ex.get("image_description", "") or ""),
                            )
                            if docx_meta:
                                image_meta = dict(image_meta or {})
                                image_meta.update(docx_meta)
                            if not converted_latex_mode:
                                formula_meta = _build_docx_formula_assets_metadata(example_title, question_text=db_problem_text)
                                if formula_meta:
                                    image_meta = dict(image_meta or {})
                                    image_meta.update(formula_meta)
                        if image_meta:
                            attached = attach_image_metadata(new_ex, image_meta)
                            if attached:
                                current_app.logger.info(f"{'[DOCX IMAGE]' if is_docx_source else '[QUESTION IMAGE]'} metadata attached question={example_title}")
                                if is_docx_source:
                                    img_assets = image_meta.get("image_assets", []) if isinstance(image_meta, dict) else []
                                    docx_attached_count += len(img_assets)
                                    docx_copied_to_question_assets += len(img_assets)
                                    for ia in img_assets:
                                        if ia.get("display_path"):
                                            docx_direct_display_images += 1
                                        if ia.get("original_format") in ("wmf", "emf"):
                                            docx_vector_images += 1
                                        if ia.get("converted_path"):
                                            docx_conversion_success += 1
                                        if ia.get("needs_image_conversion") is True and not ia.get("display_path"):
                                            docx_conversion_failed += 1
                                    formula_assets = image_meta.get("formula_assets", []) if isinstance(image_meta, dict) else []
                                    if formula_assets:
                                        docx_formula_assets_count += len(formula_assets)
                                        if image_meta.get("needs_formula_review"):
                                            docx_formula_needs_review_count += 1
                                    for fa in formula_assets:
                                        if fa.get("conversion_status") == "success":
                                            docx_conversion_success += 1
                                        elif fa.get("conversion_status") == "failed":
                                            docx_conversion_failed += 1
                                        elif fa.get("conversion_status") == "pending":
                                            docx_conversion_pending += 1
                            else:
                                current_app.logger.info(
                                    "[QUESTION IMAGE] detected but no metadata field available table=textbook_examples"
                                )
                        if is_docx_source and isinstance(image_meta, dict) and image_meta.get("has_image") and not image_meta.get("image_assets"):
                            reason = image_meta.get("image_warning", "unknown")
                            missing_image_questions.append((example_title, source_type, reason))
                            current_app.logger.info(
                                f"[DOCX IMAGE DEBUG] missing_image_candidate title={example_title} source_type={source_type} reason={reason}"
                            )
                        math_meta = _build_math_metadata(db_problem_text_raw, ex_math_meta, needs_review=needs_review)
                        for k in (
                            "needs_formula_review",
                            "formula_missing",
                            "formula_hallucination_risk",
                            "parse_warning",
                            "problem_unusable",
                            "block_boundary_error",
                            "likely_concept_explanation",
                            "skill_boundary_mismatch",
                            "has_answer_blank",
                            "question_format",
                            "needs_image_review",
                            "missing_docx_image_asset",
                            "needs_table_review",
                            "repair_log",
                        ):
                            if ex.get(k) is not None:
                                math_meta[k] = ex.get(k)
                        attach_image_metadata(new_ex, math_meta)
                        if PERM_COMB_NOTATION_RE.search(str(db_problem_text or "")):
                            current_app.logger.info(f"[DB WRITE CHECK] title={example_title} problem_text={db_problem_text}")
                        db.session.add(new_ex)
                        if source_type == "textbook_example":
                            examples_added += 1
                        else:
                            practice_questions_imported += 1
                            summary_bucket = classify_practice_source_bucket(source_type)
                            if summary_bucket == "in_class_practice":
                                in_class_practices_imported += 1
                            elif summary_bucket == "chapter_exercise":
                                chapter_exercises_imported += 1
                            elif summary_bucket == "self_assessment":
                                self_assessments_imported += 1
                            elif summary_bucket == "exam_practice":
                                exam_practices_imported += 1
                                current_app.logger.info(
                                    f"[EXAM PRACTICE IMPORT] detected title={example_title} source_type={source_type} skill_id={target_skill_id}"
                                )
                            else:
                                other_practices_imported += 1
                            if needs_review:
                                practice_questions_needs_review += 1

                    # === 範例題與練習題寫入 ===
                    for practice_idx, practice in enumerate(concept.get('practice_questions', []) or [], start=1):
                        if not isinstance(practice, dict):
                            continue

                        practice_title = get_question_title(practice) or "隨堂練習"
                        source_type = normalize_source_type_by_title(practice, default_source_type="in_class_practice")
                        practice_problem = str(
                            practice.get("problem_text", "") or practice.get("problem", "") or practice.get("question", "")
                        ).strip()
                        if not practice_problem:
                            if source_type == "self_assessment":
                                practice_problem = f"{practice_title} 憿凳蝻箏仃 [FORMULA_MISSING]"
                                practice["needs_review"] = True
                                practice["needs_formula_review"] = True
                                practice["formula_missing"] = True
                                logs = practice.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.append("preserved self_assessment item with missing stem")
                                practice["repair_log"] = logs
                            else:
                                continue

                        detected_titles.append(practice_title)
                        sub_questions = practice.get("sub_questions", []) if isinstance(practice.get("sub_questions", []), list) else []
                        practice_problem_raw = _render_sub_questions_problem(practice_problem, sub_questions)
                        segmented_text, seg_meta = segment_question_block_text(practice_problem_raw, question_title=practice_title)
                        if seg_meta.get("changed"):
                            logs = practice.get("repair_log", [])
                            if not isinstance(logs, list):
                                logs = [str(logs)]
                            if seg_meta.get("reason"):
                                logs.append(seg_meta.get("reason"))
                            practice["repair_log"] = logs
                            practice_problem_raw = segmented_text
                        block_kind = classify_non_question_block(practice_problem_raw)
                        if block_kind in ("concept_explanation", "figure_caption", "narration"):
                            logs = practice.get("repair_log", [])
                            if not isinstance(logs, list):
                                logs = [str(logs)]
                            logs.append(f"detected {block_kind}, skipped from question text")
                            practice["repair_log"] = logs
                            current_app.logger.info(
                                f"[DOCX BLOCK FILTER] skip practice title={practice_title} kind={block_kind}"
                            )
                            continue
                        converted_latex_mode = (
                            str((_DOCX_IMPORT_CONTEXT or {}).get("docx_formula_source_mode", "") or "").strip()
                            == "converted_docx_latex"
                        )
                        if not converted_latex_mode:
                            blank_norm_text, blank_meta = normalize_fill_blank_artifacts(practice_problem_raw)
                            perm_norm_text, perm_meta = normalize_permutation_combination_notation(
                                blank_norm_text,
                                volume=str(curriculum_info.get("volume", "") or ""),
                                section_title=section_title,
                            )
                            practice_problem_raw = perm_norm_text
                            if blank_meta.get("changed") or perm_meta.get("changed"):
                                logs = practice.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.extend(blank_meta.get("reasons", []))
                                logs.extend(perm_meta.get("reasons", []))
                                practice["repair_log"] = logs
                        raw_formula_block = _lookup_docx_formula_block(str(practice_title), docx_formula_blocks)
                        if raw_formula_block and re.search(r"\[FORMULA_IMAGE_\d+\]|\[WORD_EQUATION_UNPARSED\]", raw_formula_block):
                            practice["needs_review"] = True
                            practice["needs_formula_review"] = True
                            practice["formula_missing"] = True
                        practice_problem_raw = validate_problem_formula_not_hallucinated(
                            practice_title, practice, practice_problem_raw, raw_formula_block
                        )
                        if not converted_latex_mode:
                            repaired_text, repair_meta = repair_missing_single_variable_text(practice_problem_raw)
                            if repair_meta.get("applied"):
                                practice_problem_raw = repaired_text
                                logs = practice.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.append(repair_meta.get("reason"))
                                practice["repair_log"] = logs
                            elif repair_meta.get("reason") == "non_unique_candidate_variable":
                                practice["needs_review"] = True
                        if converted_latex_mode:
                            latex_fix = normalize_converted_docx_latex_text(practice_problem_raw)
                            practice_problem = str(latex_fix.get("text", practice_problem_raw) or practice_problem_raw)
                            practice_problem_norm = str(practice_problem or "").strip()
                            practice_math_meta = {}
                            current_app.logger.info(
                                "[FORMULA NORMALIZE SKIP] converted_docx_latex_preserve_latex=true field=problem_text"
                            )
                            if latex_fix.get("changes"):
                                current_app.logger.info(
                                    f"[LATEX INLINE NORMALIZE] title={practice_title} changes={len(latex_fix.get('changes', []))}"
                                )
                        else:
                            practice_problem_norm = str(normalize_math_text(practice_problem_raw) or "").strip()
                            practice_problem_before_standardize = practice_problem_norm
                            practice_problem, practice_math_meta = standardize_problem_latex(practice_problem_norm)
                            if (
                                practice_problem != practice_problem_before_standardize
                                and LATEX_DEBUG_SIGNAL_RE.search(str(practice_problem_raw or ""))
                            ):
                                current_app.logger.info(
                                    f"[LATEX STANDARDIZE] title={practice_title} "
                                    f"before={practice_problem_before_standardize} after={practice_problem}"
                                )
                        if not converted_latex_mode:
                            practice_problem_post, post_perm_meta = normalize_permutation_combination_notation(
                                practice_problem,
                                volume=str(curriculum_info.get("volume", "") or ""),
                                section_title=section_title,
                            )
                            if post_perm_meta.get("changed"):
                                logs = practice.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.extend(post_perm_meta.get("reasons", []))
                                practice["repair_log"] = logs
                                current_app.logger.info(
                                    f"[PERM COMB POST NORMALIZE] title={practice_title} before={practice_problem} after={practice_problem_post}"
                                )
                            practice_problem, post_blank_meta = normalize_fill_blank_artifacts(practice_problem_post)
                            if post_blank_meta.get("changed"):
                                logs = practice.get("repair_log", [])
                                if not isinstance(logs, list):
                                    logs = [str(logs)]
                                logs.extend(post_blank_meta.get("reasons", []))
                                practice["repair_log"] = logs
                        practice_problem, prob_meta = normalize_probability_event_notation(practice_problem)
                        if prob_meta.get("changed"):
                            logs = practice.get("repair_log", [])
                            if not isinstance(logs, list):
                                logs = [str(logs)]
                            logs.extend(prob_meta.get("reasons", []))
                            practice["repair_log"] = logs
                        practice_answer = _render_sub_questions_answer(practice.get('correct_answer', ''), sub_questions)
                        practice_solution = _render_sub_questions_solution(practice.get('detailed_solution', ''), sub_questions)
                        linked_example_title = str(practice.get("linked_example_title", "") or "").strip() or None
                        if is_section_textbook_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta):
                            needs_review = section_textbook_allowed_needs_review(practice)
                        else:
                            needs_review = bool(practice.get("needs_review", False))
                        practice["problem_text"] = practice_problem
                        practice = validate_problem_block_purity(practice)
                        if is_section_textbook_import_scope(_DOCX_IMPORT_CONTEXT, filename_meta):
                            needs_review = section_textbook_allowed_needs_review(practice)
                        else:
                            needs_review = bool(practice.get("needs_review", False))
                        if source_type == "in_class_practice":
                            linked_example_title, needs_review = _infer_linked_example_title(
                                practice_title, linked_example_title, saved_example_titles, needs_review
                            )

                        target_skill_id = str(practice.get("skill_id", "") or "").strip()
                        if not target_skill_id:
                            linked_num = _extract_title_number(linked_example_title) if linked_example_title else None
                            if linked_num is not None and linked_num in saved_example_skill_map:
                                target_skill_id = saved_example_skill_map[linked_num]
                            elif len({sid for _, sid in saved_example_order}) == 1 and saved_example_order:
                                target_skill_id = saved_example_order[0][1]
                            elif saved_example_order:
                                target_skill_id = saved_example_order[-1][1]
                                needs_review = True
                                warn_msg = (
                                    f"[PRACTICE IMPORT WARNING] title={practice_title} reason=missing_exact_linked_example"
                                )
                                current_app.logger.warning(warn_msg)
                                queue.put(f"WARN: {warn_msg}")
                            else:
                                target_skill_id = _determine_target_skill_id(clean_en_id, section_title, concept_name, practice)
                                needs_review = True
                                warn_msg = (
                                    f"[PRACTICE IMPORT WARNING] title={practice_title} reason=missing_linked_example"
                                )
                                current_app.logger.warning(warn_msg)
                                queue.put(f"WARN: {warn_msg}")

                        log_msg = (
                            f"[PRACTICE IMPORT] detected title={practice_title} source_type={source_type} "
                            f"linked_example={linked_example_title} skill_id={target_skill_id}"
                        )
                        current_app.logger.info(log_msg)
                        queue.put(f"INFO: {log_msg}")
                        if source_type == "exam_practice":
                            current_app.logger.info(
                                f"[EXAM PRACTICE IMPORT] detected title={practice_title} source_type={source_type} skill_id={target_skill_id}"
                            )

                        dedupe_hash = _normalize_problem_hash(
                            practice_problem, sub_questions=sub_questions, source_type=source_type, title=practice_title
                        )
                        source_description = _build_source_description(
                            practice_title,
                            source_type=source_type or "in_class_practice",
                            linked_example_title=linked_example_title,
                            needs_review=needs_review,
                            dedupe_hash=dedupe_hash,
                            section_context=None,
                        )

                        existing_practice = TextbookExample.query.filter_by(
                            skill_id=target_skill_id,
                            source_curriculum=curriculum_info.get('curriculum'),
                            source_volume=str(curriculum_info.get('volume')),
                            source_chapter=chapter_title,
                            source_section=section_title,
                            source_description=source_description
                        ).first()
                        if not existing_practice:
                            existing_practice = TextbookExample.query.filter_by(
                                source_curriculum=curriculum_info.get('curriculum'),
                                source_volume=str(curriculum_info.get('volume')),
                                source_chapter=chapter_title,
                                source_section=section_title,
                                source_description=source_description
                            ).first()
                        if not existing_practice:
                            existing_practice = _find_existing_duplicate_by_dedupe(
                                curriculum_info.get('curriculum'),
                                curriculum_info.get('volume'),
                                chapter_title,
                                section_title,
                                practice_title,
                                dedupe_hash,
                            )
                        if existing_practice:
                            duplicate_meta = {}
                            duplicate_formula_meta = None
                            if is_docx_source:
                                duplicate_formula_meta = _build_docx_formula_assets_metadata(
                                    practice_title, question_text=practice_problem
                                )
                            if duplicate_formula_meta:
                                duplicate_meta.update(duplicate_formula_meta)
                            for k in ("needs_review", "needs_formula_review", "formula_missing"):
                                if practice.get(k) is True:
                                    duplicate_meta[k] = True

                            changed, _reason = _merge_duplicate_existing_record(
                                existing_practice,
                                incoming_problem_text=practice_problem,
                                incoming_meta=duplicate_meta,
                                incoming_correct_answer=practice_answer,
                                incoming_detailed_solution=practice_solution,
                                title=practice_title,
                            )
                            if changed:
                                updated_duplicates += 1
                                update_msg = (
                                    f"[PRACTICE IMPORT] updated duplicate title={practice_title} reason=metadata_merge"
                                )
                                current_app.logger.info(update_msg)
                                queue.put(f"INFO: {update_msg}")
                            else:
                                practice_questions_skipped += 1
                                duplicates_skipped_count += 1
                                skip_msg = (
                                    f"[PRACTICE IMPORT] skipped duplicate title={practice_title} reason=dedupe_match"
                                )
                                current_app.logger.info(skip_msg)
                                queue.put(f"INFO: {skip_msg}")
                            continue

                        try:
                            difficulty_level = int(practice.get('difficulty_level', 1))
                        except Exception:
                            difficulty_level = 1

                        practice_row = TextbookExample(
                            skill_id=target_skill_id,
                            source_curriculum=curriculum_info.get('curriculum'),
                            source_volume=str(curriculum_info.get('volume')),
                            source_chapter=chapter_title,
                            source_section=section_title,
                            source_paragraph=concept_name,
                            source_description=source_description,
                            problem_text=practice_problem,
                            problem_type=practice.get('problem_type', 'in_class_practice'),
                            correct_answer=practice_answer,
                            detailed_solution=sanitize_detailed_solution_text(practice_solution, max_chars=500),
                            difficulty_level=difficulty_level
                        )
                        chapter_rel_dir, _, chapter_id, section_id = build_question_assets_dir(
                            curriculum_info, chapter_title, section_title
                        )
                        _ = chapter_rel_dir
                        practice_code = build_question_code(chapter_id, section_id, "practice", practice_idx)
                        image_meta = None
                        if is_pdf_source:
                            image_meta = _build_image_metadata(
                                question_title=practice_title,
                                question_text=practice_problem,
                                chapter_title=chapter_title,
                                section_title=section_title,
                                source_type=source_type,
                                question_code=practice_code,
                                force_has_image=bool(practice.get("has_image", False)),
                                image_description=str(practice.get("image_description", "") or ""),
                                source_page=practice.get("source_page"),
                                page_index=practice.get("page_index"),
                                item_payload={**practice, "_neighbor_source_pages": concept_known_pages},
                                concept_name=concept_name,
                            )
                        if is_docx_source:
                            docx_meta = _build_docx_assets_metadata(
                                practice_title,
                                chapter_title,
                                section_title,
                                source_type,
                                question_text=practice_problem,
                                image_hint=str(practice.get("image_description", "") or ""),
                            )
                            if docx_meta:
                                image_meta = dict(image_meta or {})
                                image_meta.update(docx_meta)
                            if not converted_latex_mode:
                                formula_meta = _build_docx_formula_assets_metadata(practice_title, question_text=practice_problem)
                                if formula_meta:
                                    image_meta = dict(image_meta or {})
                                    image_meta.update(formula_meta)
                        if image_meta:
                            attached = attach_image_metadata(practice_row, image_meta)
                            if attached:
                                current_app.logger.info(f"{'[DOCX IMAGE]' if is_docx_source else '[QUESTION IMAGE]'} metadata attached question={practice_title}")
                                if is_docx_source:
                                    img_assets = image_meta.get("image_assets", []) if isinstance(image_meta, dict) else []
                                    docx_attached_count += len(img_assets)
                                    docx_copied_to_question_assets += len(img_assets)
                                    for ia in img_assets:
                                        if ia.get("display_path"):
                                            docx_direct_display_images += 1
                                        if ia.get("original_format") in ("wmf", "emf"):
                                            docx_vector_images += 1
                                        if ia.get("converted_path"):
                                            docx_conversion_success += 1
                                        if ia.get("needs_image_conversion") is True and not ia.get("display_path"):
                                            docx_conversion_failed += 1
                                    formula_assets = image_meta.get("formula_assets", []) if isinstance(image_meta, dict) else []
                                    if formula_assets:
                                        docx_formula_assets_count += len(formula_assets)
                                        if image_meta.get("needs_formula_review"):
                                            docx_formula_needs_review_count += 1
                                    for fa in formula_assets:
                                        if fa.get("conversion_status") == "success":
                                            docx_conversion_success += 1
                                        elif fa.get("conversion_status") == "failed":
                                            docx_conversion_failed += 1
                                        elif fa.get("conversion_status") == "pending":
                                            docx_conversion_pending += 1
                            else:
                                current_app.logger.info(
                                    "[QUESTION IMAGE] detected but no metadata field available table=textbook_examples"
                                )
                        if is_docx_source and isinstance(image_meta, dict) and image_meta.get("has_image") and not image_meta.get("image_assets"):
                            reason = image_meta.get("image_warning", "unknown")
                            missing_image_questions.append((practice_title, source_type, reason))
                            current_app.logger.info(
                                f"[DOCX IMAGE DEBUG] missing_image_candidate title={practice_title} source_type={source_type} reason={reason}"
                            )
                        math_meta = _build_math_metadata(practice_problem_raw, practice_math_meta, needs_review=needs_review)
                        for k in (
                            "needs_formula_review",
                            "formula_missing",
                            "formula_hallucination_risk",
                            "parse_warning",
                            "problem_unusable",
                            "block_boundary_error",
                            "likely_concept_explanation",
                            "skill_boundary_mismatch",
                            "has_answer_blank",
                            "question_format",
                            "needs_image_review",
                            "missing_docx_image_asset",
                            "needs_table_review",
                            "repair_log",
                        ):
                            if practice.get(k) is not None:
                                math_meta[k] = practice.get(k)
                        attach_image_metadata(practice_row, math_meta)
                        if PERM_COMB_NOTATION_RE.search(str(practice_problem or "")):
                            current_app.logger.info(f"[DB WRITE CHECK] title={practice_title} problem_text={practice_problem}")
                        db.session.add(practice_row)

                        practice_questions_imported += 1
                        summary_bucket = classify_practice_source_bucket(source_type)
                        if summary_bucket == "in_class_practice":
                            in_class_practices_imported += 1
                            n = _extract_title_number(practice_title)
                            if n is not None:
                                in_class_nums.append(n)
                        elif summary_bucket == "chapter_exercise":
                            chapter_exercises_imported += 1
                        elif summary_bucket == "self_assessment":
                            self_assessments_imported += 1
                        elif summary_bucket == "exam_practice":
                            exam_practices_imported += 1
                        else:
                            other_practices_imported += 1
                        if needs_review:
                            practice_questions_needs_review += 1
                        saved_msg = (
                            f"[PRACTICE IMPORT] saved independent question title={practice_title} "
                            f"table=textbook_examples id=pending_commit"
                        )
                        current_app.logger.info(saved_msg)
                        queue.put(f"INFO: {saved_msg}")

        db.session.commit()
        if is_docx_source:
            ctx = _DOCX_IMPORT_CONTEXT or {}
            current_app.logger.info(f"[DOCX MODE] docx_formula_source_mode={ctx.get('docx_formula_source_mode', 'auto_detect')}")
            current_app.logger.info(f"[DOCX MODE] is_converted_latex_docx={bool(ctx.get('is_converted_latex_docx', False))}")
            current_app.logger.info(f"[DOCX MODE] latex_signal_count={int(ctx.get('latex_signal_count', 0) or 0)}")
            current_app.logger.info(f"[DOCX MODE] formula_placeholder_count={int(ctx.get('formula_placeholder_count', 0) or 0)}")
            current_app.logger.info(f"[DOCX MODE] formula_assets_extraction_skipped={bool(ctx.get('formula_assets_extraction_skipped', False))}")
            current_app.logger.info(f"[DOCX MODE] ocr_skipped={bool(ctx.get('ocr_skipped', False))}")
            current_app.logger.info(f"[DOCX MODE] pix2tex_skipped={bool(ctx.get('pix2tex_skipped', False))}")
            current_app.logger.info(f"[DOCX MODE] records_with_latex={records_with_latex}")
            current_app.logger.info(f"[DOCX MODE] records_with_placeholder={records_with_placeholder}")
            current_app.logger.info(f"[DOCX MODE] merge_guard_kept_existing={merge_guard_kept_existing}")
            current_app.logger.info(f"[DOCX MODE] merge_guard_updated_incoming={merge_guard_updated_incoming}")
            media_total = len((ctx.get("media_rel_map") or {})) if isinstance(ctx, dict) else 0
            orphan_total = len((ctx.get("orphan_images") or [])) if isinstance(ctx, dict) else 0
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] media_total={media_total}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] attached_images={docx_attached_count}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] orphan_images={orphan_total}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] copied_to_question_assets={docx_copied_to_question_assets}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] direct_display_images={docx_direct_display_images}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] vector_images={docx_vector_images}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] conversion_success={docx_conversion_success}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] conversion_failed={docx_conversion_failed}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] conversion_pending={docx_conversion_pending}")
            current_app.logger.info(f"[FORMULA ASSET SUMMARY] formula_assets={docx_formula_assets_count}")
            current_app.logger.info(f"[FORMULA ASSET SUMMARY] needs_formula_review={docx_formula_needs_review_count}")
            current_app.logger.info(f"[DOCX IMAGE SUMMARY] missing_image_questions={len(missing_image_questions)}")
            for t, s_type, reason in missing_image_questions:
                current_app.logger.warning(
                    f"[DOCX IMAGE SUMMARY WARNING] missing_image title={t} source_type={s_type} reason={reason}"
                )
            current_app.logger.info(f"[DOCX IMPORT VALIDATION] detected_titles={len(detected_titles)}")
            current_app.logger.info(f"[DOCX IMPORT VALIDATION] examples={examples_added}")
            current_app.logger.info(f"[DOCX IMPORT VALIDATION] in_class_practices={in_class_practices_imported}")
            current_app.logger.info(f"[DOCX IMPORT VALIDATION] exercises={chapter_exercises_imported}")
            if in_class_nums:
                uniq = sorted(set(in_class_nums))
                miss = [x for x in range(uniq[0], uniq[-1] + 1) if x not in uniq]
                if miss:
                    current_app.logger.warning(
                        f"[DOCX IMPORT VALIDATION WARNING] possible missing in_class_practice numbers={miss}"
                    )
        return {
            'skills_processed': skills_processed,
            'curriculums_added': curriculums_added,
            'chapters_created': chapters_created,
            'chapters_updated': chapters_updated,
            'sections_created': sections_created,
            'sections_updated': sections_updated,
            'examples_added': examples_added,
            'practice_questions_imported': practice_questions_imported,
            'in_class_practices_imported': in_class_practices_imported,
            'chapter_exercises_imported': chapter_exercises_imported,
            'self_assessments_imported': self_assessments_imported,
            'exam_practices_imported': exam_practices_imported,
            'other_practices_imported': other_practices_imported,
            'needs_review_count': practice_questions_needs_review,
            'intra_import_duplicates_merged': intra_import_duplicates_merged,
            'practice_questions_skipped': practice_questions_skipped,
            'duplicates_skipped': duplicates_skipped_count,
            'updated_duplicates': updated_duplicates,
            'processed_skill_ids': processed_skill_ids,
            'records_with_latex': records_with_latex,
            'records_with_placeholder': records_with_placeholder,
            'merge_guard_kept_existing': merge_guard_kept_existing,
            'merge_guard_updated_incoming': merge_guard_updated_incoming,
        }
    except Exception as e:
        db.session.rollback()
        tb = traceback.format_exc()
        current_app.logger.error(f"寫入資料庫中斷: {e}\n{tb}")
        queue.put(f"ERROR: 寫入資料庫中斷: {e}")
        return {}








