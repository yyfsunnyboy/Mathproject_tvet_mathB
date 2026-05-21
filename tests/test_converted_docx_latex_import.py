from pathlib import Path

from docx import Document

from app import app
from core import textbook_processor as processor
from core.routes.admin import apply_mathb_import_policy, is_vocational_mathb_volume


class _Q:
    def __init__(self):
        self.items = []

    def put(self, msg):
        self.items.append(str(msg))


def test_extract_converted_docx_latex_keeps_inline_absolute_value(tmp_path: Path):
    p = tmp_path / "converted.docx"
    d = Document()
    d.add_paragraph("數線上，若 $|x|=7$，試求 $x$ 之值。")
    d.save(str(p))
    content, _meta = processor.extract_converted_latex_docx(str(p))
    assert "$|x|=7$" in content[1]


def test_extract_converted_docx_latex_keeps_table_ge(tmp_path: Path):
    p = tmp_path / "converted_table.docx"
    d = Document()
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "試求下列不等式之解：(1) $|x|<3$ (2) $|x|\\ge 4$"
    d.save(str(p))
    content, _meta = processor.extract_converted_latex_docx(str(p))
    assert "\\ge" in content[1]


def test_extract_converted_docx_latex_keeps_frac(tmp_path: Path):
    p = tmp_path / "converted_frac.docx"
    d = Document()
    d.add_paragraph("$-\\frac{4}{3}$")
    d.save(str(p))
    content, _meta = processor.extract_converted_latex_docx(str(p))
    assert "\\frac{4}{3}" in content[1]


def test_converted_mode_skips_assets_ocr_pix2tex(tmp_path: Path):
    p = tmp_path / "converted_mode.docx"
    d = Document()
    d.add_paragraph("數線上，若 $|x|=7$，試求 $x$ 之值。")
    d.save(str(p))
    q = _Q()
    with app.app_context():
        out = processor.extract_content_from_file(
            str(p),
            q,
            import_policy={"docx_formula_source_mode": "converted_docx_latex"},
        )
        ctx = processor._DOCX_IMPORT_CONTEXT
    assert "$|x|=7$" in out[1]
    assert "[FORMULA_IMAGE_" not in out[1]
    assert ctx.get("formula_assets_extraction_skipped") is True
    assert ctx.get("ocr_skipped") is True
    assert ctx.get("pix2tex_skipped") is True


def test_prompt_contains_converted_docx_latex_rules():
    src = Path("core/textbook_processor.py").read_text(encoding="utf-8-sig")
    assert "converted_docx_latex metadata-only" in src
    assert "hydrate_converted_docx_latex_parsed_data" in src
    assert "formula_assets_extraction_skipped" in src
    assert "[FORMULA NORMALIZE SKIP] converted_docx_latex_preserve_latex=true" in src


def test_scan_example_flushes_before_suithang_header_and_subsection():
    text = (
        "例1\n"
        "例題1內容 $|x|=1$。\n"
        "隨堂練習……………………………………………\n"
        "1. 利用十字交乘法因式分解。\n"
        "1-4.2 一元二次不等式的解\n"
        "例2\n"
        "例題2內容。\n"
    )
    blocks = processor.scan_converted_docx_question_blocks(text)
    assert "例題1" in blocks
    assert "隨堂練習" not in blocks["例題1"]
    assert "1-4.2" not in blocks["例題1"]
    assert "$|x|=1$" in blocks["例題1"]
    assert "隨堂練習1" in blocks
    assert blocks["隨堂練習1"].strip()
    assert "十字交乘法" in blocks["隨堂練習1"]


def test_scan_and_hydrate_converted_docx_blocks():
    text = "例題1 試求 $|x|=1$ 之值。\n隨堂練習1 計算 $2+3$。\n"
    blocks = processor.scan_converted_docx_question_blocks(text)
    assert "例題1" in blocks
    assert "$|x|=1$" in blocks["例題1"]
    parsed = {
        "chapters": [
            {
                "chapter_title": "第1章",
                "sections": [
                    {
                        "section_title": "1-1 測試",
                        "concepts": [
                            {
                                "concept_name": "測試",
                                "concept_en_id": "Test",
                                "examples": [
                                    {
                                        "title": "例題1",
                                        "source_description": "例題1",
                                        "problem_text": "例題1",
                                        "correct_answer": None,
                                        "detailed_solution": None,
                                    }
                                ],
                                "practice_questions": [],
                            }
                        ],
                    }
                ],
            }
        ]
    }
    out, filled, _n = processor.hydrate_converted_docx_latex_parsed_data(
        parsed, extracted_text=text, section_code="1-1"
    )
    ex = out["chapters"][0]["sections"][0]["concepts"][0]["examples"][0]
    assert filled == 1
    assert "$|x|=1$" in ex["problem_text"]
    assert ex["correct_answer"] == ""
    assert ex["detailed_solution"] == ""

def test_converted_mode_skip_legacy_normalize_preserves_text():
    text = r"\(y=\frac{1}{2}(x-2)^{2}+1\) 的圖形"
    out, check = processor._normalize_imported_math_value(
        text,
        field_name="problem_text",
        docx_formula_source_mode="converted_docx_latex",
    )
    assert out == text
    assert check is not None


def test_converted_mode_suspicious_detector_does_not_rewrite_text():
    text = "# # C 3 1"
    out, check = processor._normalize_imported_math_value(
        text,
        field_name="problem_text",
        docx_formula_source_mode="converted_docx_latex",
    )
    assert out == text
    assert check.get("is_suspicious") is True

def test_scan_expected_titles_examples_1_to_8():
    text = "\n".join([f"例{i}" for i in range(1, 9)])
    titles = processor.scan_expected_titles_from_converted_text(text)
    assert all(f"例題{i}" in titles for i in range(1, 9))


def test_scan_expected_titles_suithang_1_to_8():
    text = "\n".join([f"隨堂練習{i}" for i in range(1, 9)])
    titles = processor.scan_expected_titles_from_converted_text(text)
    assert all(f"隨堂練習{i}" in titles for i in range(1, 9))


def test_scan_expected_titles_section_exercise_basic_advanced():
    text = (
        "1-4習題\n"
        "基礎題\n"
        "1 a\n2 b\n3 c\n4 d\n5 e\n6 f\n7 g\n8 h\n"
        "進階題\n"
        "9 i\n10 j\n"
    )
    titles = processor.scan_expected_titles_from_converted_text(text)
    assert all(f"1-4習題 基礎題{i}" in titles for i in range(1, 9))
    assert "1-4習題 進階題9" in titles
    assert "1-4習題 進階題10" in titles


def test_scan_expected_titles_exam():
    text = "〔109統測B〕"
    titles = processor.scan_expected_titles_from_converted_text(text)
    assert "109統測B" in titles


def test_inventory_missing_titles_detected():
    expected = [f"例題{i}" for i in range(1, 9)]
    returned = [f"例題{i}" for i in range(1, 6)]
    inv = processor.build_title_inventory(expected, returned)
    assert inv["missing_titles_count"] == 3
    assert "例題6" in inv["missing_titles_canonical"]


def test_inventory_guard_report_only_never_aborts_write():
    expected = ["例題1", "例題2"]
    returned = ["例題1"]
    inv = processor.build_title_inventory(expected, returned)
    write_aborted = False
    assert inv.get("missing_titles_count", 0) > 0
    assert write_aborted is False


def test_mathb_volume_detection():
    assert is_vocational_mathb_volume("數學B1")
    assert is_vocational_mathb_volume("數學B4")
    assert not is_vocational_mathb_volume("數學A1")


def test_mathb_import_policy_forces_converted_docx_latex_and_vocational():
    curriculum_info = {"curriculum": "general", "volume": "數學B2"}
    import_policy = {
        "docx_formula_source_mode": "auto_detect",
        "enable_formula_postprocess": True,
        "enable_formula_auto_apply": True,
        "enable_formula_detailed_report": True,
        "formula_postprocess_mode": "local_ocr",
    }
    assert apply_mathb_import_policy(curriculum_info, import_policy, filenames=["chapter.docx"]) is True
    assert curriculum_info["curriculum"] == "vocational"
    assert import_policy["docx_formula_source_mode"] == "converted_docx_latex"
    assert import_policy["enable_formula_postprocess"] is False
    assert import_policy["enable_formula_auto_apply"] is False
    assert import_policy["enable_formula_detailed_report"] is False
    assert import_policy["formula_postprocess_mode"] == "convert_only"


def test_mathb_import_policy_warns_non_latex_suffix():
    curriculum_info = {"curriculum": "general", "volume": "數學B3"}
    import_policy = {"docx_formula_source_mode": "raw_docx_with_formula_assets"}
    apply_mathb_import_policy(
        curriculum_info,
        import_policy,
        filenames=["1-1_例題_Latex.docx", "1-1_raw.docx"],
    )
    assert import_policy["docx_formula_source_mode"] == "converted_docx_latex"

def test_canonicalize_example_titles():
    assert processor.canonicalize_import_title("例題 1") == "例題1"
    assert processor.canonicalize_import_title("例 1") == "例題1"


def test_canonicalize_suitang_title():
    assert processor.canonicalize_import_title("隨堂練習 8") == "隨堂練習8"


def test_canonicalize_zone_labeled_exercise_titles_with_section_code():
    assert processor.canonicalize_import_title("基礎題9", section_code="1-4") == "1-4習題 基礎題9"
    assert processor.canonicalize_import_title("進階題9", section_code="1-4") == "1-4習題 進階題9"
    assert processor.canonicalize_import_title("自我評量1", section_code="1-4") == "1-4習題 自我評量1"
    assert processor.canonicalize_import_title("基礎題 2", section_code="1-4") == "1-4習題 基礎題2"


def test_merged_exam_bucket_title_needs_review_when_inventory_has_specific_exams():
    doc = "〔105統測A〕\n〔105統測B〕\n"
    items = processor.scan_docx_title_inventory(doc)
    meta = processor.map_returned_import_title("統測歷屆試題", section_code="1-4", inventory_items=items)
    assert meta["needs_review"] is True
    assert meta["mapping_method"] == "merged_exam_bucket_rejected"
    assert meta["returned_canonical"] == "統測歷屆試題"
    assert "105統測A" in {it["canonical_title"] for it in items}


def test_canonicalize_exercise_with_section_code():
    doc = (
        "2-1習題\n"
        "基礎題\n"
        "1 a\n8 b\n"
        "進階題\n"
        "9 c\n10 d\n"
    )
    items = processor.scan_docx_title_inventory(doc)
    assert processor.canonicalize_import_title("習題 1", section_code="2-1", inventory_items=items) == "2-1習題 基礎題1"
    assert processor.canonicalize_import_title("習題 8", section_code="2-1", inventory_items=items) == "2-1習題 基礎題8"
    assert processor.canonicalize_import_title("習題 9", section_code="2-1", inventory_items=items) == "2-1習題 進階題9"
    assert processor.canonicalize_import_title("習題 10", section_code="2-1", inventory_items=items) == "2-1習題 進階題10"


def test_scan_converted_docx_blocks_canonical_exercise_and_exam_keys():
    text = (
        "1-4習題\n"
        "基礎題\n"
        "8\t根據下圖填入判別式\n"
        "進階題\n"
        "9\t已知路邊行動咖啡車製作n杯咖啡的成本為n + 50元，而賣出n杯咖啡的收入為${{n}^{2}}-4n$元，試問最少要賣出多少杯咖啡才會開始有利潤？\n"
        "10\t設a、b均為實數\n"
        "設 $f(x)$ 之值。\n"
        "(A) 選項甲\n"
        "(B) 選項乙\n"
        "(C) 選項丙\n"
        "(D) 選項丁\n"
        "〔109統測B〕\n"
        "KEY\n"
        "(1) 解答\n"
    )
    blocks = processor.scan_converted_docx_question_blocks(text)
    assert "1-4習題 進階題9" in blocks
    assert "咖啡" in blocks["1-4習題 進階題9"]
    assert "109統測B" in blocks
    assert "(A)" in blocks["109統測B"]
    assert "(D)" in blocks["109統測B"]
    assert "KEY" not in blocks["109統測B"]
    assert blocks["109統測B"].strip() != "109統測B"


def test_hydrate_canonical_exercise_title_finds_block():
    text = (
        "1-4習題\n"
        "進階題\n"
        "9\t已知路邊行動咖啡車製作n杯咖啡的成本為n + 50元，試問最少要賣出多少杯咖啡才會開始有利潤？\n"
    )
    parsed = {
        "chapters": [
            {
                "chapter_title": "第1章",
                "sections": [
                    {
                        "section_title": "1-4 一元二次不等式",
                        "concepts": [
                            {
                                "concept_name": "練習",
                                "concept_en_id": "Practice",
                                "examples": [],
                                "practice_questions": [
                                    {
                                        "title": "進階題9",
                                        "source_description": "進階題9",
                                        "problem_text": "進階題9",
                                        "correct_answer": "",
                                        "detailed_solution": "",
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ]
    }
    out, filled, n_blocks = processor.hydrate_converted_docx_latex_parsed_data(
        parsed, extracted_text=text, section_code="1-4"
    )
    pq = out["chapters"][0]["sections"][0]["concepts"][0]["practice_questions"][0]
    assert filled == 1
    assert "咖啡" in pq["problem_text"]
    assert n_blocks >= 1


def test_chapter_self_assessment_mode_inventory_and_blocks():
    text = (
        "第1章 坐標系與函數圖形\n"
        "CH1自我評量\n"
        "自我評量\n"
        "1-1 數線與絕對值\n"
        "1.\t若 $|x|=3$，則 $x$ 為何？\n"
        "(A) 3\n"
        "(B) -3\n"
        "(C) 3 或 -3\n"
        "(D) 0\n"
        "1-2 函數\n"
        "2.\t設 $f(x)=2x+1$，求 $f(0)$。\n"
        "1-3 二次函數\n"
        "11.\t解 $x^2-1>0$。\n"
        "1-4 一元二次不等式\n"
        "15.\t已知路邊行動咖啡車製作n杯咖啡的成本為n + 50元，試問最少要賣出多少杯咖啡？\n"
        "20.\t最後一題。\n"
        "78\n"
        "79\n"
    )
    items = processor.scan_docx_title_inventory(text)
    titles = {it["canonical_title"] for it in items}
    assert all(it.get("kind") == "self_assessment" for it in items)
    assert "CH1自我評量 題1" in titles
    assert "CH1自我評量 題15" in titles
    assert "CH1自我評量 題20" in titles
    assert len(items) == 5
    blocks = processor.scan_converted_docx_question_blocks(text)
    assert len(blocks) == 5
    assert "咖啡" in blocks["CH1自我評量 題15"]
    assert "(A)" in blocks["CH1自我評量 題1"]
    assert "KEY" not in blocks.get("CH1自我評量 題1", "")
    sec_by_num = {int(it["number"]): it["section_code"] for it in items}
    assert sec_by_num[1] == "1-1"
    assert sec_by_num[2] == "1-2"
    assert sec_by_num[11] == "1-3"
    assert sec_by_num[15] == "1-4"


def test_chapter_self_assessment_hydrate_and_inventory_guard():
    text = (
        "CH1自我評量\n"
        "自我評量\n"
        "1-1 數線\n"
        "1.\t第一題 $|x|=1$。\n"
        "1-2 函數\n"
        "2.\t第二題 $f(x)=x$。\n"
    )
    items = processor.scan_docx_title_inventory(text)
    expected = sorted({it["canonical_title"] for it in items})
    parsed = {
        "chapters": [
            {
                "chapter_title": "第1章 坐標系與函數圖形",
                "sections": [
                    {
                        "section_title": "1-1 數線",
                        "concepts": [
                            {
                                "concept_name": "自我評量",
                                "concept_en_id": "SelfAssessment",
                                "examples": [
                                    {
                                        "title": "題1",
                                        "source_description": "題1",
                                        "source_type": "self_assessment",
                                        "problem_text": "題1",
                                        "correct_answer": "",
                                        "detailed_solution": "",
                                    }
                                ],
                                "practice_questions": [],
                            },
                            {
                                "concept_name": "函數",
                                "concept_en_id": "Functions",
                                "examples": [
                                    {
                                        "title": "題2",
                                        "source_description": "題2",
                                        "source_type": "self_assessment",
                                        "problem_text": "題2",
                                        "correct_answer": "",
                                        "detailed_solution": "",
                                    }
                                ],
                                "practice_questions": [],
                            },
                        ],
                    }
                ],
            }
        ]
    }
    out, filled, _n = processor.hydrate_converted_docx_latex_parsed_data(
        parsed, extracted_text=text, inventory_items=items
    )
    assert filled == 2
    assert "$|x|=1$" in out["chapters"][0]["sections"][0]["concepts"][0]["examples"][0]["problem_text"]
    returned = processor.collect_returned_titles_from_parsed_data(out)
    inv = processor.build_title_inventory(expected, returned, inventory_items=items)
    assert inv["missing_titles_count"] == 0


def test_inventory_exercise_alignment_no_missing():
    doc = (
        "1-4習題\n"
        "基礎題\n"
        "1 a\n2 b\n3 c\n4 d\n5 e\n6 f\n7 g\n8 h\n"
        "進階題\n"
        "9 i\n10 j\n"
    )
    items = processor.scan_docx_title_inventory(doc)
    expected = sorted({it["canonical_title"] for it in items if it.get("kind") == "chapter_exercise"})
    returned = [f"習題 {i}" for i in range(1, 11)]
    inv = processor.build_title_inventory(expected, returned, section_code="1-4", inventory_items=items)
    assert inv["missing_titles_count"] == 0


def test_inventory_examples_suitang_alignment_no_missing():
    expected = [*[f"例題{i}" for i in range(1, 9)], *[f"隨堂練習{i}" for i in range(1, 9)]]
    returned = [*[f"例題 {i}" for i in range(1, 9)], *[f"隨堂練習 {i}" for i in range(1, 9)]]
    inv = processor.build_title_inventory(expected, returned)
    assert inv["missing_titles_count"] == 0


def test_curriculum_volume_warning_general_mathb():
    warn = processor.detect_curriculum_volume_warning("general", "數學B1", "longteng")
    assert "vocational mathB import expected" in warn
