# -*- coding: utf-8 -*-
import os
import pytest
from docx import Document
from flask import Flask

import core.textbook_processor as processor
from core.question_image_assets import (
    build_question_asset_dir,
    build_question_asset_filename,
    safe_slug,
)


def test_build_question_asset_dir_vocational():
    # "數學B4", "1 排列組合", "1-2 直線排列"
    p = build_question_asset_dir("vocational", "longteng", "\u6578\u5b78B4", "1 \u6392\u5217\u7d44\u5408", "1-2 \u76f4\u7dda\u6392\u5217")
    assert p == "uploads/question_assets/vocational/longteng/\u6578\u5b78B4/ch01_\u6392\u5217\u7d44\u5408/sec_1-2_\u76f4\u7dda\u6392\u5217"


def test_build_question_asset_dir_junior():
    # "數學2上", "第3章 一元二次方程式", "3-1 因式分解法"
    p = build_question_asset_dir("junior", "kangxuan", "\u6578\u5b782\u4e0a", "\u7b2c3\u7ae0 \u4e00\u5143\u4e8c\u6b21\u65b9\u7a0b\u5f0f", "3-1 \u56e0\u5f0f\u5206\u89e3\u6cd5")
    assert p == "uploads/question_assets/junior/kangxuan/\u6578\u5b782\u4e0a/ch03_\u4e00\u5143\u4e8c\u6b21\u65b9\u7a0b\u5f0f/sec_3-1_\u56e0\u5f0f\u5206\u89e3\u6cd5"


def test_safe_slug_windows_chars_and_chinese():
    # '1-2 直線排列 < > : " / \ | ? *'
    s = safe_slug('1-2 \u76f4\u7dda\u6392\u5217 < > : " / \\ | ? *')
    assert s == "1-2_\u76f4\u7dda\u6392\u5217"


def test_build_question_asset_filename_not_conflict_by_hash():
    # "例題7"
    f1 = build_question_asset_filename("textbook_example", "\u4f8b\u984c7", "a1b2c3d4", 1, "png")
    f2 = build_question_asset_filename("textbook_example", "\u4f8b\u984c7", "f3e9aa21", 1, "png")
    assert f1 != f2
    assert f1.endswith(".png") and f2.endswith(".png")


def test_b1_coordinate_guard_normalizes_point_sup_sub_forms():
    # "已知點 {}^{a}P_{b} 在第二象限內，設 A(-1,2), B(3,3), C^{2}_{1} 為三頂點。"
    text = "\u5df2\u77e5\u9ede {}^{a}P_{b} \u5728\u7b2c\u4e8c\u8c61\u9650\u5167\uff0c\u8a2d A(-1,2), B(3,3), C^{2}_{1} \u70ba\u4e09\u9802\u9ede\u3002"
    normalized, _meta = processor.normalize_permutation_combination_notation(
        text, volume="\u6578\u5b78B1", section_title="1-2 \u5e73\u9762\u5750\u6a19\u7cfb\u8207\u7dda\u578b\u51fd\u6578"
    )
    assert "P(a,b)" in normalized
    assert "C(1,2)" in normalized
    assert "{}^{a}P_{b}" not in normalized
    assert "C^{2}_{1}" not in normalized


def test_b1_coordinate_guard_keeps_existing_coordinate_points():
    # "點 C(3,1)、P(a,b)、Q(b,a)、R(-b,a^2) 在平面上。"
    text = "\u9ede C(3,1)\u3001P(a,b)\u3001Q(b,a)\u3001R(-b,a^2) \u5728\u5e73\u9762\u4e0a\u3002"
    normalized, _meta = processor.normalize_permutation_combination_notation(
        text, volume="\u6578\u5b78B1", section_title="1-2 \u5e73\u9762\u5750\u6a19\u7cfb\u8207\u7dda\u578b\u51fd\u6578"
    )
    assert "C(3,1)" in normalized
    assert "P(a,b)" in normalized
    assert "Q(b,a)" in normalized
    assert "R(-b,a^2)" in normalized


def test_b4_combination_notation_not_affected_by_b1_coordinate_guard():
    # "從 7 人中選 3 人排列，共有 P(7,3) 種；任取 3 人，共有 C(7,3) 種。"
    text = "\u5f9e 7 \u4eba\u4e2d\u9078 3 \u4eba\u6392\u5217\uff0c\u5171\u6709 P(7,3) \u7a2e\uff1b\u4efb\u53d6 3 \u4eba\uff0c\u5171\u6709 C(7,3) \u7a2e\u3002"
    normalized, _meta = processor.normalize_permutation_combination_notation(
        text, volume="\u6578\u5b78B4", section_title="1-2 \u76f4\u7dda\u6392\u5217"
    )
    assert "P^{7}_{3}" in normalized
    assert "C^{7}_{3}" in normalized


def test_is_section_exposition_title_kewenneirong():
    """'課文內容' must be identified as a section-exposition title."""
    # "課文內容"
    assert processor._is_section_exposition_title("\u8ab2\u6587\u5167\u5bb9")


def test_is_section_exposition_title_kewenshuo():
    """'課文說明' must be identified as a section-exposition title."""
    # "課文說明"
    assert processor._is_section_exposition_title("\u8ab2\u6587\u8aaa\u660e")


def test_is_section_exposition_title_shuoming():
    """'說明' must be identified as a section-exposition title."""
    # "說明"
    assert processor._is_section_exposition_title("\u8aaa\u660e")


def test_is_section_exposition_title_example_not_exposition():
    """'例題1' must NOT be an exposition title."""
    # "例題1"
    assert not processor._is_section_exposition_title("\u4f8b\u984c1")


def test_is_section_exposition_title_jichu_not_exposition():
    """'1-1習題 基礎題1' must NOT be an exposition title."""
    # "1-1習題 基礎題1"
    assert not processor._is_section_exposition_title("1-1\u7fd2\u984c \u57fa\u790e\u984c1")


def test_is_section_exposition_title_practice_not_exposition():
    """'隨堂練習1' must NOT be an exposition title."""
    # "隨堂練習1"
    assert not processor._is_section_exposition_title("\u96a8\u5802\u7df4\u7fd21")


def test_is_section_exposition_title_empty_not_exposition():
    """Empty string must not be flagged as exposition."""
    assert not processor._is_section_exposition_title("")


def test_normalize_source_type_exposition_returns_section_exposition():
    """normalize_source_type_by_title on '課文內容' item must return 'section_exposition'."""
    # "課文內容", "若 a > b，則 a+c > b+c"
    item = {"source_description": "\u8ab2\u6587\u5167\u5bb9", "problem_text": "\u82e5 a > b\uff0c\u5247 a+c > b+c"}
    result = processor.normalize_source_type_by_title(item, default_source_type="textbook_example")
    assert result == "section_exposition"


def test_normalize_source_type_liti_returns_textbook_example():
    """normalize_source_type_by_title on '例題1' must return 'textbook_example'."""
    # "例題1", "解不等式 |x-1| < 2"
    item = {"source_description": "\u4f8b\u984c1", "problem_text": "\u89e3\u4e0d\u7b49\u5f0f |x-1| < 2"}
    result = processor.normalize_source_type_by_title(item, default_source_type="textbook_example")
    assert result == "textbook_example"


def test_normalize_source_type_jichu_returns_basic_exercise():
    """normalize_source_type_by_title on '1-1習題 基礎題1' must return 'basic_exercise'."""
    # "1-1習題 基礎題1", "解不等式"
    item = {"source_description": "1-1\u7fd2\u984c \u57fa\u790e\u984c1", "problem_text": "\u89e3\u4e0d\u7b49\u5f0f"}
    result = processor.normalize_source_type_by_title(item, default_source_type="textbook_example")
    assert result == "basic_exercise"


def test_normalize_source_type_practice_returns_in_class_practice():
    """normalize_source_type_by_title on '隨堂練習1' must return 'in_class_practice'."""
    # "隨堂練習1", "求解"
    item = {"source_description": "\u96a8\u5802\u7df4\u7fd21", "problem_text": "\u6c42\u89e3"}
    result = processor.normalize_source_type_by_title(item, default_source_type="in_class_practice")
    assert result == "in_class_practice"


def test_extract_converted_latex_docx():
    doc_path = "test_temp_extract.docx"
    doc = Document()
    doc.add_paragraph("This is a paragraph.")
    doc.add_paragraph("This has some LaTeX: \\(x+1\\).")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "Table cell text"
    doc.save(doc_path)

    try:
        content_by_page, doc_meta = processor.extract_converted_latex_docx(doc_path)
        assert 1 in content_by_page
        text = content_by_page[1]
        assert "This is a paragraph." in text
        assert "This has some LaTeX: \\(x+1\\)." in text
        assert "Table cell text" in text
        assert doc_meta["paragraph_count"] >= 2
        assert doc_meta["table_count"] == 1
    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_detect_converted_latex_docx():
    text_with_latex = "Let \\(a = 1\\) and \\(b = 2\\). Therefore, \\(a + b = 3\\). Also \\frac{1}{2}."
    meta = processor.detect_converted_latex_docx(text_with_latex)
    assert meta["is_converted_latex_docx"] is True
    assert meta["latex_signal_count"] >= 3

    text_with_placeholders = "[FORMULA_IMAGE_1] [FORMULA_IMAGE_2] [FORMULA_IMAGE_3]"
    meta2 = processor.detect_converted_latex_docx(text_with_placeholders)
    assert meta2["is_converted_latex_docx"] is False
    assert meta2["formula_placeholder_count"] == 3


def test_extract_content_from_file():
    doc_path = "test_temp_content.docx"
    doc = Document()
    doc.add_paragraph("This is a paragraph with LaTeX: \\(x^2 + y^2 = z^2\\), \\(a \\le b\\), \\(c \\ge d\\).")
    doc.save(doc_path)

    app = Flask(__name__)
    class DummyQueue:
        def __init__(self):
            self.msgs = []
        def put(self, msg):
            self.msgs.append(msg)

    try:
        with app.app_context():
            res = processor.extract_content_from_file(doc_path, DummyQueue())
            assert 1 in res
            assert "This is a paragraph with LaTeX" in res[1]
    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)
