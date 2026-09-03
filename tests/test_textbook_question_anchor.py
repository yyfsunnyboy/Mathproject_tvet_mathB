# -*- coding: utf-8 -*-
from pathlib import Path

import pytest

from core.textbook_question_anchor import (
    build_anchor_id,
    build_anchors_from_block_meta,
    build_pdf_match_stub,
    build_question_anchor,
    build_text_fingerprint,
    canonicalize_question_type,
    collect_b2_11_question_anchors,
    detect_anchor_id_collisions,
    extract_question_number,
    normalize_fingerprint_text,
    normalize_question_label,
    question_anchor_notes_payload,
    source_order_is_monotonic,
    summarize_anchor_collection,
    validate_question_anchor,
)


def test_normalize_question_label_example_and_practice():
    assert normalize_question_label("例題 1") == "例1"
    assert normalize_question_label("例 1") == "例1"
    assert normalize_question_label("隨堂練習 1") == "隨堂練習1"
    assert normalize_question_label("1-1 習題 基礎題 3") == "1-1習題 基礎題3"


def test_example_and_in_class_practice_do_not_share_anchor_id():
    example = build_question_anchor(
        question_label="例1",
        problem_text="已知角 A",
        source_order=1,
        volume="數學B2",
        section="1-1",
        source_type="textbook_example",
    )
    practice = build_question_anchor(
        question_label="隨堂練習1",
        problem_text="已知角 A",
        source_order=2,
        volume="數學B2",
        section="1-1",
        source_type="in_class_practice",
    )
    assert example["question_type"] == "example"
    assert practice["question_type"] == "in_class_practice"
    assert example["question_number"] == "1"
    assert practice["question_number"] == "1"
    assert example["anchor_id"] != practice["anchor_id"]
    assert example["anchor_id"] == "vocational_math_B2_1-1_example_001_001"
    assert practice["anchor_id"] == "vocational_math_B2_1-1_in_class_practice_001_002"


def test_duplicate_numbers_across_types_and_source_order_tie_breaker():
    first = build_question_anchor(
        question_label="例1",
        problem_text="第一題",
        source_order=3,
        volume="數學B2",
        section="1-1",
        source_type="textbook_example",
    )
    second = build_question_anchor(
        question_label="例1",
        problem_text="重複出現的例1",
        source_order=7,
        volume="數學B2",
        section="1-1",
        source_type="textbook_example",
    )
    assert first["anchor_id"] != second["anchor_id"]
    assert first["anchor_id"].endswith("_003")
    assert second["anchor_id"].endswith("_007")


def test_self_assessment_does_not_collide_with_section_example():
    example = build_question_anchor(
        question_label="例1",
        problem_text="例題",
        source_order=1,
        volume="數學B2",
        section="1-1",
        source_type="textbook_example",
    )
    assessment = build_question_anchor(
        question_label="CH1自我評量 題1",
        problem_text="評量",
        source_order=1,
        volume="數學B2",
        section="1-1",
        source_type="self_assessment",
    )
    assert canonicalize_question_type("self_assessment") == "self_assessment"
    assert assessment["question_type"] == "self_assessment"
    assert example["anchor_id"] != assessment["anchor_id"]


def test_malformed_missing_question_number_uses_x_and_source_order():
    row = build_question_anchor(
        question_label="補充說明",
        problem_text="沒有題號的區塊",
        source_order=4,
        volume="數學B2",
        section="1-1",
        source_type="unknown",
        question_type="unknown",
    )
    assert extract_question_number("補充說明") == ""
    assert row["question_number"] == ""
    assert "_x_004" in row["anchor_id"]
    ok, errors = validate_question_anchor(row)
    assert ok, errors


def test_fingerprint_whitespace_and_line_breaks_do_not_change_hash():
    a = build_text_fingerprint("已知  角 A\n等於 90")
    b = build_text_fingerprint("已知 角 A 等於 90")
    c = build_text_fingerprint("已知 角 A\r\n等於 90")
    assert a == b == c


def test_fingerprint_preserves_chinese_and_latex_contribution():
    chinese = normalize_fingerprint_text("角度的基本性質 $\\theta$")
    assert "角度的基本性質" in chinese
    assert "\\theta" in chinese
    with_latex = build_text_fingerprint("求 $\\theta$")
    without_latex = build_text_fingerprint("求 theta")
    assert with_latex != without_latex
    assert build_text_fingerprint("角度") != build_text_fingerprint("角")


def test_same_input_twice_identical_anchor_and_fingerprint():
    kwargs = dict(
        question_label="例3",
        problem_text="若 $\\angle A=30^\\circ$",
        source_order=7,
        curriculum="vocational",
        publisher="longteng",
        volume="數學B2",
        chapter="1",
        section="1-1",
        source_type="textbook_example",
    )
    first = build_question_anchor(**kwargs)
    second = build_question_anchor(**kwargs)
    assert first == second
    assert first["text_fingerprint"] == second["text_fingerprint"]


def test_validate_and_collision_detection():
    rows = build_anchors_from_block_meta(
        {
            "例1": {
                "anchor": "例1",
                "source_type": "textbook_example",
                "problem_text": "第一題",
                "section_code": "1-1",
            },
            "隨堂練習1": {
                "anchor": "隨堂練習1",
                "source_type": "in_class_practice",
                "problem_text": "練習",
                "section_code": "1-1",
            },
            "1-1習題 基礎題 1": {
                "anchor": "1-1習題 基礎題 1",
                "source_type": "textbook_exercise",
                "problem_text": "習題",
                "section_code": "1-1",
            },
        },
        {"curriculum": "vocational", "publisher": "longteng", "volume": "數學B2", "chapter_index": 1},
    )
    assert len(rows) == 3
    assert source_order_is_monotonic(rows)
    assert detect_anchor_id_collisions(rows) == []
    summary = summarize_anchor_collection(rows)
    assert summary["collision_count"] == 0
    assert all(validate_question_anchor(row)[0] for row in rows)
    stub = build_pdf_match_stub(rows[0])
    assert stub["pdf_page"] is None
    assert stub["anchor_id"] == rows[0]["anchor_id"]
    notes = question_anchor_notes_payload(rows[0])
    assert notes["question_anchor"]["anchor_id"] == rows[0]["anchor_id"]


def test_source_order_tie_breaker_on_duplicate_meta_ids():
    duplicate_id = build_anchor_id(
        curriculum="vocational",
        volume="數學B2",
        section="1-1",
        question_type="example",
        question_number="1",
        source_order=1,
    )
    collided = [
        {
            "anchor_id": duplicate_id,
            "curriculum": "vocational",
            "publisher": "longteng",
            "volume": "數學B2",
            "chapter": "1",
            "section": "1-1",
            "question_type": "example",
            "question_number": "1",
            "source_order": 1,
            "question_label": "例1",
            "text_fingerprint": build_text_fingerprint("a"),
        },
        {
            "anchor_id": duplicate_id,
            "curriculum": "vocational",
            "publisher": "longteng",
            "volume": "數學B2",
            "chapter": "1",
            "section": "1-1",
            "question_type": "example",
            "question_number": "1",
            "source_order": 1,
            "question_label": "例1",
            "text_fingerprint": build_text_fingerprint("b"),
        },
    ]
    assert detect_anchor_id_collisions(collided) == [duplicate_id]
    from core.textbook_question_anchor import _apply_occurrence_tie_breakers

    fixed = _apply_occurrence_tie_breakers(collided)
    assert detect_anchor_id_collisions(fixed) == []
    assert fixed[1]["occurrence_index"] == 2
    assert fixed[1]["anchor_id"].endswith("_occ02")


def test_real_phase1_phase2_docx_produces_stable_anchors(tmp_path: Path):
    from docx import Document

    from core.textbook_question_anchor import collect_question_anchors_from_latex_docx

    path = tmp_path / "第一章 1-1 角度的基本性質-課本_Latex.docx"
    doc = Document()
    doc.add_paragraph("例1")
    doc.add_paragraph(r"已知 $\angle A=30^{\circ}$")
    doc.add_paragraph("隨堂練習1")
    doc.add_paragraph("求此角的餘角")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "例2 表格中的角"
    doc.save(str(path))

    info = {
        "curriculum": "vocational",
        "publisher": "longteng",
        "grade": 10,
        "volume": "數學B2",
        "section_code": "1-1",
        "chapter_index": 1,
        "source_scope": "section_textbook",
    }
    first = collect_question_anchors_from_latex_docx(path, info)
    second = collect_question_anchors_from_latex_docx(path, info)
    assert first.get("status") == "ok", first
    assert first["phase2_block_count"] >= 2
    labels = [row["question_label"] for row in first["anchors"]]
    assert "例1" in labels
    assert "隨堂練習1" in labels
    assert first["summary"]["collision_count"] == 0
    assert first["summary"]["empty_fingerprint_count"] == 0
    assert [row["anchor_id"] for row in first["anchors"]] == [
        row["anchor_id"] for row in second["anchors"]
    ]
    assert [row["text_fingerprint"] for row in first["anchors"]] == [
        row["text_fingerprint"] for row in second["anchors"]
    ]


def test_table_cell_block_can_be_anchored():
    rows = build_anchors_from_block_meta(
        {
            "例2": {
                "anchor": "例2",
                "source_type": "textbook_example",
                "problem_text": "表格內公式 $\\angle ABC$",
                "section_code": "1-1",
            }
        },
        {"volume": "數學B2", "section_code": "1-1", "chapter_index": 1},
    )
    assert len(rows) == 1
    assert rows[0]["question_label"] == "例2"
    assert rows[0]["text_fingerprint"]


@pytest.mark.integration
def test_b2_11_phase2_anchors_if_latex_available():
    project_root = Path(__file__).resolve().parents[1]
    try:
        from app import app as flask_app
    except Exception as exc:
        flask_app = None
        _ = exc
    result = collect_b2_11_question_anchors(project_root, app=flask_app)
    if result.get("status") == "missing_latex_docx":
        pytest.skip("B2 1-1 *_Latex.docx not found in textbook_import/source/vocational/math_B2/")

    print(result.get("report") or result)
    assert result.get("status") == "ok", result
    anchors = result["anchors"]
    assert anchors
    assert result["summary"]["collision_count"] == 0
    assert result["summary"]["source_order_monotonic"] is True
    assert result["summary"]["empty_fingerprint_count"] == 0

    second = collect_b2_11_question_anchors(project_root, app=flask_app)
    first_ids = [row["anchor_id"] for row in anchors]
    second_ids = [row["anchor_id"] for row in second["anchors"]]
    first_fp = [row["text_fingerprint"] for row in anchors]
    second_fp = [row["text_fingerprint"] for row in second["anchors"]]
    assert first_ids == second_ids
    assert first_fp == second_fp
