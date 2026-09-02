# -*- coding: utf-8 -*-
"""Unit tests for deterministic MathType / EQ → LaTeX converter."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from core.mtef import MTEF, equation_native_to_latex, mtef_bytes_to_latex
from core.mtef.record import MtAST, MtChar, MtLine, MtTmpl, RecordType, SelectorType
from core.textbook_importer_v3_docx import find_reference_docx_in_storage
from core.textbook_mathtype_converter import (
    convert_docx_mathtype_to_latex_docx,
    convert_eq_instruction_to_latex,
    wrap_latex_for_v2,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
O_NS = "urn:schemas-microsoft-com:office:office"
V_NS = "urn:schemas-microsoft-com:vml"
REL_PKG_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _ast_char(mtcode: int, typeface: int = 128 + 8) -> MtAST:
    ch = MtChar()
    ch.mtcode = mtcode
    ch.typeface = typeface
    return MtAST(RecordType.CHAR, ch, None)


def _ast_line(children: list[MtAST], null: bool = False) -> MtAST:
    line = MtLine()
    line.null = null
    node = MtAST(RecordType.LINE, line, None)
    node.children = children
    return node


def _render_ast(root_children: list[MtAST]) -> str:
    eqn = MTEF()
    eqn.Valid = True
    root = MtAST(RecordType.ROOT, None, None)
    root.children = root_children
    eqn.ast = root
    return eqn.Translate()


def test_wrap_latex_for_v2_uses_inline_delimiters():
    assert wrap_latex_for_v2(r"$ \frac{1}{2} $") == r"\(\frac{1}{2}\)"
    assert wrap_latex_for_v2(r"\alpha") == r"\(\alpha\)"
    assert wrap_latex_for_v2("  ") == ""


def test_simple_token_and_operators():
    latex = _render_ast([_ast_line([_ast_char(ord("1")), _ast_char(ord("+")), _ast_char(ord("2"))])])
    assert "1" in latex and "+" in latex and "2" in latex


def test_fraction_template():
    tmpl = MtTmpl()
    tmpl.selector = SelectorType.tmFRACT
    node = MtAST(RecordType.TMPL, tmpl, None)
    node.children = [
        _ast_line([_ast_char(ord("1"))]),
        _ast_line([_ast_char(ord("2"))]),
    ]
    latex = _render_ast([_ast_line([node])])
    assert r"\frac" in latex
    assert "1" in latex and "2" in latex


def test_superscript_and_subscript_templates():
    sup = MtTmpl()
    sup.selector = SelectorType.tmSUP
    sup_node = MtAST(RecordType.TMPL, sup, None)
    # children: [sub-slot, sup-slot] per upstream makeLatex
    sup_node.children = [_ast_line([]), _ast_line([_ast_char(ord("2"))])]
    latex_sup = _render_ast([_ast_line([_ast_char(ord("x")), sup_node])])
    assert "^" in latex_sup

    sub = MtTmpl()
    sub.selector = SelectorType.tmSUB
    sub_node = MtAST(RecordType.TMPL, sub, None)
    sub_node.children = [_ast_line([_ast_char(ord("1"))]), _ast_line([])]
    latex_sub = _render_ast([_ast_line([_ast_char(0x03B8, 128 + 4), sub_node])])
    assert "_" in latex_sub


def test_sqrt_template():
    tmpl = MtTmpl()
    tmpl.selector = SelectorType.tmROOT
    node = MtAST(RecordType.TMPL, tmpl, None)
    node.children = [_ast_line([_ast_char(ord("2"))]), _ast_line([])]
    latex = _render_ast([_ast_line([node])])
    assert r"\sqrt" in latex


def test_greek_and_degree_char_maps():
    latex_pi = _render_ast([_ast_line([_ast_char(0x03C0, 128 + 4)])])
    assert r"\pi" in latex_pi
    latex_deg = _render_ast([_ast_line([_ast_char(0x00B0, 128 + 6)])])
    assert r"\circ" in latex_deg


def test_trig_as_function_chars():
    # typeface function (fnFUNCTION=2)
    chars = [_ast_char(ord(c), 128 + 2) for c in "sin"]
    latex = _render_ast([_ast_line(chars + [_ast_char(ord("x"), 128 + 3)])])
    assert "sin" in latex


def test_nested_fraction_expression():
    inner = MtTmpl()
    inner.selector = SelectorType.tmFRACT
    inner_node = MtAST(RecordType.TMPL, inner, None)
    inner_node.children = [
        _ast_line([_ast_char(0x03C0, 128 + 4)]),
        _ast_line([_ast_char(ord("2"))]),
    ]
    outer = MtTmpl()
    outer.selector = SelectorType.tmFRACT
    outer_node = MtAST(RecordType.TMPL, outer, None)
    outer_node.children = [
        _ast_line([inner_node]),
        _ast_line([_ast_char(ord("3"))]),
    ]
    latex = _render_ast([_ast_line([outer_node])])
    assert latex.count(r"\frac") >= 2


def test_eq_field_fraction_and_root():
    latex, err = convert_eq_instruction_to_latex(r"EQ \f(1,2)")
    assert err is None
    assert latex == r"\frac{1}{2}"
    latex, err = convert_eq_instruction_to_latex(r"EQ \r(x)")
    assert err is None
    assert r"\sqrt" in latex


def test_malformed_mtef_graceful_failure():
    latex, meta = equation_native_to_latex(b"not-an-equation-native-stream")
    assert latex == ""
    assert meta.get("error")


def _minimal_docx_with_ole(ole_bytes: bytes, *, in_table: bool = False) -> bytes:
    content_types = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="{CT_NS}">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Default Extension="bin" ContentType="application/vnd.openxmlformats-officedocument.oleObject"/>
  <Default Extension="wmf" ContentType="image/x-wmf"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    root_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{REL_PKG_NS}">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    doc_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{REL_PKG_NS}">
  <Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="embeddings/oleObject1.bin"/>
  <Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image1.wmf"/>
</Relationships>"""
    object_xml = f"""
      <w:r>
        <w:object>
          <v:shape><v:imagedata r:id="rId4"/></v:shape>
          <o:OLEObject Type="Embed" ProgID="Equation.DSMT4" r:id="rId5"/>
        </w:object>
      </w:r>"""
    if in_table:
        body = f"""
    <w:tbl>
      <w:tblPr/>
      <w:tblGrid><w:gridCol w:w="5000"/></w:tblGrid>
      <w:tr>
        <w:tc>
          <w:tcPr><w:tcW w:w="5000" w:type="dxa"/></w:tcPr>
          <w:p>
            <w:r><w:t>cell</w:t></w:r>
            {object_xml}
          </w:p>
        </w:tc>
      </w:tr>
    </w:tbl>"""
    else:
        body = f"""
    <w:p>
      <w:r><w:t>before</w:t></w:r>
      {object_xml}
      <w:r><w:t>after</w:t></w:r>
    </w:p>"""
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="{W_NS}" xmlns:r="{R_NS}" xmlns:o="{O_NS}" xmlns:v="{V_NS}">
  <w:body>
    {body}
    <w:sectPr/>
  </w:body>
</w:document>"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", root_rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/embeddings/oleObject1.bin", ole_bytes)
        zf.writestr("word/media/image1.wmf", b"WMF")
    return buf.getvalue()


def test_ole_relationship_resolution_and_equation_native_extraction(tmp_path: Path):
    ref = find_reference_docx_in_storage(PROJECT_ROOT)
    if ref is None:
        pytest.skip("reference B2 1-1 DOCX missing")
    with zipfile.ZipFile(ref) as zf:
        ole_bytes = zf.read("word/embeddings/oleObject10.bin")
    latex, meta = mtef_bytes_to_latex(ole_bytes)
    assert meta.get("valid")
    assert r"\circ" in latex or "60" in latex

    src = tmp_path / "sample.docx"
    out = tmp_path / "sample_Latex.docx"
    src.write_bytes(_minimal_docx_with_ole(ole_bytes))
    before = src.read_bytes()
    report = convert_docx_mathtype_to_latex_docx(src, out)
    assert src.read_bytes() == before
    assert report["converted_ok"] == 1
    assert out.is_file()
    # python-docx can open output
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert r"\(" in text


def test_table_cell_formula_conversion(tmp_path: Path):
    ref = find_reference_docx_in_storage(PROJECT_ROOT)
    if ref is None:
        pytest.skip("reference B2 1-1 DOCX missing")
    with zipfile.ZipFile(ref) as zf:
        ole_bytes = zf.read("word/embeddings/oleObject10.bin")
    src = tmp_path / "table.docx"
    out = tmp_path / "table_Latex.docx"
    src.write_bytes(_minimal_docx_with_ole(ole_bytes, in_table=True))
    report = convert_docx_mathtype_to_latex_docx(src, out)
    assert report["converted_ok"] == 1
    doc = Document(str(out))
    cell_text = doc.tables[0].cell(0, 0).text
    assert r"\(" in cell_text


def test_existing_v2_style_extractor_can_read_output(tmp_path: Path):
    ref = find_reference_docx_in_storage(PROJECT_ROOT)
    if ref is None:
        pytest.skip("reference B2 1-1 DOCX missing")
    out = tmp_path / "ref_Latex.docx"
    report = convert_docx_mathtype_to_latex_docx(ref, out)
    assert report["converted_ok"] >= 139
    assert report["equation_native_ok"] == 142
    assert report["original_unchanged"] is True

    doc = Document(str(out))
    lines: list[str] = []
    for block in doc.element.body.iterchildren():
        if block.tag.endswith("}p"):
            t = Paragraph(block, doc).text.strip()
            if t:
                lines.append(t)
        elif block.tag.endswith("}tbl"):
            tbl = Table(block, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        t = para.text.strip()
                        if t:
                            lines.append(t)
    joined = "\n".join(lines)
    assert r"\(" in joined
    assert r"\frac" in joined or r"\pi" in joined or r"\overline" in joined
    assert "MATH_PARSE_FAILED" in joined or report["converted_failed"] == 0
