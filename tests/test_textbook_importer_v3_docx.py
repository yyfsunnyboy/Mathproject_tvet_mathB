# -*- coding: utf-8 -*-
import io
import os
import zipfile
from pathlib import Path

import pytest
from docx import Document

from core.textbook_importer_v3_docx import (
    REFERENCE_STATISTICS,
    compare_reference_statistics,
    find_reference_docx_in_storage,
    format_parse_summary_text,
    parse_docx_structure,
    parse_docx_summary,
    parse_reference_docx_from_storage,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
O_NS = "urn:schemas-microsoft-com:office:office"
V_NS = "urn:schemas-microsoft-com:vml"
REL_PKG_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

EXPECTED_REFERENCE_STATS = REFERENCE_STATISTICS


def _build_minimal_docx_zip(document_xml: str, *, extra_files: dict[str, bytes] | None = None) -> bytes:
    content_types = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="{CT_NS}">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    root_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{REL_PKG_NS}">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    doc_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{REL_PKG_NS}">
  <Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="embeddings/oleObject1.bin"/>
  <Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image1.wmf"/>
</Relationships>"""

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", root_rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/embeddings/oleObject1.bin", b"OLE")
        zf.writestr("word/media/image1.wmf", b"WMF")
        for path, data in (extra_files or {}).items():
            zf.writestr(path, data)
    return buf.getvalue()


def test_ordered_blocks_preserve_paragraph_table_order(tmp_path: Path):
    p = tmp_path / "ordered.docx"
    d = Document()
    d.add_paragraph("Paragraph A")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "Cell text"
    d.add_paragraph("Paragraph B")
    d.save(str(p))

    parsed = parse_docx_structure(str(p))
    block_types = [block["type"] for block in parsed["blocks"]]
    assert block_types == ["paragraph", "table", "paragraph"]
    assert parsed["blocks"][0]["plain_text"] == "Paragraph A"
    assert parsed["blocks"][2]["plain_text"] == "Paragraph B"
    assert parsed["blocks"][1]["rows"][0]["cells"][0]["paragraphs"][0]["plain_text"] == "Cell text"


def test_eq_field_complete_recognition():
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="{W_NS}" xmlns:r="{R_NS}">
  <w:body>
    <w:p>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> EQ \\f(1,2) </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>display</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>
    <w:sectPr/>
  </w:body>
</w:document>"""
    parsed = parse_docx_structure(_build_minimal_docx_zip(document_xml))
    eq_fields = [f for f in parsed["formulas"] if f["kind"] == "eq_field"]
    assert len(eq_fields) == 1
    assert "EQ" in eq_fields[0]["instruction"]
    assert eq_fields[0]["display_text"] == "display"


def test_mathtype_ole_discovery_and_media_classification():
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="{W_NS}" xmlns:r="{R_NS}" xmlns:o="{O_NS}" xmlns:v="{V_NS}">
  <w:body>
    <w:p>
      <w:r>
        <w:object>
          <v:shape><v:imagedata r:id="rId4"/></v:shape>
          <o:OLEObject Type="Embed" ProgID="Equation.DSMT4" r:id="rId5"/>
        </w:object>
      </w:r>
    </w:p>
    <w:tbl>
      <w:tr>
        <w:tc>
          <w:p>
            <w:r>
              <w:object>
                <v:shape><v:imagedata r:id="rId4"/></v:shape>
                <o:OLEObject Type="Embed" ProgID="Equation.DSMT4" r:id="rId5"/>
              </w:object>
            </w:r>
          </w:p>
        </w:tc>
      </w:tr>
    </w:tbl>
    <w:sectPr/>
  </w:body>
</w:document>"""
    parsed = parse_docx_structure(
        _build_minimal_docx_zip(
            document_xml,
            extra_files={"word/media/image2.png": b"PNG"},
        )
    )
    ole_formulas = parsed.get("mathtype_oles") or [f for f in parsed["formulas"] if f["kind"] == "mathtype_ole"]
    assert len(ole_formulas) == 2
    assert ole_formulas[0]["prog_id"] == "Equation.DSMT4"
    assert ole_formulas[0]["embedding_path"] == "word/embeddings/oleObject1.bin"
    assert ole_formulas[1]["location"]["table_index"] == 0
    assert parsed["summary"]["mathtype_ole_in_table_cells"] == 1

    media_types = {m["path"]: m["type"] for m in parsed["media"]}
    assert media_types["word/media/image1.wmf"] == "formula_preview"
    assert media_types["word/media/image2.png"] == "independent_image"
    assert parsed["summary"]["independent_images"] == 1
    assert parsed["summary"]["formula_preview_media_refs"] == 1


def test_parse_docx_summary_omits_blocks():
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="{W_NS}">
  <w:body><w:p><w:r><w:t>hello</w:t></w:r></w:p><w:sectPr/></w:body>
</w:document>"""
    parsed = parse_docx_summary(_build_minimal_docx_zip(document_xml), filename="hello.docx")
    assert parsed["filename"] == "hello.docx"
    assert "blocks" not in parsed
    assert parsed["summary"]["paragraphs"] == 1


@pytest.mark.integration
def test_reference_textbook_statistics_if_available():
    project_root = Path(__file__).resolve().parents[1]
    path = find_reference_docx_in_storage(project_root)
    env_path = os.environ.get("TEXTBOOK_IMPORTER_V3_TEST_DOCX", "").strip()
    if env_path:
        path = Path(env_path)
    if path is None or not path.is_file():
        pytest.skip(
            "Reference DOCX not found in textbook_import/source/vocational/math_B2/ "
            "matching '第一章 1-1 角度的基本性質-課本'"
        )

    parsed = parse_docx_summary(str(path), filename=path.name)
    summary = parsed["summary"]
    print(format_parse_summary_text(parsed))
    diffs = compare_reference_statistics(summary)
    if diffs:
        details = "\n".join(
            f"{item['metric']}: expected {item['expected']}, actual {item['actual']}, diff {item['difference']}"
            for item in diffs
        )
        pytest.fail("Reference textbook statistics mismatch:\n" + details)
