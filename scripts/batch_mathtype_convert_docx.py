#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Windows-only DOCX MathType -> LaTeX converter via Word COM automation."""

from __future__ import annotations

import argparse
import fnmatch
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional

from docx import Document

LATEX_SIGNALS = (
    "$",
    "\\(",
    "\\)",
    "\\[",
    "\\]",
    "\\frac",
    "\\sqrt",
    "\\le",
    "\\ge",
    "\\left",
    "\\right",
    "\\overline",
    "\\triangle",
)

AUTO_MACRO_CANDIDATES = (
    "MathTypeCommands.ConvertEquations",
    "MathTypeCommands.ConvertEquation",
    "ConvertEquations",
    "MTConvertEquations",
    "MathType.ConvertEquations",
)

KEYWORDS = ("mathtype", "convert", "equation", "tex", "latex")

PYWIN32_INSTALL_HINT = (
    "pywin32 is required for Word automation.\n"
    "Please install it in the current venv:\n"
    "python -m pip install pywin32\n"
    "python -m pywin32_postinstall -install"
)

PYWINAUTO_HINT = "pywinauto is optional for UI fallback: python -m pip install pywinauto"


@dataclass(frozen=True)
class ScanResult:
    filename: str
    input_path: Path
    status: str
    reason: str
    output_path: Optional[Path] = None


@dataclass
class FileResult:
    index: int
    filename: str
    status: str
    reason: str
    strategy_used: str
    command_used: str
    output_path: str
    latex_signal_count_before: int
    latex_signal_count_current: int
    latex_signal_count_saved: int


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Batch convert DOCX MathType equations to LaTeX text DOCX.")
    parser.add_argument("--input-folder")
    parser.add_argument("--output-folder")
    parser.add_argument("--pattern", default="*.docx")
    parser.add_argument("--interactive", action="store_true")
    parser.add_argument("--batch", action="store_true")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--limit", type=int, default=None)
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--continue-on-fail", action="store_true")
    parser.add_argument("--report", default=None)

    parser.add_argument("--discover-mathtype", action="store_true")
    parser.add_argument("--auto", action="store_true")
    parser.add_argument("--semi-auto", action="store_true")

    parser.add_argument("--macro-name", default=None)
    parser.add_argument("--command-onaction", default=None)
    parser.add_argument("--command-caption", default=None)
    parser.add_argument("--allow-ui-automation", action="store_true")
    return parser.parse_args()


def check_pywin32_available() -> bool:
    try:
        import win32com.client  # type: ignore # noqa: F401
        return True
    except ImportError:
        return False


def infer_mode(args: argparse.Namespace) -> str:
    if args.discover_mathtype:
        return "discover"
    if args.batch:
        return "batch"
    return "interactive"


def is_already_latex_name(stem: str) -> bool:
    return "latex" in stem.lower()


def build_output_path(input_path: Path, output_folder: Path) -> Path:
    return output_folder / f"{input_path.stem}_Latex.docx"


def has_existing_latex_sibling(path: Path) -> bool:
    sibling_dir = path.parent
    variants = (
        f"{path.stem}_Latex.docx",
        f"{path.stem}_latex.docx",
        f"{path.stem}Latex.docx",
        f"{path.stem}latex.docx",
    )
    return any((sibling_dir / name).exists() for name in variants)


def classify_file(path: Path, output_folder: Path, overwrite: bool) -> ScanResult:
    name = path.name
    lower = name.lower()

    if lower.endswith(".pdf"):
        return ScanResult(name, path, "skipped_pdf", "file_ext_pdf")
    if lower.endswith(".doc") and not lower.endswith(".docx"):
        return ScanResult(name, path, "skipped_non_docx", "file_ext_doc")
    if lower.endswith(".tmp"):
        return ScanResult(name, path, "skipped_non_docx", "file_ext_tmp")
    if name.startswith("~$"):
        return ScanResult(name, path, "skipped_word_temp", "word_temp_file")
    if not lower.endswith(".docx"):
        return ScanResult(name, path, "skipped_non_docx", "not_docx")
    if is_already_latex_name(path.stem):
        return ScanResult(name, path, "skipped_already_latex", "filename_contains_latex")
    if has_existing_latex_sibling(path):
        return ScanResult(name, path, "skipped_existing_latex_sibling", "input_folder_already_has_latex_version")

    output_path = build_output_path(path, output_folder)
    alt_output = output_folder / f"{path.stem}_latex.docx"
    if not overwrite and (output_path.exists() or alt_output.exists()):
        return ScanResult(name, path, "skipped_existing_output", "existing_output", output_path)

    return ScanResult(name, path, "candidate", "ok", output_path)


def scan_input_folder(input_folder: Path, output_folder: Path, pattern: str, overwrite: bool) -> list[ScanResult]:
    results: list[ScanResult] = []
    for path in sorted(input_folder.iterdir(), key=lambda p: p.name.lower()):
        if not path.is_file():
            continue
        if not fnmatch.fnmatch(path.name, pattern):
            continue
        results.append(classify_file(path, output_folder, overwrite))
    return results


def count_latex_signals_in_text(text: str) -> int:
    return sum(text.count(signal) for signal in LATEX_SIGNALS)


def count_latex_signals_docx(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return count_latex_signals_in_text("\n".join(chunks))


def _safe_get(obj: Any, name: str, default: Any = "") -> Any:
    try:
        return getattr(obj, name)
    except Exception:
        return default


def collect_mathtype_diagnostics(app) -> dict[str, Any]:
    out: dict[str, Any] = {
        "macro_attempts": [],
        "commandbar_candidates": [],
        "commandbar_attempts": [],
        "ui_automation_steps": [],
        "errors": [],
        "com_addins": [],
        "addins_templates": [],
    }
    try:
        for addin in app.COMAddIns:
            out["com_addins"].append(
                {
                    "description": str(_safe_get(addin, "Description", "")),
                    "progid": str(_safe_get(addin, "ProgId", "")),
                    "connect": str(_safe_get(addin, "Connect", "")),
                    "guid": str(_safe_get(addin, "Guid", "")),
                }
            )
    except Exception as exc:
        out["errors"].append(f"com_addins_error:{exc}")

    try:
        for addin in app.AddIns:
            out["addins_templates"].append(
                {
                    "name": str(_safe_get(addin, "Name", "")),
                    "path": str(_safe_get(addin, "Path", "")),
                    "installed": str(_safe_get(addin, "Installed", "")),
                }
            )
    except Exception as exc:
        out["errors"].append(f"addins_error:{exc}")

    return out


def score_command_candidate(candidate: dict[str, str]) -> int:
    text = " ".join(
        [
            candidate.get("control_caption", ""),
            candidate.get("control_onaction", ""),
            candidate.get("commandbar_name", ""),
        ]
    ).lower()
    score = 0
    if "mathtype" in text and "convert" in text and "equation" in text:
        score += 20
    if "convert equations" in text:
        score += 10
    if "latex" in text or "tex" in text:
        score += 6
    if candidate.get("control_onaction", ""):
        score += 4
    return score


def discover_commandbar_candidates(app) -> list[dict[str, str]]:
    candidates: list[dict[str, str]] = []
    for bar in app.CommandBars:
        bar_name = str(_safe_get(bar, "Name", ""))
        for ctrl in bar.Controls:
            cap = str(_safe_get(ctrl, "Caption", ""))
            action = str(_safe_get(ctrl, "OnAction", ""))
            cid = str(_safe_get(ctrl, "Id", ""))
            ctype = str(_safe_get(ctrl, "Type", ""))
            probe = f"{bar_name} {cap} {action}".lower()
            if any(k in probe for k in KEYWORDS):
                candidates.append(
                    {
                        "commandbar_name": bar_name,
                        "control_caption": cap,
                        "control_onaction": action,
                        "control_id": cid,
                        "control_type": ctype,
                    }
                )
    return sorted(candidates, key=score_command_candidate, reverse=True)


def get_doc_text(doc) -> str:
    return str(_safe_get(_safe_get(doc, "Content", None), "Text", ""))


def try_macro_strategy(app, args: argparse.Namespace, diagnostics: dict[str, Any]) -> tuple[bool, str]:
    macro_list = [args.macro_name] if args.macro_name else list(AUTO_MACRO_CANDIDATES)
    for macro in macro_list:
        if not macro:
            continue
        diagnostics["macro_attempts"].append(macro)
        try:
            app.Run(macro)
            return True, f"macro:{macro}"
        except Exception as exc:
            diagnostics["errors"].append(f"macro_failed:{macro}:{exc}")
    return False, ""


def try_commandbar_strategy(app, args: argparse.Namespace, diagnostics: dict[str, Any]) -> tuple[bool, str]:
    candidates = discover_commandbar_candidates(app)
    diagnostics["commandbar_candidates"] = candidates

    filtered = candidates
    if args.command_caption:
        filtered = [c for c in candidates if args.command_caption.lower() in c.get("control_caption", "").lower()]

    for cand in filtered:
        caption = cand.get("control_caption", "")
        onaction = cand.get("control_onaction", "")
        diagnostics["commandbar_attempts"].append(f"{caption}|{onaction}")
        try:
            for bar in app.CommandBars:
                for ctrl in bar.Controls:
                    if str(_safe_get(ctrl, "Caption", "")) == caption and str(_safe_get(ctrl, "OnAction", "")) == onaction:
                        ctrl.Execute()
                        return True, f"control:{caption}|{onaction}"
        except Exception as exc:
            diagnostics["errors"].append(f"commandbar_failed:{caption}:{exc}")
    return False, ""


def try_ui_automation_fallback(app, diagnostics: dict[str, Any]) -> tuple[bool, str]:
    diagnostics["ui_automation_steps"].append("start_ui_fallback")
    try:
        try:
            from pywinauto import Application  # type: ignore # noqa: F401
            diagnostics["ui_automation_steps"].append("pywinauto_available")
        except Exception:
            diagnostics["errors"].append(PYWINAUTO_HINT)

        app.Visible = True
        app.Activate()
        diagnostics["ui_automation_steps"].append("activate_word")

        # last resort, only enabled by explicit flag.
        try:
            import win32com.client  # type: ignore

            shell = win32com.client.Dispatch("WScript.Shell")
            shell.AppActivate("Word")
            time.sleep(0.2)
            shell.SendKeys("%")
            diagnostics["ui_automation_steps"].append("sendkeys_placeholder_triggered")
            return True, "ui_sendkeys_fallback"
        except Exception as exc:
            diagnostics["errors"].append(f"ui_sendkeys_failed:{exc}")
            return False, ""
    except Exception as exc:
        diagnostics["errors"].append(f"ui_automation_failed:{exc}")
        return False, ""


def perform_auto_conversion(app, doc, args: argparse.Namespace, output_path: Path, diagnostics: dict[str, Any]) -> tuple[str, str, str, int, int, int, str]:
    before = count_latex_signals_in_text(get_doc_text(doc))

    ok, command_used = try_macro_strategy(app, args, diagnostics)
    strategy_used = "macro"

    if not ok:
        ok, command_used = try_commandbar_strategy(app, args, diagnostics)
        strategy_used = "commandbar"

    if not ok:
        if not args.allow_ui_automation:
            return (
                "failed_ui_automation_not_allowed",
                "auto_strategies_exhausted_without_ui_automation",
                "none",
                before,
                before,
                0,
                strategy_used,
            )
        ok, command_used = try_ui_automation_fallback(app, diagnostics)
        strategy_used = "ui_automation"
        if not ok:
            return ("failed_ui_automation", "ui_automation_failed", command_used, before, before, 0, strategy_used)

    time.sleep(2.0)
    current = count_latex_signals_in_text(get_doc_text(doc))
    if current == 0 or current <= before:
        return (
            "failed_no_latex_signal_after_convert",
            "latex_signal_not_increased_after_convert",
            command_used,
            before,
            current,
            0,
            strategy_used,
        )

    doc.SaveAs2(str(output_path.resolve()), FileFormat=16)
    saved = count_latex_signals_docx(output_path) if output_path.exists() else 0
    if saved <= 0:
        return (
            "failed_save_validation",
            "latex_signal_count_saved_zero",
            command_used,
            before,
            current,
            saved,
            strategy_used,
        )
    return ("converted", "ok", command_used, before, current, saved, strategy_used)


def try_convert_one_auto(input_path: Path, output_path: Path, args: argparse.Namespace) -> tuple[str, str, str, int, int, int, dict[str, Any], str]:
    import win32com.client  # type: ignore

    app = None
    doc = None
    diagnostics: dict[str, Any] = {
        "macro_attempts": [],
        "commandbar_candidates": [],
        "commandbar_attempts": [],
        "ui_automation_steps": [],
        "errors": [],
        "com_addins": [],
        "addins_templates": [],
    }
    try:
        app = win32com.client.Dispatch("Word.Application")
        app.Visible = True
        app.DisplayAlerts = 0
        diagnostics.update(collect_mathtype_diagnostics(app))
        doc = app.Documents.Open(str(input_path.resolve()))

        status, reason, command_used, before, current, saved, strategy_used = perform_auto_conversion(
            app, doc, args, output_path, diagnostics
        )
        if status == "converted":
            return (status, reason, command_used, before, current, saved, diagnostics, strategy_used)

        if status == "failed_ui_automation_not_allowed":
            return (
                "failed_no_auto_command",
                reason,
                command_used,
                before,
                current,
                saved,
                diagnostics,
                strategy_used,
            )
        return (status, reason, command_used, before, current, saved, diagnostics, strategy_used)
    except Exception as exc:
        diagnostics["errors"].append(f"word_runtime_error:{exc}")
        return (
            "failed_no_auto_command",
            f"word_runtime_error:{exc}",
            "",
            0,
            0,
            0,
            diagnostics,
            "none",
        )
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass


def try_convert_one_semi_auto(input_path: Path, output_path: Path) -> tuple[str, str, str, int, int, int, dict[str, Any], str]:
    import win32com.client  # type: ignore

    app = None
    doc = None
    diagnostics = {
        "macro_attempts": [],
        "commandbar_candidates": [],
        "commandbar_attempts": [],
        "ui_automation_steps": [],
        "errors": [],
        "com_addins": [],
        "addins_templates": [],
    }
    try:
        app = win32com.client.Dispatch("Word.Application")
        app.Visible = True
        app.DisplayAlerts = 0
        diagnostics.update(collect_mathtype_diagnostics(app))
        doc = app.Documents.Open(str(input_path.resolve()))

        before = count_latex_signals_in_text(get_doc_text(doc))
        print("請在已開啟的 Word 文件中手動執行 MathType → Convert Equations，完成後按 Enter")
        _ = input()
        current = count_latex_signals_in_text(get_doc_text(doc))

        if current == 0 or current <= before:
            return (
                "failed_no_latex_signal_after_convert",
                "latex_signal_not_increased_after_convert",
                "manual",
                before,
                current,
                0,
                diagnostics,
                "semi_auto",
            )

        doc.SaveAs2(str(output_path.resolve()), FileFormat=16)
        saved = count_latex_signals_docx(output_path) if output_path.exists() else 0
        if saved > 0:
            return ("converted", "ok", "manual", before, current, saved, diagnostics, "semi_auto")
        return (
            "failed_save_validation",
            "latex_signal_count_saved_zero",
            "manual",
            before,
            current,
            saved,
            diagnostics,
            "semi_auto",
        )
    except KeyboardInterrupt:
        return ("cancelled_by_user", "cancelled_by_user", "manual", 0, 0, 0, diagnostics, "semi_auto")
    except Exception as exc:
        diagnostics["errors"].append(f"word_runtime_error:{exc}")
        return ("failed_no_auto_command", f"word_runtime_error:{exc}", "manual", 0, 0, 0, diagnostics, "semi_auto")
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass


def discover_mathtype_only() -> tuple[bool, dict[str, Any]]:
    if not check_pywin32_available():
        return False, {"errors": [PYWIN32_INSTALL_HINT]}
    import win32com.client  # type: ignore

    app = None
    try:
        app = win32com.client.Dispatch("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        diags = collect_mathtype_diagnostics(app)
        diags["commandbar_candidates"] = discover_commandbar_candidates(app)
        return True, diags
    except Exception as exc:
        return True, {"errors": [f"word_runtime_error:{exc}"]}
    finally:
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass


def write_report(
    report_path: Path,
    mode: str,
    scan_results: list[ScanResult],
    file_results: list[FileResult],
    diagnostics: dict[str, Any],
    auto: bool,
    allow_ui_automation: bool,
    selected_file: str,
) -> None:
    converted_count = sum(1 for r in file_results if r.status == "converted")
    failed_count = sum(1 for r in file_results if r.status.startswith("failed"))
    strategy_used = next((r.strategy_used for r in file_results if r.strategy_used), "")
    before = next((r.latex_signal_count_before for r in file_results if r.latex_signal_count_before >= 0), 0)
    current = next((r.latex_signal_count_current for r in file_results if r.latex_signal_count_current >= 0), 0)
    saved = next((r.latex_signal_count_saved for r in file_results if r.latex_signal_count_saved >= 0), 0)

    lines: list[str] = []
    lines.append("# MathType DOCX Batch Convert Report")
    lines.append("")
    lines.append("## Summary")
    lines.append(f"- mode: `{mode}`")
    lines.append(f"- auto: `{str(auto).lower()}`")
    lines.append(f"- allow_ui_automation: `{str(allow_ui_automation).lower()}`")
    lines.append(f"- strategy_used: `{strategy_used}`")
    lines.append(f"- selected_file: `{selected_file}`")
    lines.append(f"- converted_count: `{converted_count}`")
    lines.append(f"- failed_count: `{failed_count}`")
    lines.append(f"- latex_signal_count_before: `{before}`")
    lines.append(f"- latex_signal_count_current: `{current}`")
    lines.append(f"- latex_signal_count_saved: `{saved}`")

    lines.append("")
    lines.append("## Per File")
    lines.append("| index | filename | status | reason | strategy_used | command_used | output_path | latex_signal_count_before | latex_signal_count_current | latex_signal_count_saved |")
    lines.append("|---:|---|---|---|---|---|---|---:|---:|---:|")
    for r in file_results:
        lines.append(
            f"| {r.index} | {r.filename} | {r.status} | {r.reason} | {r.strategy_used} | {r.command_used} | {r.output_path} | {r.latex_signal_count_before} | {r.latex_signal_count_current} | {r.latex_signal_count_saved} |"
        )

    lines.append("")
    lines.append("## Diagnostics")
    for key in ("macro_attempts", "commandbar_candidates", "commandbar_attempts", "ui_automation_steps", "errors"):
        lines.append(f"### {key}")
        value = diagnostics.get(key, [])
        if value:
            for item in value:
                lines.append(f"- `{item}`")
        else:
            lines.append("- (none)")

    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8-sig")


def pick_interactive_candidate(candidates: list[ScanResult]) -> list[ScanResult]:
    print("Selectable DOCX candidates:")
    for i, c in enumerate(candidates, start=1):
        print(f"[{i}] {c.filename}")
    selected = input("Enter index to convert: ").strip()
    if not selected.isdigit():
        raise ValueError("Invalid selection: must be number")
    idx = int(selected)
    if idx < 1 or idx > len(candidates):
        raise ValueError("Invalid selection: out of range")
    return [candidates[idx - 1]]


def main() -> int:
    args = parse_args()
    mode = infer_mode(args)

    if mode == "discover":
        ok, diagnostics = discover_mathtype_only()
        if not ok:
            print(PYWIN32_INSTALL_HINT)
        if args.report:
            write_report(Path(args.report), "discover", [], [], diagnostics, args.auto, args.allow_ui_automation, "")
        return 0

    if not args.input_folder or not args.output_folder:
        print("ERROR: --input-folder and --output-folder are required unless --discover-mathtype is used.", file=sys.stderr)
        return 2

    if not args.dry_run and not check_pywin32_available():
        print(PYWIN32_INSTALL_HINT)

    input_folder = Path(args.input_folder)
    output_folder = Path(args.output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    scan_results = scan_input_folder(input_folder, output_folder, args.pattern, args.overwrite)
    candidates = [r for r in scan_results if r.status == "candidate"]
    if args.limit is not None and args.limit >= 0:
        candidates = candidates[: args.limit]

    selected = pick_interactive_candidate(candidates) if (mode == "interactive" and not args.dry_run and candidates) else candidates

    results: list[FileResult] = []
    diagnostics: dict[str, Any] = {
        "macro_attempts": [],
        "commandbar_candidates": [],
        "commandbar_attempts": [],
        "ui_automation_steps": [],
        "errors": [],
    }
    selected_file = selected[0].filename if (mode == "interactive" and selected) else ""

    idx = 1
    for sr in scan_results:
        if sr.status != "candidate":
            results.append(FileResult(idx, sr.filename, sr.status, sr.reason, "", "", str(sr.output_path) if sr.output_path else "", 0, 0, 0))
            idx += 1

    for sr in selected:
        if args.dry_run:
            results.append(FileResult(idx, sr.filename, "candidate", "dry_run", "", "", str(sr.output_path), 0, 0, 0))
            idx += 1
            continue

        if not check_pywin32_available():
            results.append(FileResult(idx, sr.filename, "failed_pywin32_missing", "pywin32_not_installed", "", "", str(sr.output_path), 0, 0, 0))
            idx += 1
            continue

        if args.semi_auto and not args.auto:
            status, reason, cmd, before, current, saved, diag, strategy = try_convert_one_semi_auto(sr.input_path, sr.output_path)
        else:
            status, reason, cmd, before, current, saved, diag, strategy = try_convert_one_auto(sr.input_path, sr.output_path, args)

        diagnostics = diag
        results.append(
            FileResult(idx, sr.filename, status, reason, strategy, cmd, str(sr.output_path), before, current, saved)
        )
        idx += 1

        if mode == "batch" and status.startswith("failed") and not args.continue_on_fail:
            break

    if args.report:
        write_report(Path(args.report), mode, scan_results, results, diagnostics, args.auto, args.allow_ui_automation, selected_file)
        print(f"Report written: {args.report}")

    converted_count = sum(1 for r in results if r.status == "converted")
    failed_count = sum(1 for r in results if r.status.startswith("failed"))
    print(f"Done. converted={converted_count}, failed={failed_count}, total={len(results)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
