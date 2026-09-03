# -*- coding: utf-8 -*-
"""B2 1-1 PDF visual alignment dry-run (NO DB write).

Reuses:
- core.question_image_assets.render_pdf_page_to_image / question_needs_image
- core.textbook_question_anchor label/type helpers
- PyMuPDF text layer (words/blocks) + get_images/get_drawings

Does NOT:
- write TextbookExample / notes.image_assets / uploads/question_assets
- OCR / pix2tex / Gemini / Phase1-4
- B1 hardcoded page fallback maps
"""

from __future__ import annotations

import json
import re
import sqlite3
import sys
import unicodedata
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from config import Config  # noqa: E402
from core.question_image_assets import (  # noqa: E402
    question_needs_image,
    render_pdf_page_to_image,
    safe_slug,
)
from core.textbook_question_anchor import (  # noqa: E402
    build_question_anchor,
    normalize_question_label,
)

SECTION_DB = "1-1 角度的基本性質"
SECTION_CODE = "1-1"
VOLUME = "數學B2"
CURRICULUM = "vocational"

# Canonical source_order for this section (DOCX / Phase2 authority).
CANONICAL_ORDER: list[str] = [
    "例1",
    "隨堂練習1",
    "例2",
    "隨堂練習2",
    "例3",
    "隨堂練習3",
    "例4",
    "隨堂練習4",
    "108統測B",
    "1-1習題 基礎題 1",
    "1-1習題 基礎題 2",
    "1-1習題 基礎題 3",
    "1-1習題 基礎題 4",
    "1-1習題 基礎題 5",
    "1-1習題 基礎題 6",
    "1-1習題 基礎題 7",
    "1-1習題 基礎題 8",
    "1-1習題 進階題 9",
    "1-1習題 進階題 10",
]

# Primary unique locators verified against B2 1-1 PDF text layer
# (degrees appear as "c"; π often as "r").
LABEL_LOCATORS: dict[str, dict[str, Any]] = {
    "例1": {"must_any": ["試將135c"], "page_hint": [5]},
    "隨堂練習1": {"must_any": ["570c-", "試將570c"], "page_hint": [6]},
    "例2": {"must_any": ["林媽媽製作了"], "page_hint": [7]},
    "隨堂練習2": {"must_any": ["半徑為9公分"], "page_hint": [7]},
    "例3": {"must_any": ["下列何者與50c互為同界角"], "page_hint": [10]},
    "隨堂練習3": {"must_any": ["下列何者與6r互為同界角"], "page_hint": [10]},
    "例4": {"must_any": ["11350c", "1350c21100"], "page_hint": [11]},
    "隨堂練習4": {"must_any": ["1750c21030", "750c21030"], "page_hint": [11]},
    "108統測B": {"must_any": ["2019c", "【108統測B】"], "page_hint": [12]},
    "1-1習題 基礎題1": {"must_any": ["對照表"], "page_hint": [13]},
    "1-1習題 基礎題2": {"must_any": ["化成以弧度為單位"], "page_hint": [13]},
    "1-1習題 基礎題3": {"must_any": ["化成以度為單位"], "page_hint": [13]},
    "1-1習題 基礎題4": {"must_any": ["半徑為12公分"], "page_hint": [13]},
    "1-1習題 基礎題5": {"must_any": ["古典時鐘", "鐘擺長為12公分"], "page_hint": [13]},
    "1-1習題 基礎題6": {"must_any": ["下列何者為60c的同界角"], "page_hint": [13]},
    "1-1習題 基礎題7": {"must_any": ["390c", "1580c"], "page_hint": [13]},
    "1-1習題 基礎題8": {"must_any": ["1415r", "15r2311r"], "page_hint": [14]},
    "1-1習題 進階題9": {"must_any": ["圓形花圃", "農藝科同學"], "page_hint": [14]},
    "1-1習題 進階題10": {"must_any": ["摺扇橋"], "page_hint": [14]},
}

HEADING_PAT = re.compile(
    r"1-1\.\d|1–1\.\d|習題|進階題|基礎題|同界角|扇形|有向角|角的度量"
)


def now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def find_pdf() -> Path:
    folder = ROOT / "textbook_import" / "source" / "vocational" / "math_B2"
    pdfs = [p for p in folder.glob("*.pdf") if "Latex" not in p.name]
    if not pdfs:
        raise FileNotFoundError(f"no PDF in {folder}")
    return pdfs[0]


def normalize_pdf_text(text: str) -> str:
    t = unicodedata.normalize("NFKC", str(text or ""))
    t = t.replace("°", "c").replace("˚", "c").replace("º", "c")
    t = t.replace("−", "-").replace("–", "-").replace("—", "-")
    t = t.replace("π", "r")  # many PDFs render pi-ish as r in text layer
    t = re.sub(r"\s+", "", t)
    return t


def normalize_query_text(text: str) -> str:
    t = unicodedata.normalize("NFKC", str(text or ""))
    t = re.sub(r"\\\(.*?\\\)", " ", t)
    t = re.sub(r"\\\[.*?\\\]", " ", t)
    t = t.replace("°", "c").replace("˚", "c")
    t = re.sub(r"[{}^_\\]", "", t)
    t = re.sub(r"\s+", "", t)
    return t


def locator_for_label(label: str) -> dict[str, Any]:
    lab = normalize_question_label(label)
    if lab in LABEL_LOCATORS:
        return LABEL_LOCATORS[lab]
    # tolerate spaced 基礎題/進階題 forms
    for key, loc in LABEL_LOCATORS.items():
        if normalize_question_label(key) == lab:
            return loc
    return {}


def load_examples() -> list[dict[str, Any]]:
    conn = sqlite3.connect(Config.db_path)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        """
        SELECT id, skill_id, source_description, problem_type, problem_text, notes,
               source_curriculum, source_volume, source_chapter, source_section
        FROM textbook_examples
        WHERE source_curriculum=? AND source_volume=? AND source_section=?
        """,
        (CURRICULUM, VOLUME, SECTION_DB),
    ).fetchall()
    conn.close()
    order_map = {normalize_question_label(x): i + 1 for i, x in enumerate(CANONICAL_ORDER)}
    items: list[dict[str, Any]] = []
    for r in rows:
        notes = {}
        try:
            notes = json.loads(r["notes"] or "{}")
        except Exception:
            notes = {}
        qa = notes.get("question_anchor") if isinstance(notes.get("question_anchor"), dict) else {}
        label = normalize_question_label(str(r["source_description"] or qa.get("question_label") or ""))
        source_order = qa.get("source_order") or order_map.get(label)
        source_type = str(r["problem_type"] or qa.get("source_type") or "")
        if not qa.get("anchor_id"):
            rebuilt = build_question_anchor(
                curriculum=CURRICULUM,
                publisher="longteng",
                volume=VOLUME,
                chapter="1",
                section=SECTION_CODE,
                source_type=source_type,
                question_label=label,
                source_order=int(source_order or 0),
                problem_text=str(r["problem_text"] or ""),
                occurrence_index=1,
            )
            qa = {**rebuilt, "_rebuilt_for_dryrun": True}
        items.append(
            {
                "textbook_example_id": r["id"],
                "skill_id": r["skill_id"],
                "source_description": label,
                "source_type": source_type,
                "source_order": int(source_order or 9999),
                "problem_text": r["problem_text"] or "",
                "anchor": qa,
                "anchor_id": qa.get("anchor_id"),
                "notes_had_anchor": bool(notes.get("question_anchor")),
            }
        )
    items.sort(key=lambda x: (x["source_order"], x["textbook_example_id"]))
    return items


def build_page_index(doc) -> list[dict[str, Any]]:
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text") or ""
        words = page.get_text("words") or []
        blocks = page.get_text("blocks") or []
        images = []
        for info in page.get_images(full=True) or []:
            xref = info[0]
            try:
                rects = page.get_image_rects(xref) or []
            except Exception:
                rects = []
            for rect in rects:
                images.append(
                    {
                        "xref": xref,
                        "bbox": [float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1)],
                        "area": float(abs((rect.x1 - rect.x0) * (rect.y1 - rect.y0))),
                    }
                )
        drawings = []
        try:
            raw_draws = page.get_drawings() or []
        except Exception:
            raw_draws = []
        for d in raw_draws:
            r = d.get("rect")
            if r is None:
                continue
            x0, y0, x1, y1 = float(r.x0), float(r.y0), float(r.x1), float(r.y1)
            area = abs((x1 - x0) * (y1 - y0))
            if area < 80:  # skip tiny strokes
                continue
            drawings.append({"bbox": [x0, y0, x1, y1], "area": area})
        pages.append(
            {
                "page": i + 1,  # 1-based
                "width": float(page.rect.width),
                "height": float(page.rect.height),
                "text": text,
                "norm_text": normalize_pdf_text(text),
                "words": words,
                "blocks": blocks,
                "images": images,
                "drawings": drawings,
                "char_count": len(text),
            }
        )
    return pages


def find_phrase_y(page: dict[str, Any], phrase: str) -> float | None:
    if not phrase:
        return None
    words = page["words"]
    candidates = [phrase]
    for n in (12, 8, 6, 4):
        if len(phrase) > n:
            candidates.append(phrase[:n])
    for target in candidates:
        if not words:
            if target in page["norm_text"]:
                return 0.0
            continue
        for w in words:
            nw = normalize_pdf_text(str(w[4]))
            if target in nw:
                return float(w[1])
        for i in range(len(words)):
            acc = ""
            y0 = float(words[i][1])
            for j in range(i, min(i + 16, len(words))):
                if abs(float(words[j][1]) - y0) > 18:
                    break
                acc += normalize_pdf_text(str(words[j][4]))
                if target in acc:
                    return y0
    return None


def score_match(page: dict[str, Any], phrases: list[str], min_page: int) -> tuple[float, dict[str, Any]]:
    _ = (page, phrases, min_page)
    return 0.0, {"hits": []}


def match_questions(items: list[dict[str, Any]], pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Match using verified unique text-layer anchors + page hints (no OCR)."""
    results = []
    last_page = 3
    last_y = -1.0

    for item in items:
        label = normalize_question_label(item["source_description"])
        loc = locator_for_label(label)
        must_any = [normalize_pdf_text(x) for x in (loc.get("must_any") or [])]
        page_hint = set(int(x) for x in (loc.get("page_hint") or []))

        best = None  # (score, page_no, y, hits, method)
        for page in pages:
            hits = [m for m in must_any if m and m in page["norm_text"]]
            if not hits:
                continue
            # y from longest hit
            y = None
            for h in sorted(hits, key=len, reverse=True):
                y = find_phrase_y(page, h)
                if y is not None:
                    break
            if y is None:
                y = 40.0
            score = 0.92 if len(hits) >= 1 else 0.0
            if len(max(hits, key=len)) >= 8:
                score = 0.97
            if page_hint and page["page"] in page_hint:
                score = min(0.99, score + 0.02)
                method = "unique_phrase+page_hint+order"
            else:
                # outside hint: still allow but lower
                score = min(score, 0.8)
                method = "unique_phrase+order"
            # soft order penalty
            if page["page"] < last_page - 1:
                score -= 0.4
            elif page["page"] == last_page and y + 6 < last_y and label not in ("隨堂練習2", "隨堂練習3", "隨堂練習4"):
                # same-page later questions should be below previous
                if y < last_y - 20:
                    score -= 0.15
            if score <= 0:
                continue
            cand = (score, page["page"], float(y), hits, method)
            if best is None or cand[0] > best[0] or (cand[0] == best[0] and page_hint and cand[1] in page_hint):
                # prefer page_hint on ties
                if best is not None and cand[0] == best[0] and page_hint:
                    if best[1] in page_hint and cand[1] not in page_hint:
                        pass
                    else:
                        best = cand
                else:
                    best = cand

        if best is None:
            results.append(
                {
                    **item,
                    "pdf_match": None,
                    "match_method": "unmatched",
                    "match_score": 0.0,
                    "needs_review": True,
                    "reason": "no_unique_phrase_hit",
                }
            )
            continue

        score, page_no, y, hits, method = best
        last_page, last_y = page_no, y
        results.append(
            {
                **item,
                "pdf_match": {
                    "page": page_no,
                    "question_start_y": y,
                    "hits": hits,
                },
                "match_method": method,
                "match_score": round(float(score), 3),
                "needs_review": bool(score < 0.75),
                "reason": "matched_via_text_layer",
            }
        )
    return results


def intersect_area(a: list[float], b: list[float]) -> float:
    x0 = max(a[0], b[0])
    y0 = max(a[1], b[1])
    x1 = min(a[2], b[2])
    y1 = min(a[3], b[3])
    if x1 <= x0 or y1 <= y0:
        return 0.0
    return float((x1 - x0) * (y1 - y0))


def union_bbox(boxes: list[list[float]]) -> list[float] | None:
    if not boxes:
        return None
    return [
        min(b[0] for b in boxes),
        min(b[1] for b in boxes),
        max(b[2] for b in boxes),
        max(b[3] for b in boxes),
    ]


def assign_regions(matches: list[dict[str, Any]], pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    # Sort by source_order already
    for i, row in enumerate(matches):
        pm = row.get("pdf_match")
        if not pm:
            row["regions"] = []
            row["question_bbox"] = None
            row["cross_page"] = False
            continue
        page_no = int(pm["page"])
        y0 = float(pm.get("question_start_y") or 0)
        page = pages[page_no - 1]
        y1 = page["height"] - 36
        next_row = matches[i + 1] if i + 1 < len(matches) else None
        cross_page_suspected = False
        if next_row and next_row.get("pdf_match"):
            np = next_row["pdf_match"]
            if int(np["page"]) == page_no:
                y1 = max(y0 + 36, float(np.get("question_start_y") or y1) - 4)
            elif int(np["page"]) > page_no:
                # Do NOT auto-ingest intermediate pages (causes visual bleed).
                # Keep current page remainder only; flag for human review if gap > 1 page.
                y1 = page["height"] - 36
                if int(np["page"]) - page_no >= 1:
                    cross_page_suspected = int(np["page"]) - page_no >= 2
        regions = [{"page": page_no, "bbox": [36.0, max(0.0, y0 - 8), page["width"] - 36.0, y1]}]
        for reg in regions:
            p = pages[reg["page"] - 1]
            x0, yy0, x1, yy1 = reg["bbox"]
            reg["bbox"] = [
                max(0.0, min(x0, p["width"])),
                max(0.0, min(yy0, p["height"])),
                max(0.0, min(x1, p["width"])),
                max(0.0, min(yy1, p["height"])),
            ]
        row["regions"] = regions
        row["question_bbox"] = regions[0]["bbox"]
        row["cross_page"] = False
        row["cross_page_suspected"] = cross_page_suspected
        pm["question_bbox"] = row["question_bbox"]
        pm["regions"] = regions
    return matches


def detect_visuals(matches: list[dict[str, Any]], pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    for row in matches:
        regions = row.get("regions") or []
        text_flag = question_needs_image(row.get("problem_text") or "")
        img_hits = []
        draw_hits = []
        for reg in regions:
            page = pages[reg["page"] - 1]
            qb = reg["bbox"]
            qarea = max(1.0, abs((qb[2] - qb[0]) * (qb[3] - qb[1])))
            for im in page["images"]:
                inter = intersect_area(qb, im["bbox"])
                if inter <= 0:
                    continue
                if im["area"] < 1200:
                    continue
                if inter / max(im["area"], 1) < 0.25 and inter / qarea < 0.02:
                    continue
                img_hits.append({"page": reg["page"], **im, "intersect": inter})
            # cluster drawings
            for d in page["drawings"]:
                inter = intersect_area(qb, d["bbox"])
                if inter <= 0:
                    continue
                if d["area"] < 2500:
                    continue
                # ignore full-page-ish frames
                if d["area"] > 0.55 * page["width"] * page["height"]:
                    continue
                draw_hits.append({"page": reg["page"], **d, "intersect": inter})

        visual_required: Any = False
        visual_type = None
        visual_bbox = None
        reason = "text_and_formula_only"
        needs_review = bool(row.get("needs_review"))

        significant_draws = [
            d
            for d in draw_hits
            if d["area"] >= 15000
            and (d["bbox"][2] - d["bbox"][0]) >= 100
            and (d["bbox"][3] - d["bbox"][1]) >= 70
        ]
        significant_imgs = []
        for im in img_hits:
            if im["area"] < 5000:
                continue
            cy = (im["bbox"][1] + im["bbox"][3]) / 2.0
            page_h = pages[im["page"] - 1]["height"]
            if cy < 55 or cy > page_h - 40:
                continue
            for reg in regions:
                if reg["page"] != im["page"]:
                    continue
                qb = reg["bbox"]
                if qb[1] - 10 <= cy <= qb[3] + 10:
                    significant_imgs.append(im)
                    break

        if significant_imgs:
            visual_required = True
            visual_type = "embedded_image"
            visual_bbox = union_bbox([i["bbox"] for i in significant_imgs])
            reason = "embedded_image_in_question_region"
        elif text_flag and significant_draws:
            visual_required = True
            visual_type = "diagram"
            top = sorted(significant_draws, key=lambda d: -d["area"])[:8]
            visual_bbox = union_bbox([d["bbox"] for d in top])
            reason = "figure_keyword_and_vector_drawing"
        elif text_flag:
            # Soften drawing gate when problem text explicitly says 如圖
            soft_draws = [
                d
                for d in draw_hits
                if d["area"] >= 6000
                and (d["bbox"][2] - d["bbox"][0]) >= 60
                and (d["bbox"][3] - d["bbox"][1]) >= 50
            ]
            if soft_draws:
                visual_required = True
                visual_type = "diagram"
                top = sorted(soft_draws, key=lambda d: -d["area"])[:10]
                visual_bbox = union_bbox([d["bbox"] for d in top])
                reason = "figure_keyword_and_soft_vector_cluster"
            else:
                visual_required = "needs_review"
                visual_type = "diagram_suspected"
                reason = "text_says_figure_but_no_clear_image_or_drawing_cluster"
                needs_review = True
        elif significant_draws and not text_flag:
            # In this PDF, formula glyphs are often vector paths — do not auto-claim.
            visual_required = False
            reason = "formula_or_decoration_vectors_only"
        else:
            visual_required = False
            reason = "text_and_formula_only"

        # Cap visual bbox inside question region; if clip fails, demote rather than invent.
        if visual_bbox and regions:
            rb = regions[0]["bbox"]
            inter = intersect_area(visual_bbox, rb)
            if inter < 0.15 * max(
                1.0,
                abs((visual_bbox[2] - visual_bbox[0]) * (visual_bbox[3] - visual_bbox[1])),
            ):
                clipped = [
                    max(visual_bbox[0], rb[0]),
                    max(visual_bbox[1], rb[1]),
                    min(visual_bbox[2], rb[2]),
                    min(visual_bbox[3], rb[3]),
                ]
                if clipped[2] > clipped[0] + 5 and clipped[3] > clipped[1] + 5:
                    visual_bbox = clipped
                else:
                    visual_bbox = None
                    if text_flag:
                        visual_required = "needs_review"
                        reason = "visual_bbox_uncertain"
                        needs_review = True
                    else:
                        visual_required = False
                        reason = "image_outside_reliable_question_band"

        row["visual_required"] = visual_required
        row["visual_type"] = visual_type
        row["visual_bbox"] = visual_bbox
        row["visual_evidence"] = {
            "image_count": len(significant_imgs),
            "drawing_count": len(significant_draws),
            "text_figure_keyword": text_flag,
        }
        row["reason"] = reason
        row["needs_review"] = needs_review or (row.get("match_score") or 0) < 0.75
        if isinstance(visual_required, str):
            row["needs_review"] = True
    return matches


def render_previews(
    matches: list[dict[str, Any]],
    pdf_path: Path,
    out_dir: Path,
    dpi: int = 120,
) -> None:
    import fitz
    from PIL import Image, ImageDraw, ImageFont

    out_dir.mkdir(parents=True, exist_ok=True)
    pages_dir = out_dir / "pages"
    crops_dir = out_dir / "crops"
    pages_dir.mkdir(exist_ok=True)
    crops_dir.mkdir(exist_ok=True)

    # render needed pages once
    needed = sorted({r["pdf_match"]["page"] for r in matches if r.get("pdf_match")})
    page_imgs: dict[int, Path] = {}
    for p in needed:
        dest = pages_dir / f"page_{p:03d}.png"
        render_pdf_page_to_image(str(pdf_path), p - 1, str(dest), dpi=dpi)
        page_imgs[p] = dest

    zoom = dpi / 72.0
    for row in matches:
        pm = row.get("pdf_match")
        if not pm:
            row["preview_crop"] = None
            continue
        page_no = int(pm["page"])
        img_path = page_imgs[page_no]
        # prefer visual bbox crop if required, else question region
        bbox = None
        if row.get("visual_required") is True and row.get("visual_bbox"):
            bbox = row["visual_bbox"]
            tag = "visual"
        else:
            bbox = row.get("question_bbox")
            tag = "question"
        if not bbox:
            row["preview_crop"] = None
            continue
        with Image.open(img_path) as im:
            x0, y0, x1, y1 = [int(v * zoom) for v in bbox]
            x0 = max(0, x0 - 4)
            y0 = max(0, y0 - 4)
            x1 = min(im.width, x1 + 4)
            y1 = min(im.height, y1 + 4)
            if x1 <= x0 + 5 or y1 <= y0 + 5:
                row["preview_crop"] = None
                continue
            crop = im.crop((x0, y0, x1, y1))
            # draw thin border for review
            draw = ImageDraw.Draw(crop)
            draw.rectangle([0, 0, crop.width - 1, crop.height - 1], outline=(220, 40, 40), width=2)
            anchor = safe_slug(row.get("anchor_id") or row.get("source_description") or "q")
            out = crops_dir / f"{anchor}_{tag}_preview.png"
            crop.save(out)
            row["preview_crop"] = str(out.relative_to(ROOT)).replace("\\", "/")
            row["preview_tag"] = tag


def build_contact_sheet(matches: list[dict[str, Any]], out_path: Path) -> Path:
    from PIL import Image, ImageDraw, ImageFont

    cell_w, cell_h = 320, 280
    cols = 4
    rows_n = (len(matches) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * cell_w, rows_n * cell_h), (245, 245, 245))
    draw = ImageDraw.Draw(sheet)
    try:
        font = ImageFont.truetype("arial.ttf", 12)
        font_sm = ImageFont.truetype("arial.ttf", 10)
    except Exception:
        font = ImageFont.load_default()
        font_sm = font

    for idx, row in enumerate(matches):
        r, c = divmod(idx, cols)
        x0, y0 = c * cell_w, r * cell_h
        draw.rectangle([x0, y0, x0 + cell_w - 1, y0 + cell_h - 1], outline=(180, 180, 180))
        label = row.get("source_description") or ""
        page = (row.get("pdf_match") or {}).get("page")
        score = row.get("match_score")
        vr = row.get("visual_required")
        header = f"{row.get('source_order')}. {label}"
        meta = f"p={page} score={score} visual={vr}"
        draw.text((x0 + 6, y0 + 4), header[:42], fill=(20, 20, 20), font=font)
        draw.text((x0 + 6, y0 + 20), meta[:48], fill=(60, 60, 60), font=font_sm)
        preview = row.get("preview_crop")
        if preview:
            try:
                with Image.open(ROOT / preview) as im:
                    im = im.convert("RGB")
                    max_w, max_h = cell_w - 12, cell_h - 48
                    im.thumbnail((max_w, max_h))
                    sheet.paste(im, (x0 + 6, y0 + 40))
            except Exception as exc:
                draw.text((x0 + 6, y0 + 50), f"preview err: {exc}"[:40], fill=(180, 0, 0), font=font_sm)
        else:
            draw.text((x0 + 6, y0 + 50), "(no preview)", fill=(120, 120, 120), font=font_sm)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(out_path)
    return out_path


def mathtype_failure_report(items: list[dict[str, Any]]) -> dict[str, Any]:
    latex = next((ROOT / "textbook_import/source/vocational/math_B2").glob("*_Latex.docx"))
    from docx import Document

    doc = Document(str(latex))
    paras = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.append(cell.text)
    failures = {}
    for idx in (18, 19, 75):
        token = f"[MATH_PARSE_FAILED_{idx}]"
        ctx = next((p for p in paras if token in p), "")
        # classify
        belonging = "教材說明"
        hit_q = None
        for it in items:
            if token in (it.get("problem_text") or ""):
                belonging = "題目"
                hit_q = it.get("source_description")
                break
        # heuristics: if context is conversion exposition / 同界角 definition
        if belonging != "題目":
            if "一周角" in ctx or "由此可得" in ctx:
                belonging = "教材說明"
            elif "最小正同界角" in ctx or "最大負同界角" in ctx:
                belonging = "教材說明"
        failures[f"failure_{idx}"] = {
            "token": token,
            "classification": belonging,
            "question_label": hit_q,
            "context80": (ctx or "")[:120],
        }
    return failures


def to_public_row(row: dict[str, Any]) -> dict[str, Any]:
    pm = row.get("pdf_match")
    pdf_match = None
    if pm:
        pdf_match = {
            "page": pm.get("page"),
            "question_bbox": row.get("question_bbox"),
            "regions": row.get("regions") or [],
            "match_method": row.get("match_method"),
            "match_score": row.get("match_score"),
            "hits": pm.get("hits"),
            "question_start_y": pm.get("question_start_y"),
        }
    return {
        "anchor_id": row.get("anchor_id"),
        "textbook_example_id": row.get("textbook_example_id"),
        "source_order": row.get("source_order"),
        "source_description": row.get("source_description"),
        "source_type": row.get("source_type"),
        "skill_id": row.get("skill_id"),
        "notes_had_anchor": row.get("notes_had_anchor"),
        "problem_text80": (row.get("problem_text") or "")[:80],
        "pdf_match": pdf_match,
        "visual_required": row.get("visual_required"),
        "visual_type": row.get("visual_type"),
        "visual_bbox": row.get("visual_bbox"),
        "visual_evidence": row.get("visual_evidence"),
        "cross_page": bool(row.get("cross_page")),
        "cross_page_suspected": bool(row.get("cross_page_suspected")),
        "needs_review": bool(row.get("needs_review")),
        "reason": row.get("reason"),
        "preview_crop": row.get("preview_crop"),
    }


def main() -> int:
    import fitz

    pdf_path = find_pdf()
    debug_root = ROOT / "textbook_import" / "debug" / "pdf_visual" / "B2_1-1"
    debug_root.mkdir(parents=True, exist_ok=True)

    items = load_examples()
    doc = fitz.open(str(pdf_path))
    pages = build_page_index(doc)
    usable_pages = sum(1 for p in pages if p["char_count"] >= 50)
    text_layer_usable = usable_pages >= max(1, int(0.6 * len(pages)))

    matches = match_questions(items, pages)
    matches = assign_regions(matches, pages)
    matches = detect_visuals(matches, pages)

    # MathType failures → mark question needs_review if applicable
    mt = mathtype_failure_report(items)
    for key, info in mt.items():
        if info.get("classification") == "題目" and info.get("question_label"):
            for row in matches:
                if row.get("source_description") == info["question_label"]:
                    row["needs_review"] = True
                    row["reason"] = f"{row.get('reason')}; mathtype_{key}_in_question"

    render_previews(matches, pdf_path, debug_root, dpi=120)
    contact = debug_root / "contact_sheet.png"
    build_contact_sheet(matches, contact)

    public_rows = [to_public_row(r) for r in matches]
    matched = [r for r in public_rows if r.get("pdf_match")]
    high = [r for r in matched if (r.get("pdf_match") or {}).get("match_score", 0) >= 0.90]
    mid = [r for r in matched if 0.75 <= (r.get("pdf_match") or {}).get("match_score", 0) < 0.90]
    low = [r for r in matched if (r.get("pdf_match") or {}).get("match_score", 0) < 0.75]
    unmatched = [r for r in public_rows if not r.get("pdf_match")]
    needs_review = [r for r in public_rows if r.get("needs_review")]
    visual_true = [r for r in public_rows if r.get("visual_required") is True]
    visual_review = [r for r in public_rows if r.get("visual_required") == "needs_review"]
    cross = [r for r in public_rows if r.get("cross_page")]
    vtypes = Counter(r.get("visual_type") or "none" for r in public_rows if r.get("visual_required") is True)

    summary = {
        "pdf_page_count": len(pages),
        "pdf_text_layer_usable": text_layer_usable,
        "usable_text_pages": usable_pages,
        "matched": len(matched),
        "high_confidence_ge_0_90": len(high),
        "review_recommended_0_75_0_89": len(mid),
        "low_confidence_lt_0_75": len(low),
        "needs_review": len(needs_review),
        "visual_required_true": len(visual_true),
        "visual_required_needs_review": len(visual_review),
        "visual_type_distribution": dict(vtypes),
        "unmatched": len(unmatched),
        "cross_page": len(cross),
        "notes_missing_anchor_count": sum(1 for r in items if not r.get("notes_had_anchor")),
    }

    payload = {
        "status": "ok",
        "generated_at": now_iso(),
        "mode": "pdf_visual_dryrun",
        "db_write": False,
        "pdf_path": str(pdf_path),
        "section": SECTION_DB,
        "summary": summary,
        "mathtype_failures": mt,
        "questions": public_rows,
        "artifacts": {
            "debug_dir": str(debug_root.relative_to(ROOT)).replace("\\", "/"),
            "contact_sheet": str(contact.relative_to(ROOT)).replace("\\", "/"),
            "preview_crops": [r.get("preview_crop") for r in public_rows if r.get("preview_crop")],
        },
        "reuse_notes": {
            "render": "core.question_image_assets.render_pdf_page_to_image",
            "text_figure_heuristic": "core.question_image_assets.question_needs_image",
            "anchor_helpers": "core.textbook_question_anchor",
            "not_used": [
                "scripts/pdf_visual_enrich_imported_section.py OCR/pix2tex",
                "B1 hardcoded page_fallback_map_for_b1_11",
            ],
        },
        "page_stats": [
            {
                "page": p["page"],
                "chars": p["char_count"],
                "images": len(p["images"]),
                "drawings": len(p["drawings"]),
            }
            for p in pages
        ],
    }

    out_json = ROOT / "textbook_import" / "debug" / "B2_1-1_pdf_visual_dryrun.json"
    out_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    summary_txt = ROOT / "textbook_import" / "debug" / "B2_1-1_pdf_visual_dryrun.summary.txt"
    lines = [
        "B2 1-1 PDF visual dry-run",
        f"generated_at: {payload['generated_at']}",
        f"pdf: {pdf_path.name}",
        f"pages: {summary['pdf_page_count']} text_usable={summary['pdf_text_layer_usable']}",
        f"matched: {summary['matched']}/19 high={summary['high_confidence_ge_0_90']} mid={summary['review_recommended_0_75_0_89']} low={summary['low_confidence_lt_0_75']} unmatched={summary['unmatched']}",
        f"needs_review: {summary['needs_review']}",
        f"visual_required: {summary['visual_required_true']} (needs_review visual: {summary['visual_required_needs_review']})",
        f"visual_types: {summary['visual_type_distribution']}",
        f"cross_page: {summary['cross_page']}",
        f"mathtype: {json.dumps(mt, ensure_ascii=False)}",
        f"json: {out_json}",
        f"contact: {contact}",
        "",
        "per-question:",
    ]
    for r in public_rows:
        pm = r.get("pdf_match") or {}
        lines.append(
            f"- {r['source_order']:>2} {r['source_description']:<18} page={pm.get('page')} "
            f"score={pm.get('match_score')} visual={r.get('visual_required')} "
            f"review={r.get('needs_review')} reason={r.get('reason')}"
        )
    summary_txt.write_text("\n".join(lines) + "\n", encoding="utf-8")
    doc.close()
    print(json.dumps({"json": str(out_json), "summary": summary}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
